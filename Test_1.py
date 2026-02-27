import glob
import tkinter as tk
from tkinter import ttk, messagebox, font as tkfont
from tkinter import filedialog
from PIL import Image, ImageTk
from matplotlib import pyplot as plt
import pandas as pd
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import re  
from matplotlib.animation import FuncAnimation
import os
import subprocess
from docx import Document
from docx.shared import Pt, RGBColor
from docx2pdf import convert
import tempfile
import fitz
import logging
from datetime import datetime, time
import webbrowser
import shutil
import os
import uuid
import bcrypt
import shutil
import smtplib
import sqlite3
import datetime
from threading import Timer
from contextlib import contextmanager
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext
from tkinter import font as tkfont
from cryptography.fernet import Fernet

# Variables globales pour stocker les tableaux croisés
pivot_techniques = None
pivot_comptables = None
pivot_compte_41 = None
tableau_listing_police_invalide = None
tableau_listing_valide = None
tableau_listing_police_invalide_comptable = None
tableau_listing_valide_comptable = None
df=None

# ------------------------ CONSTANTS ------------------------
DB_FILE = "admin_system.db"
BACKUP_INTERVAL = 3600  # 1 heure en secondes
PERMISSIONS_LIST = [
    "admin", "user_manage", "content_manage", 
    "settings_manage", "logs_view", "logs_manage"
]
class LoadingScreen:
    """Écran de chargement animé"""
    def __init__(self, root):
        self.root = root
        self.loading_window = tk.Toplevel(root)
        self.loading_window.title("Chargement")
        self.loading_window.geometry("400x300")
        self.loading_window.configure(bg='#140c3d')
        self.loading_window.overrideredirect(True)
        self.loading_window.attributes('-topmost', True)
        
        # Centrer la fenêtre
        self.loading_window.update_idletasks()
        width = 400
        height = 300
        x = (self.loading_window.winfo_screenwidth() // 2) - (width // 2)
        y = (self.loading_window.winfo_screenheight() // 2) - (height // 2)
        self.loading_window.geometry(f'{width}x{height}+{x}+{y}')
        
        # Fond avec dégradé
        self.canvas = tk.Canvas(self.loading_window, width=400, height=300, 
                               bg='#140c3d', highlightthickness=0)
        self.canvas.pack(fill=tk.BOTH, expand=True)
        self.draw_gradient()
        
        # Logo ou icône
        self.canvas.create_text(200, 80, text="📊", 
                               font=('Segoe UI', 48), fill='white')
        
        # Texte principal
        self.canvas.create_text(200, 130, text="Data Analyzer Pro",
                               font=('Segoe UI', 20, 'bold'), fill='white')
        
        # Texte secondaire
        self.canvas.create_text(200, 160, text="Chargement de l'application...",
                               font=('Segoe UI', 10), fill='#bbbbbb')
        
        # Barre de progression animée
        self.progress_frame = tk.Frame(self.canvas, bg='#140c3d')  # Corrigé: couleur hexadécimale
        self.progress_frame.place(x=100, y=200, width=200, height=20)
        
        self.progress_canvas = tk.Canvas(self.progress_frame, width=200, height=20,
                                        bg='#1a1a2e', highlightthickness=0)
        self.progress_canvas.pack()
        
        self.progress_bar = self.progress_canvas.create_rectangle(0, 0, 0, 20,
                                                                fill='#4a148c', outline='')
        
        # Animation de points
        self.dots = self.canvas.create_text(200, 230, text="",
                                          font=('Segoe UI', 14), fill='white')
        self.dot_count = 0
        self.animate_dots()
        
        # Animation de la barre de progression
        self.animate_progress()
        
    def draw_gradient(self):
        """Dessine un dégradé sur le canvas"""
        for i in range(300):
            ratio = i / 300
            r = int(20 * (1 - ratio) + 74 * ratio)
            g = int(12 * (1 - ratio) + 20 * ratio)
            b = int(61 * (1 - ratio) + 46 * ratio)
            color = f'#{r:02x}{g:02x}{b:02x}'
            self.canvas.create_line(0, i, 400, i, fill=color)
    
    def animate_dots(self):
        """Animation des points de chargement"""
        dots = "." * (self.dot_count % 4)
        self.canvas.itemconfig(self.dots, text=f"Chargement{dots}")
        self.dot_count += 1
        if self.loading_window.winfo_exists():
            self.loading_window.after(300, self.animate_dots)
    
    def animate_progress(self):
        """Animation de la barre de progression"""
        current_width = self.progress_canvas.coords(self.progress_bar)[2]
        if current_width < 200:
            new_width = current_width + 2
            self.progress_canvas.coords(self.progress_bar, 0, 0, new_width, 20)
            self.loading_window.after(20, self.animate_progress)
    
    def destroy(self):
        """Ferme l'écran de chargement avec animation"""
        for i in range(10, -1, -1):
            if self.loading_window.winfo_exists():
                self.loading_window.attributes('-alpha', i/10)
                self.loading_window.update()
                time.sleep(0.02)
        self.loading_window.destroy()

# ------------------------ FIRST RUN SETUP ------------------------
class FirstRunSetup:
    def __init__(self, root):
        self.root = root
        self.setup_complete = self._check_setup()
        
        if not self.setup_complete:
            self._show_login_window()
        else:
            self._show_login_window()

    def _check_setup(self):
        """Vérifie si l'admin existe déjà"""
        if not os.path.exists(DB_FILE):
            return False
            
        try:
            with sqlite3.connect(DB_FILE) as conn:
                cursor = conn.cursor()
                cursor.execute("SELECT username FROM users WHERE username='admin'")
                return cursor.fetchone() is not None
        except:
            return False

    def _show_setup_window(self):
        """Affiche la fenêtre de configuration initiale"""
        self.window = tk.Toplevel(self.root)
        self.window.title("Configuration Initiale - Création du compte Admin")
        self.window.geometry("500x400")
        self.window.resizable(False, False)
        self.window.protocol("WM_DELETE_WINDOW", self._prevent_close)
        
        # Style
        title_font = tkfont.Font(family='Segoe UI', size=16, weight='bold')
        label_font = tkfont.Font(family='Segoe UI', size=10)
        
        # Cadre principal
        main_frame = ttk.Frame(self.window)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Titre
        ttk.Label(main_frame, 
                 text="Création du compte administrateur", 
                 font=title_font).pack(pady=(0, 20))
        
        # Formulaire
        form_frame = ttk.Frame(main_frame)
        form_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(form_frame, text="Nom d'utilisateur:", font=label_font).pack(anchor=tk.W)
        self.username_entry = ttk.Entry(form_frame, font=label_font)
        self.username_entry.pack(fill=tk.X, pady=(0, 10))
        self.username_entry.insert(0, "admin")  # Nom d'utilisateur par défaut
        
        ttk.Label(form_frame, text="Mot de passe:", font=label_font).pack(anchor=tk.W)
        self.password_entry = ttk.Entry(form_frame, show="*", font=label_font)
        self.password_entry.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(form_frame, text="Confirmer le mot de passe:", font=label_font).pack(anchor=tk.W)
        self.confirm_entry = ttk.Entry(form_frame, show="*", font=label_font)
        self.confirm_entry.pack(fill=tk.X, pady=(0, 20))
        
        # Bouton
        btn = ttk.Button(main_frame, 
                        text="Créer le compte", 
                        command=self._create_admin_account,
                        style='Accent.TButton')
        btn.pack(fill=tk.X, pady=(10, 0))

    def _prevent_close(self):
        """Empêche la fermeture de la fenêtre sans créer de compte"""
        messagebox.showwarning("Configuration requise", 
                             "Vous devez créer un compte administrateur pour continuer")

    def _create_admin_account(self):
        """Crée le compte administrateur"""
        username = self.username_entry.get()
        password = self.password_entry.get()
        confirm = self.confirm_entry.get()
        
        if not username or not password:
            messagebox.showerror("Erreur", "Tous les champs sont obligatoires")
            return
            
        if password != confirm:
            messagebox.showerror("Erreur", "Les mots de passe ne correspondent pas")
            return
            
        if len(password) < 8:
            messagebox.showerror("Erreur", "Le mot de passe doit contenir au moins 8 caractères")
            return
            
        # Création de la base de données et du compte admin
        try:
            with sqlite3.connect(DB_FILE) as conn:
                cursor = conn.cursor()
                
                # Création des tables
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS users (
                        username TEXT PRIMARY KEY,
                        password TEXT NOT NULL,
                        permissions TEXT,
                        last_login TEXT,
                        status TEXT,
                        role TEXT
                    )
                """)
                
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS logs (
                        id TEXT PRIMARY KEY,
                        timestamp TEXT NOT NULL,
                        action TEXT NOT NULL,
                        details TEXT
                    )
                """)
                
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS history (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        timestamp TEXT NOT NULL,
                        username TEXT NOT NULL,
                        action_type TEXT NOT NULL,
                        target_user TEXT,
                        details TEXT
                    )
                """)
                
                # Insertion de l'admin
                cursor.execute(
                    "INSERT INTO users VALUES (?, ?, ?, ?, ?)",
                    (username, PasswordManager.hash_password(password), "all", "", "active", "admin")
                )
                conn.commit()
                
            messagebox.showinfo("Succès", "Compte administrateur créé avec succès")
            self.setup_complete = True
            self.window.destroy()
            self._show_login_window()
            
        except Exception as e:
            messagebox.showerror("Erreur", f"Échec de la création du compte: {str(e)}")

    def _show_login_window(self):
        """Affiche la fenêtre de connexion"""
        # Cache la fenêtre principale pendant la connexion
        self.root.withdraw()
        LoginWindow(self.root, self._on_login_success)

    def _on_login_success(self, username):
        """Callback après une connexion réussie"""
        # Vérifie le rôle de l'utilisateur
        with DatabaseHandler(DB_FILE).connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT permissions FROM users WHERE username=?",
                (username,)
            )
            result = cursor.fetchone()
        
        if result and "all" in result[0]:  # Si admin
            self.root.deiconify()
            AdminSystem(self.root, username)
        else:  # Si user
            self.root.deiconify()
            DashboardApp(self.root)

# ------------------------ DATABASE HANDLER ------------------------
class DatabaseHandler:
    def __init__(self, db_file):
        self.db_file = db_file
    
    @contextmanager
    def connection(self):
        conn = sqlite3.connect(self.db_file)
        try:
            yield conn
        finally:
            conn.close()

# ------------------------ SECURITY MANAGER ------------------------
class PasswordManager:
    @staticmethod
    def hash_password(password):
        """Hash a password using bcrypt"""
        return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
    
    @staticmethod
    def verify_password(hashed_password, user_password):
        """Verify a password against its hash"""
        return bcrypt.checkpw(user_password.encode('utf-8'), hashed_password.encode('utf-8'))

class BackupManager:
    def __init__(self, db_file, interval=BACKUP_INTERVAL):
        self.db_file = db_file
        self.interval = interval
        self.backup_key = self._get_or_create_key()
        self._timer = None
    
    def _get_or_create_key(self):
        """Generate or retrieve encryption key"""
        key_file = "backup_key.key"
        if os.path.exists(key_file):
            with open(key_file, "rb") as f:
                return f.read()
        else:
            key = Fernet.generate_key()
            with open(key_file, "wb") as f:
                f.write(key)
            return key
    
    def start(self):
        """Start the backup timer"""
        self._timer = Timer(self.interval, self.create_backup)
        self._timer.start()
    
    def create_backup(self):
        """Create an encrypted database backup"""
        backup_dir = "backups"
        os.makedirs(backup_dir, exist_ok=True)
        
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        temp_file = os.path.join(backup_dir, f"temp_backup_{timestamp}.db")
        backup_file = os.path.join(backup_dir, f"backup_{timestamp}.db.enc")
        
        try:
            shutil.copy2(self.db_file, temp_file)
            fernet = Fernet(self.backup_key)
            with open(temp_file, "rb") as f:
                data = f.read()
            encrypted = fernet.encrypt(data)
            with open(backup_file, "wb") as f:
                f.write(encrypted)
            os.remove(temp_file)
            self._clean_old_backups(backup_dir)
        except Exception as e:
            print(f"Backup error: {str(e)}")
        finally:
            self.start()
    
    def _clean_old_backups(self, backup_dir, keep=5):
        """Keep only the most recent backups"""
        backups = sorted(
            [os.path.join(backup_dir, f) for f in os.listdir(backup_dir) 
             if f.startswith("backup_") and f.endswith(".enc")],
            key=os.path.getmtime
        )
        for old_backup in backups[:-keep]:
            os.remove(old_backup)

# ------------------------ NOTIFICATION SERVICE ------------------------
class EmailNotifier:
    def __init__(self, config):
        self.config = config
    
    def send(self, subject, body):
        """Send an email notification"""
        try:
            msg = MIMEMultipart()
            msg['From'] = self.config['email_from']
            msg['To'] = self.config['email_to']
            msg['Subject'] = subject
            msg.attach(MIMEText(body, 'plain'))
            
            with smtplib.SMTP(
                self.config['smtp_server'], 
                self.config['smtp_port']
            ) as server:
                server.starttls()
                server.login(
                    self.config['username'],
                    self.config['password']
                )
                server.send_message(msg)
        except Exception as e:
            print(f"Failed to send email: {str(e)}")

# ------------------------ USER INTERFACE ------------------------
class LoginWindow:
    def __init__(self, root, on_success):
        self.root = root
        self.on_success = on_success
        self.window = tk.Toplevel(root)
        self.window.title("Authentification")
        self.window.geometry("400x400")
        self.window.resizable(False, False)
        
        self._setup_ui()
        self.attempts = 0
        self.locked = False
    
    def _setup_ui(self):
        """Setup login interface"""
        # Style
        title_font = tkfont.Font(family='Segoe UI', size=18, weight='bold')
        label_font = tkfont.Font(family='Segoe UI', size=10)
        
        # Cadre principal
        main_frame = ttk.Frame(self.window)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=30, pady=30)
        
        # Logo et titre
        ttk.Label(main_frame, 
                 text="Connexion", 
                 font=title_font).pack(pady=(0, 30))
        
        # Formulaire
        form_frame = ttk.Frame(main_frame)
        form_frame.pack(fill=tk.X)
        
        self.username_var = tk.StringVar()
        self.password_var = tk.StringVar()
        self.role_var = tk.StringVar()
        
        ttk.Label(form_frame, text="Nom d'utilisateur :", font=label_font).pack(anchor=tk.W)
        ttk.Entry(form_frame, textvariable=self.username_var, font=label_font).pack(fill=tk.X, pady=(0, 15))
        
        ttk.Label(form_frame, text="Mot de passe :", font=label_font).pack(anchor=tk.W)
        ttk.Entry(form_frame, textvariable=self.password_var, show="*", font=label_font).pack(fill=tk.X, pady=(0, 20))

        ttk.Label(form_frame, text="Rôle :", font=label_font).pack(anchor=tk.W)
        ttk.Entry(form_frame, textvariable=self.role_var, font=label_font).pack(fill=tk.X, pady=(0, 15))
        
        # Bouton de connexion
        btn = ttk.Button(main_frame, 
                        text="Se connecter", 
                        command=self.authenticate,
                        style='Accent.TButton')
        btn.pack(fill=tk.X, pady=(10, 0))
    
    # Dans la classe LoginWindow, modifiez la méthode authenticate comme suit :
    def authenticate(self):
        """Authenticate the user"""
        if self.locked:
            messagebox.showerror("Compte verrouillé", "Trop de tentatives. Veuillez réessayer plus tard.")
            return
            
        username = self.username_var.get()
        password = self.password_var.get()
        role = self.role_var.get()  # Récupérer le rôle
        
        with DatabaseHandler(DB_FILE).connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT password, status FROM users WHERE username=?",
                (username,)
            )
            result = cursor.fetchone()
            
        if not result:
            self._handle_failed_attempt()
            return
            
        hashed_pwd, status = result
        if status != "active":
            messagebox.showerror("Erreur", "Ce compte est désactivé")
            return
            
        if not PasswordManager.verify_password(hashed_pwd, password):
            self._handle_failed_attempt()
            return
        
        # Après authentification réussie, vérifier le rôle
        if role == "user":
            # Lancer l'interface utilisateur (DashboardApp)
            self.root.deiconify()
            DashboardApp(self.root)
        elif role == "admin":
            # Lancer l'interface admin (AdminSystem)
            self.on_success(username)
        
        self.window.destroy()
    
    def _handle_failed_attempt(self):
        """Handle a failed login attempt"""
        self.attempts += 1
        messagebox.showerror("Erreur", "Identifiants incorrects")
        if self.attempts >= 3:
            self.locked = True
            messagebox.showerror("Verrouillé", "Trop de tentatives. Compte temporairement verrouillé.")

class AdminSystem:
    def __init__(self, root, username):
        self.root = root
        self.current_user = username
        self.users = {}
        self.logs = []
        
        self._setup_window()
        self._init_services()
        self._setup_ui()
        self.load_data()
        self.log_action("Connexion", f"Utilisateur {username} connecté")
    
    def _setup_window(self):
        """Configure main window"""
        self.root.title(f"Système d'Administration Sécurisé - Connecté en tant que {self.current_user}")
        self.root.geometry("1200x800")
        self.root.minsize(1000, 600)
        
        # Configurer l'icône de l'application
        try:
            self.root.iconbitmap('admin_icon.ico')
        except:
            pass
    
    def _init_services(self):
        """Initialize services"""
        self.db_handler = DatabaseHandler(DB_FILE)
        self.backup_manager = BackupManager(DB_FILE)
        self.backup_manager.start()
        
        self.email_notifier = EmailNotifier({
            'smtp_server': 'smtp.example.com',
            'smtp_port': 587,
            'email_from': 'admin@example.com',
            'email_to': 'admin.alerts@example.com',
            'username': 'admin@example.com',
            'password': 'votre_mot_de_passe'
        })
    
    def _setup_ui(self):
        """Setup user interface"""
        self.status_var = tk.StringVar(value="Prêt")
        
        self._create_main_frame()
        self._create_menu()
        self._create_status_bar()
    
    def _create_main_frame(self):
        """Create main application frame"""
        self.main_frame = ttk.Frame(self.root)
        self.main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.notebook = ttk.Notebook(self.main_frame)
        self.notebook.pack(fill=tk.BOTH, expand=True)
        
        # Add tabs
        self._create_user_tab()
        self._create_permissions_tab()
        self._create_logs_tab()
        self._create_security_tab()
    
    def _create_user_tab(self):
        """Create user management tab"""
        self.user_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.user_frame, text="Gestion des Utilisateurs")
        
        # Barre d'outils
        toolbar = ttk.Frame(self.user_frame)
        toolbar.pack(fill=tk.X, pady=(0, 10))
        
        btn_add = ttk.Button(
            toolbar, 
            text="Ajouter Utilisateur", 
            command=self.show_add_user_dialog,
            style='Accent.TButton'
        )
        btn_add.pack(side=tk.LEFT, padx=(0, 5))
        
        btn_edit = ttk.Button(
            toolbar, 
            text="Modifier", 
            command=self.show_edit_user_dialog
        )
        btn_edit.pack(side=tk.LEFT, padx=5)
        
        btn_delete = ttk.Button(
            toolbar, 
            text="Supprimer", 
            command=self.delete_user
        )
        btn_delete.pack(side=tk.LEFT, padx=5)
        
        # Treeview pour afficher les utilisateurs
        tree_frame = ttk.Frame(self.user_frame)
        tree_frame.pack(fill=tk.BOTH, expand=True)
        
        columns = ("username", "status", "last_login", "permissions")
        self.user_tree = ttk.Treeview(tree_frame, columns=columns, show="headings", selectmode="browse")
        
        # Configurer les colonnes
        self.user_tree.heading("username", text="Nom d'utilisateur")
        self.user_tree.heading("status", text="Statut")
        self.user_tree.heading("last_login", text="Dernière connexion")
        self.user_tree.heading("permissions", text="Permissions")
        
        self.user_tree.column("username", width=200, anchor=tk.W)
        self.user_tree.column("status", width=100, anchor=tk.CENTER)
        self.user_tree.column("last_login", width=150, anchor=tk.CENTER)
        self.user_tree.column("permissions", width=300, anchor=tk.W)
        
        scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self.user_tree.yview)
        self.user_tree.configure(yscrollcommand=scrollbar.set)
        
        self.user_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.user_tree.bind("<Double-1>", lambda e: self.show_edit_user_dialog())
    
    def _create_permissions_tab(self):
        """Create permissions management tab"""
        self.permission_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.permission_frame, text="Gestion des Permissions")
        
        # Section de sélection d'utilisateur
        top_frame = ttk.Frame(self.permission_frame)
        top_frame.pack(fill=tk.X, pady=(0, 15))
        
        ttk.Label(top_frame, text="Utilisateur:", font=('Segoe UI', 10)).pack(side=tk.LEFT, padx=(0, 5))
        
        self.user_var = tk.StringVar()
        self.user_combobox = ttk.Combobox(top_frame, textvariable=self.user_var, state="readonly")
        self.user_combobox.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        
        btn_load = ttk.Button(top_frame, text="Charger", command=self.load_user_permissions)
        btn_load.pack(side=tk.LEFT)
        
        # Section des permissions
        perm_frame = ttk.Frame(self.permission_frame)
        perm_frame.pack(fill=tk.BOTH, expand=True)
        
        # Permissions disponibles
        avail_frame = ttk.Frame(perm_frame)
        avail_frame.grid(row=0, column=0, sticky=tk.NSEW, padx=(0, 10))
        
        ttk.Label(avail_frame, text="Permissions disponibles:", font=('Segoe UI', 10, 'bold')).pack(anchor=tk.W)
        self.available_perms = tk.Listbox(avail_frame, selectmode=tk.MULTIPLE, height=10,
                                        font=('Segoe UI', 10), borderwidth=1, relief=tk.SOLID)
        self.available_perms.pack(fill=tk.BOTH, expand=True)
        
        # Boutons de transfert
        button_frame = ttk.Frame(perm_frame)
        button_frame.grid(row=0, column=1, sticky=tk.NS)
        
        ttk.Button(button_frame, text=">", command=self.add_permission, width=3).pack(pady=5)
        ttk.Button(button_frame, text="<", command=self.remove_permission, width=3).pack(pady=5)
        
        # Permissions de l'utilisateur
        user_frame = ttk.Frame(perm_frame)
        user_frame.grid(row=0, column=2, sticky=tk.NSEW)
        
        ttk.Label(user_frame, text="Permissions de l'utilisateur:", font=('Segoe UI', 10, 'bold')).pack(anchor=tk.W)
        self.user_perms = tk.Listbox(user_frame, selectmode=tk.MULTIPLE, height=10,
                                   font=('Segoe UI', 10), borderwidth=1, relief=tk.SOLID)
        self.user_perms.pack(fill=tk.BOTH, expand=True)
        
        # Bouton de sauvegarde
        btn_save = ttk.Button(perm_frame, 
                            text="Sauvegarder les permissions",
                            command=self.save_user_permissions,
                            style='Accent.TButton')
        btn_save.grid(row=1, column=0, columnspan=3, pady=(15, 0))
        
        # Configuration du grid
        perm_frame.columnconfigure(0, weight=1)
        perm_frame.columnconfigure(2, weight=1)
        perm_frame.rowconfigure(0, weight=1)
        
        # Remplir les permissions disponibles
        for perm in PERMISSIONS_LIST:
            self.available_perms.insert(tk.END, perm)
    
    def _create_logs_tab(self):
        """Create logs management tab"""
        self.log_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.log_frame, text="Journal des Activités")
        
        # Barre de filtres
        filter_frame = ttk.Frame(self.log_frame)
        filter_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(filter_frame, text="Filtrer par:", font=('Segoe UI', 10)).pack(side=tk.LEFT, padx=(0, 5))
        
        self.log_action_var = tk.StringVar()
        self.log_action_combobox = ttk.Combobox(filter_frame, textvariable=self.log_action_var)
        self.log_action_combobox.pack(side=tk.LEFT, padx=5)
        
        ttk.Label(filter_frame, text="Date:", font=('Segoe UI', 10)).pack(side=tk.LEFT, padx=(10, 5))
        self.log_date_var = tk.StringVar()
        self.log_date_combobox = ttk.Combobox(filter_frame, textvariable=self.log_date_var)
        self.log_date_combobox.pack(side=tk.LEFT, padx=5)
        
        btn_filter = ttk.Button(filter_frame, text="Filtrer", command=self.filter_logs)
        btn_filter.pack(side=tk.LEFT, padx=(10, 0))
        
        btn_reset = ttk.Button(filter_frame, text="Réinitialiser", command=self.reset_log_filters)
        btn_reset.pack(side=tk.LEFT, padx=5)
        
        # Zone de texte pour les logs
        self.log_text = scrolledtext.ScrolledText(
            self.log_frame, 
            wrap=tk.WORD, 
            font=('Consolas', 10),
            padx=10,
            pady=10
        )
        self.log_text.pack(fill=tk.BOTH, expand=True)
    
    def _create_security_tab(self):
        """Create security settings tab"""
        self.security_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.security_frame, text="Paramètres de Sécurité")
        
        main_frame = ttk.Frame(self.security_frame)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Politique de mot de passe
        ttk.Label(main_frame, 
                 text="Politique de mot de passe", 
                 font=('Segoe UI', 11, 'bold')).grid(row=0, column=0, sticky=tk.W, pady=(0, 10))
        
        self.pwd_min_length = tk.IntVar(value=8)
        ttk.Label(main_frame, text="Longueur minimale:").grid(row=1, column=0, sticky=tk.W)
        ttk.Spinbox(main_frame, from_=6, to=20, textvariable=self.pwd_min_length).grid(row=1, column=1, sticky=tk.W, pady=5)
        
        self.pwd_require_upper = tk.BooleanVar(value=True)
        ttk.Checkbutton(main_frame, text="Requiert des majuscules", variable=self.pwd_require_upper).grid(row=2, column=0, columnspan=2, sticky=tk.W)
        
        self.pwd_require_special = tk.BooleanVar(value=True)
        ttk.Checkbutton(main_frame, text="Requiert des caractères spéciaux", variable=self.pwd_require_special).grid(row=3, column=0, columnspan=2, sticky=tk.W)
        
        self.pwd_expiry_days = tk.IntVar(value=90)
        ttk.Label(main_frame, text="Expiration (jours):").grid(row=4, column=0, sticky=tk.W)
        ttk.Spinbox(main_frame, from_=30, to=365, textvariable=self.pwd_expiry_days).grid(row=4, column=1, sticky=tk.W, pady=5)
        
        # Verrouillage de compte
        ttk.Label(main_frame, 
                 text="Verrouillage de compte", 
                 font=('Segoe UI', 11, 'bold')).grid(row=5, column=0, sticky=tk.W, pady=(20, 10))
        
        self.lockout_attempts = tk.IntVar(value=5)
        ttk.Label(main_frame, text="Tentatives avant verrouillage:").grid(row=6, column=0, sticky=tk.W)
        ttk.Spinbox(main_frame, from_=3, to=10, textvariable=self.lockout_attempts).grid(row=6, column=1, sticky=tk.W, pady=5)
        
        self.lockout_duration = tk.IntVar(value=30)
        ttk.Label(main_frame, text="Durée de verrouillage (minutes):").grid(row=7, column=0, sticky=tk.W)
        ttk.Spinbox(main_frame, from_=1, to=1440, textvariable=self.lockout_duration).grid(row=7, column=1, sticky=tk.W, pady=5)
        
        # 2FA
        ttk.Label(main_frame, 
                 text="Authentification à deux facteurs (2FA)", 
                 font=('Segoe UI', 11, 'bold')).grid(row=8, column=0, sticky=tk.W, pady=(20, 10))
        
        self.two_factor_enabled = tk.BooleanVar(value=True)
        ttk.Checkbutton(main_frame, text="Activer 2FA pour les administrateurs", variable=self.two_factor_enabled).grid(row=9, column=0, columnspan=2, sticky=tk.W)
        
        # Bouton de sauvegarde
        btn_save = ttk.Button(main_frame, 
                            text="Sauvegarder les paramètres",
                            command=self.save_security_settings,
                            style='Accent.TButton')
        btn_save.grid(row=10, column=0, columnspan=2, pady=(30, 0))
    
    def _create_menu(self):
        """Create menu bar"""
        menubar = tk.Menu(self.root)
        
        # Menu Fichier
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="Actualiser", command=self.refresh_data)
        file_menu.add_separator()
        file_menu.add_command(label="Quitter", command=self.root.quit)
        menubar.add_cascade(label="Fichier", menu=file_menu)
        
        # Menu Aide
        help_menu = tk.Menu(menubar, tearoff=0)
        help_menu.add_command(label="À propos", command=self.show_about)
        menubar.add_cascade(label="Aide", menu=help_menu)
        
        self.root.config(menu=menubar)
    
    def _create_status_bar(self):
        """Create status bar"""
        status_bar = ttk.Frame(self.root, height=25)
        status_bar.pack(fill=tk.X, padx=10, pady=(0, 10))
        
        ttk.Label(status_bar, textvariable=self.status_var).pack(side=tk.LEFT)
        ttk.Label(status_bar, text=f"Utilisateurs: {len(self.users)} | Logs: {len(self.logs)}").pack(side=tk.RIGHT)
    
    def load_data(self):
        """Load data from database"""
        with self.db_handler.connection() as conn:
            cursor = conn.cursor()
            
            cursor.execute("SELECT * FROM users")
            self.users = {
                row[0]: {
                    "password": row[1],
                    "permissions": row[2].split(",") if row[2] else [],
                    "last_login": row[3],
                    "status": row[4]
                } for row in cursor.fetchall()
            }
            
            cursor.execute("SELECT * FROM logs ORDER BY timestamp DESC")
            self.logs = [
                {"id": row[0], "timestamp": row[1], "action": row[2], "details": row[3]}
                for row in cursor.fetchall()
            ]
        
        self.refresh_user_list()
        self.refresh_user_combobox()
        self.display_logs()
    
    def save_data(self):
        """Save data to database"""
        with self.db_handler.connection() as conn:
            cursor = conn.cursor()
            
            for username, data in self.users.items():
                cursor.execute(
                    """INSERT OR REPLACE INTO users VALUES (?, ?, ?, ?, ?)""",
                    (
                        username,
                        data["password"],
                        ",".join(data["permissions"]),
                        data["last_login"],
                        data["status"]
                    )
                )
            
            for log in self.logs:
                cursor.execute(
                    """INSERT OR REPLACE INTO logs VALUES (?, ?, ?, ?)""",
                    (log["id"], log["timestamp"], log["action"], log["details"])
                )
            
            conn.commit()
    
    def log_action(self, action, details=""):
        """Log an action"""
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        log_entry = {
            "id": str(uuid.uuid4()),
            "timestamp": timestamp,
            "action": action,
            "details": details
        }
        self.logs.append(log_entry)
        self.save_data()
    
    def log_history(self, action_type, target_user=None, details=""):
        """Record an action in history"""
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        with self.db_handler.connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                """INSERT INTO history 
                (timestamp, username, action_type, target_user, details)
                VALUES (?, ?, ?, ?, ?)""",
                (timestamp, self.current_user, action_type, target_user, details)
            )
            conn.commit()
    
    def refresh_user_list(self):
        """Refresh user list"""
        self.user_tree.delete(*self.user_tree.get_children())
        
        for i, (username, data) in enumerate(self.users.items()):
            self.user_tree.insert("", tk.END, values=(
                username,
                data["status"],
                data["last_login"] if data["last_login"] else "Jamais",
                ", ".join(data["permissions"]) if data["permissions"] else "Aucune"
            ))
        
        self.status_var.set(f"{len(self.users)} utilisateurs chargés")
    
    def refresh_user_combobox(self):
        """Refresh user combobox"""
        self.user_combobox['values'] = list(self.users.keys())
    
    def load_user_permissions(self):
        """Load user permissions"""
        username = self.user_var.get()
        if not username:
            return
            
        self.user_perms.delete(0, tk.END)
        
        for perm in self.users[username]["permissions"]:
            self.user_perms.insert(tk.END, perm)
    
    def add_permission(self):
        """Add permission to user"""
        username = self.user_var.get()
        if not username:
            messagebox.showwarning("Avertissement", "Veuillez sélectionner un utilisateur")
            return
            
        selected = self.available_perms.curselection()
        if not selected:
            return
            
        for index in selected:
            perm = self.available_perms.get(index)
            if perm not in self.users[username]["permissions"]:
                self.users[username]["permissions"].append(perm)
                self.user_perms.insert(tk.END, perm)
    
    def remove_permission(self):
        """Remove permission from user"""
        username = self.user_var.get()
        if not username:
            messagebox.showwarning("Avertissement", "Veuillez sélectionner un utilisateur")
            return
            
        selected = self.user_perms.curselection()
        if not selected:
            return
            
        for index in reversed(selected):
            perm = self.user_perms.get(index)
            if perm in self.users[username]["permissions"]:
                self.users[username]["permissions"].remove(perm)
            self.user_perms.delete(index)
    
    def save_user_permissions(self):
        """Save user permissions"""
        username = self.user_var.get()
        if not username:
            messagebox.showwarning("Avertissement", "Veuillez sélectionner un utilisateur")
            return
            
        self.save_data()
        self.log_action("Modification permissions", f"Permissions modifiées pour: {username}")
        messagebox.showinfo("Succès", "Permissions sauvegardées avec succès")
        self.refresh_user_list()
    
    def display_logs(self):
        """Display logs"""
        self.log_text.configure(state='normal')
        self.log_text.delete(1.0, tk.END)
        
        for log in self.logs:
            self.log_text.insert(tk.END, f"[{log['timestamp']}] {log['action']}: {log['details']}\n")
        
        self.log_text.configure(state='disabled')
        self.log_text.see(tk.END)
        
        # Update filter comboboxes
        actions = sorted(list(set(log['action'] for log in self.logs)))
        self.log_action_combobox['values'] = ["Toutes"] + actions
        self.log_action_combobox.current(0)
        
        dates = sorted(list(set(log['timestamp'][:10] for log in self.logs)), reverse=True)
        self.log_date_combobox['values'] = ["Toutes"] + dates
        self.log_date_combobox.current(0)
    
    def filter_logs(self):
        """Filter logs"""
        action = self.log_action_var.get()
        date = self.log_date_var.get()
        
        filtered_logs = self.logs
        
        if action != "Toutes":
            filtered_logs = [log for log in filtered_logs if log['action'] == action]
            
        if date != "Toutes":
            filtered_logs = [log for log in filtered_logs if log['timestamp'].startswith(date)]
        
        self.log_text.configure(state='normal')
        self.log_text.delete(1.0, tk.END)
        
        for log in filtered_logs:
            self.log_text.insert(tk.END, f"[{log['timestamp']}] {log['action']}: {log['details']}\n")
        
        self.log_text.configure(state='disabled')
        self.log_text.see(tk.END)
    
    def reset_log_filters(self):
        """Reset log filters"""
        self.log_action_combobox.current(0)
        self.log_date_combobox.current(0)
        self.display_logs()
    
    def save_security_settings(self):
        """Save security settings"""
        messagebox.showinfo("Succès", "Paramètres de sécurité sauvegardés")
        self.log_action("Modification sécurité", "Paramètres de sécurité mis à jour")
    
    def show_add_user_dialog(self):
        """Show add user dialog"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Ajouter un utilisateur")
        dialog.resizable(False, False)
        dialog.grab_set()
        
        # Style
        label_font = tkfont.Font(family='Segoe UI', size=10)
        entry_font = tkfont.Font(family='Segoe UI', size=10)
        
        # Variables
        username = tk.StringVar()
        password = tk.StringVar()
        confirm_password = tk.StringVar()
        status = tk.StringVar(value="active")
        
        # Titre
        ttk.Label(dialog, 
                 text="Ajouter un nouvel utilisateur", 
                 font=('Segoe UI', 12, 'bold')).grid(row=0, column=0, columnspan=2, pady=(0, 15))
        
        # Formulaire
        ttk.Label(dialog, text="Nom d'utilisateur:", font=label_font).grid(row=1, column=0, padx=5, pady=5, sticky=tk.E)
        ttk.Entry(dialog, textvariable=username, font=entry_font).grid(row=1, column=1, padx=5, pady=5, sticky=tk.W)
        
        ttk.Label(dialog, text="Mot de passe:", font=label_font).grid(row=2, column=0, padx=5, pady=5, sticky=tk.E)
        ttk.Entry(dialog, textvariable=password, show="*", font=entry_font).grid(row=2, column=1, padx=5, pady=5, sticky=tk.W)
        
        ttk.Label(dialog, text="Confirmer mot de passe:", font=label_font).grid(row=3, column=0, padx=5, pady=5, sticky=tk.E)
        ttk.Entry(dialog, textvariable=confirm_password, show="*", font=entry_font).grid(row=3, column=1, padx=5, pady=5, sticky=tk.W)
        
        ttk.Label(dialog, text="Statut:", font=label_font).grid(row=4, column=0, padx=5, pady=5, sticky=tk.E)
        ttk.Combobox(dialog, textvariable=status, values=["active", "disabled"], state="readonly", font=entry_font).grid(row=4, column=1, padx=5, pady=5, sticky=tk.W)
        
        # Boutons
        button_frame = ttk.Frame(dialog)
        button_frame.grid(row=5, column=0, columnspan=2, pady=(10, 0))
        
        ttk.Button(button_frame, 
                  text="Ajouter", 
                  command=lambda: self.add_user(
                      username.get(), 
                      password.get(), 
                      confirm_password.get(), 
                      status.get(), 
                      dialog
                  ),
                  style='Accent.TButton').pack(side=tk.LEFT, padx=5)
        
        ttk.Button(button_frame, text="Annuler", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def add_user(self, username, password, confirm_password, status, dialog):
        """Add new user"""
        if not username or not password:
            messagebox.showerror("Erreur", "Tous les champs sont obligatoires")
            return
            
        if password != confirm_password:
            messagebox.showerror("Erreur", "Les mots de passe ne correspondent pas")
            return
            
        if username in self.users:
            messagebox.showerror("Erreur", "Ce nom d'utilisateur existe déjà")
            return
            
        if len(password) < 8:
            messagebox.showerror("Erreur", "Le mot de passe doit contenir au moins 8 caractères")
            return
            
        if not any(c.isupper() for c in password):
            messagebox.showerror("Erreur", "Le mot de passe doit contenir au moins une majuscule")
            return
            
        if not any(c.isdigit() for c in password):
            messagebox.showerror("Erreur", "Le mot de passe doit contenir au moins un chiffre")
            return
            
        self.users[username] = {
            "password": PasswordManager.hash_password(password),
            "permissions": [],
            "last_login": "",
            "status": status
        }
        
        self.save_data()
        self.log_action("Ajout utilisateur", f"Utilisateur ajouté: {username}")
        self.refresh_user_list()
        dialog.destroy()
        messagebox.showinfo("Succès", f"Utilisateur {username} ajouté avec succès")
    
    def show_edit_user_dialog(self):
        """Show edit user dialog"""
        selection = self.user_tree.selection()
        if not selection:
            messagebox.showwarning("Avertissement", "Veuillez sélectionner un utilisateur")
            return
            
        username = self.user_tree.item(selection[0])['values'][0]
        user_data = self.users[username]
        
        dialog = tk.Toplevel(self.root)
        dialog.title(f"Modifier l'utilisateur {username}")
        dialog.resizable(False, False)
        dialog.grab_set()
        
        # Style
        label_font = tkfont.Font(family='Segoe UI', size=10)
        entry_font = tkfont.Font(family='Segoe UI', size=10)
        
        # Variables
        password = tk.StringVar()
        confirm_password = tk.StringVar()
        status = tk.StringVar(value=user_data["status"])
        
        # Titre
        ttk.Label(dialog, 
                 text=f"Modifier l'utilisateur: {username}", 
                 font=('Segoe UI', 12, 'bold')).grid(row=0, column=0, columnspan=2, pady=(0, 15))
        
        # Formulaire
        ttk.Label(dialog, text="Nouveau mot de passe:", font=label_font).grid(row=1, column=0, padx=5, pady=5, sticky=tk.E)
        ttk.Entry(dialog, textvariable=password, show="*", font=entry_font).grid(row=1, column=1, padx=5, pady=5, sticky=tk.W)
        
        ttk.Label(dialog, text="Confirmer mot de passe:", font=label_font).grid(row=2, column=0, padx=5, pady=5, sticky=tk.E)
        ttk.Entry(dialog, textvariable=confirm_password, show="*", font=entry_font).grid(row=2, column=1, padx=5, pady=5, sticky=tk.W)
        
        ttk.Label(dialog, text="Statut:", font=label_font).grid(row=3, column=0, padx=5, pady=5, sticky=tk.E)
        ttk.Combobox(dialog, textvariable=status, values=["active", "disabled"], state="readonly", font=entry_font).grid(row=3, column=1, padx=5, pady=5, sticky=tk.W)
        
        # Boutons
        button_frame = ttk.Frame(dialog)
        button_frame.grid(row=4, column=0, columnspan=2, pady=(10, 0))
        
        ttk.Button(button_frame, 
                  text="Enregistrer", 
                  command=lambda: self.update_user(
                      username, 
                      password.get(), 
                      confirm_password.get(), 
                      status.get(), 
                      dialog
                  ),
                  style='Accent.TButton').pack(side=tk.LEFT, padx=5)
        
        ttk.Button(button_frame, text="Annuler", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def update_user(self, username, password, confirm_password, status, dialog):
        """Update user"""
        if password and password != confirm_password:
            messagebox.showerror("Erreur", "Les mots de passe ne correspondent pas")
            return
            
        if password:
            if len(password) < 8:
                messagebox.showerror("Erreur", "Le mot de passe doit contenir au moins 8 caractères")
                return
                
            if not any(c.isupper() for c in password):
                messagebox.showerror("Erreur", "Le mot de passe doit contenir au moins une majuscule")
                return
                
            if not any(c.isdigit() for c in password):
                messagebox.showerror("Erreur", "Le mot de passe doit contenir au moins un chiffre")
                return
                
            self.users[username]["password"] = PasswordManager.hash_password(password)
            
        self.users[username]["status"] = status
        
        self.save_data()
        self.log_action("Modification utilisateur", f"Utilisateur modifié: {username}")
        self.refresh_user_list()
        dialog.destroy()
        messagebox.showinfo("Succès", f"Utilisateur {username} modifié avec succès")
    
    def delete_user(self):
        """Delete user"""
        selection = self.user_tree.selection()
        if not selection:
            messagebox.showwarning("Avertissement", "Veuillez sélectionner un utilisateur")
            return
            
        username = self.user_tree.item(selection[0])['values'][0]
        
        if username == "admin":
            messagebox.showerror("Erreur", "Impossible de supprimer le compte admin")
            return
            
        if messagebox.askyesno("Confirmation", f"Êtes-vous sûr de vouloir supprimer l'utilisateur {username} ?"):
            del self.users[username]
            self.save_data()
            self.log_action("Suppression utilisateur", f"Utilisateur supprimé: {username}")
            self.refresh_user_list()
            messagebox.showinfo("Succès", f"Utilisateur {username} supprimé avec succès")
    
    def refresh_data(self):
        """Refresh all data"""
        self.load_data()
        self.status_var.set("Données actualisées")
    
    def show_about(self):
        """Show about dialog"""
        about_window = tk.Toplevel(self.root)
        about_window.title("À propos")
        about_window.geometry("400x250")
        about_window.resizable(False, False)
        
        ttk.Label(about_window, 
                 text="Système d'Administration Sécurisé", 
                 font=('Segoe UI', 14, 'bold')).pack(pady=(20, 10))
        
        ttk.Label(about_window, 
                 text="Version 1.0\n\nUn système complet pour gérer les utilisateurs,\nles permissions et surveiller les activités.",
                 font=('Segoe UI', 10)).pack(pady=10)
        
        ttk.Label(about_window, 
                 text="Développé avec Python et Tkinter\n© 2023 - Tous droits réservés",
                 font=('Segoe UI', 9)).pack(pady=(20, 0))
        
        ttk.Button(about_window, 
                  text="Fermer", 
                  command=about_window.destroy,
                  style='Accent.TButton').pack(pady=20)


class GradientFrame(tk.Canvas):
    """Un cadre avec un fond dégradé"""
    def __init__(self, parent, color1="#1e3c72", color2="#2a5298", **kwargs):
        tk.Canvas.__init__(self, parent, **kwargs)
        self._color1 = self._hex_to_rgb(color1)
        self._color2 = self._hex_to_rgb(color2)
        self.bind("<Configure>", self._draw_gradient)
    
    def _hex_to_rgb(self, hex_color):
        """Convertit une couleur hex en tuple RGB"""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        
    def _draw_gradient(self, event=None):
        """Dessine le dégradé"""
        self.delete("gradient")
        width = self.winfo_width()
        height = self.winfo_height()
        
        for i in range(height):
            # Interpolation linéaire entre les deux couleurs
            r = int(self._color1[0] + (self._color2[0] - self._color1[0]) * i / height)
            g = int(self._color1[1] + (self._color2[1] - self._color1[1]) * i / height)
            b = int(self._color1[2] + (self._color2[2] - self._color1[2]) * i / height)
            color = f"#{r:02x}{g:02x}{b:02x}"
            self.create_line(0, i, width, i, tags=("gradient",), fill=color)
            
        self.lower("gradient")

class DropdownMenu(tk.Frame):
    """Menu déroulant personnalisé"""
    def __init__(self, parent, title, items, **kwargs):
        tk.Frame.__init__(self, parent, **kwargs)
        self._is_open = False
        self._items = items
        
        # Bouton principal
        self.main_button = tk.Button(
            self, 
            text=title,
            command=self.toggle,
            bg="#4b8bbe",
            fg="white",
            relief="flat",
            font=("Helvetica", 10, "bold"),
            padx=10,
            pady=5,
            anchor="w"
        )
        self.main_button.pack(fill="x")
        
        # Frame pour les items (initialement cachée)
        self.items_frame = tk.Frame(self, bg="#f0f0f0")
        
    def toggle(self):
        if self._is_open:
            self.items_frame.pack_forget()
        else:
            self.items_frame.pack(fill="x", pady=(5,0))
            # Ajouter les items si pas déjà fait
            if len(self.items_frame.winfo_children()) == 0:
                for text, command in self._items:
                    btn = tk.Button(
                        self.items_frame,
                        text=text,
                        command=lambda c=command: [c(), self.toggle()],
                        bg="#e0e0e0",
                        fg="black",
                        relief="flat",
                        font=("Helvetica", 9),
                        padx=20,
                        pady=3,
                        anchor="w"
                    )
                    btn.pack(fill="x", pady=1)
        self._is_open = not self._is_open

class CertificatApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Générateur de Certificats Pro+")
        self.geometry("1200x800")
        
        # Configuration des logs
        self.setup_logging()
        
        # Variables
        self.word_path = None
        self.excel_path = None
        self.df = None
        self.temp_pdf_path = None
        self.current_preview = None
        self.edited_values = {}  # Stocke les modifications manuelles
        
        # Polices et couleurs
        self.fonts = {
            "system": [],
            "current": "Arial"
        }

        # Charger les polices système d'abord
        self.load_system_fonts()  # <-- Appeler cette méthode avant setup_ui()
    
        self.colors = {
            "texte": "#000000",
            "fond": "#FFFFFF",
            "bouton": "#4CAF50"
        }
        
        # Interface
        self.setup_ui()
        
        # Charger les polices système
        self.load_system_fonts()

        # Dans la méthode __init__, ajoutez :
        self.champs_cotes = ["N° Assuré", "N° Police", "N° Référence", "Intermédiaire", "Tél", "Tél WhatApps", 
                             "Nom(s) et Prénoms", "Date de Naissance", "Sexe", "Effet", "Echéance", "Durée (mois)", 
                             "Fractionnement", "Date de souscription", "Périodicité"]

        self.champs_dessous = ["Garantie", "Capital (FCFA)", "Primes Périodes (FCFA)", 
                               "Prime nette", "Accessoires", "Prime Totale"] 
    
    def setup_logging(self):
        """Configure le système de logs"""
        logging.basicConfig(
            filename='certificat_generator.log',
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        logging.info("Application démarrée")
    
    def load_system_fonts(self):
        """Charge les polices système disponibles"""
        self.fonts["system"] = list(tkfont.families())  # <-- Remplit le tableau des polices
        self.fonts["current"] = "Arial"  # <-- Police par défaut
    
    def setup_ui(self):
        """Configure l'interface utilisateur"""
        # Style
        style = ttk.Style()
        style.configure("TNotebook.Tab", padding=[10, 5])
        style.configure("TButton", padding=6)
        
        # Cadre principal avec fond dégradé
        self.main_frame = GradientFrame(self, color1="#1e3c72", color2="#2a5298")
        self.main_frame.pack(fill="both", expand=True)
        
        # Notebook (onglets)
        self.notebook = ttk.Notebook(self.main_frame)
        self.notebook.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Onglet Certificat
        self.setup_certificate_tab()
        
        # Onglet Excel
        self.setup_excel_tab()
        
        # Onglet Personnalisation
        self.setup_customization_tab()
        
        # Onglet Génération
        self.setup_generation_tab()
        
        # Onglet Logs
        self.setup_logs_tab()
    
    def setup_certificate_tab(self):
        """Configure l'onglet Certificat"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Modèle Certificat")
        
        # Frame supérieure pour les boutons
        btn_frame = ttk.Frame(tab)
        btn_frame.pack(pady=10, fill="x")
        
        # Boutons d'importation
        ttk.Button(
            btn_frame, 
            text="📄 Importer Modèle Word",
            command=self.importer_word
        ).pack(side="left", padx=5)
        
        self.preview_btn = ttk.Button(
            btn_frame,
            text="👁️ Prévisualiser",
            command=self.previsualiser_certificat,
            state="disabled"
        )
        self.preview_btn.pack(side="left", padx=5)
        
        # Frame pour l'affichage
        display_frame = ttk.Frame(tab)
        display_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Canvas pour l'aperçu PDF
        self.canvas_preview = tk.Canvas(
            display_frame, 
            bg="white", 
            width=700, 
            height=500
        )
        self.canvas_preview.pack(fill="both", expand=True)
        
        # Barre de défilement
        scroll_y = ttk.Scrollbar(
            display_frame, 
            orient="vertical", 
            command=self.canvas_preview.yview
        )
        scroll_y.pack(side="right", fill="y")
        self.canvas_preview.configure(yscrollcommand=scroll_y.set)
    
    def setup_excel_tab(self):
        """Configure l'onglet Excel"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Données Excel")
        
        # Bouton d'importation
        ttk.Button(
            tab,
            text="📊 Importer Fichier Excel/CSV",
            command=self.importer_excel
        ).pack(pady=10)
        
        # Treeview avec barres de défilement
        frame = ttk.Frame(tab)
        frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        scroll_y = ttk.Scrollbar(frame)
        scroll_y.pack(side="right", fill="y")
        
        scroll_x = ttk.Scrollbar(frame, orient="horizontal")
        scroll_x.pack(side="bottom", fill="x")
        
        self.tree_excel = ttk.Treeview(
            frame,
            yscrollcommand=scroll_y.set,
            xscrollcommand=scroll_x.set
        )
        self.tree_excel.pack(fill="both", expand=True)
        
        scroll_y.config(command=self.tree_excel.yview)
        scroll_x.config(command=self.tree_excel.xview)
        
        # Menu contextuel pour l'édition
        self.tree_menu = tk.Menu(self, tearoff=0)
        self.tree_menu.add_command(
            label="Modifier cette valeur",
            command=self.edit_tree_value
        )
        self.tree_excel.bind("<Button-3>", self.show_tree_menu)
    
    def setup_customization_tab(self):
        """Configure l'onglet de personnalisation"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Personnalisation")
        
        # Police
        ttk.Label(tab, text="Police du texte:").pack(pady=(10, 0))
        
        self.font_var = tk.StringVar(value="Arial")
        font_dropdown = ttk.Combobox(
            tab,
            textvariable=self.font_var,
            values=self.fonts["system"],
            state="readonly"
        )
        font_dropdown.pack(pady=5)
        
        # Taille de police
        ttk.Label(tab, text="Taille de police:").pack(pady=(10, 0))
        self.font_size = tk.IntVar(value=12)
        ttk.Spinbox(
            tab,
            from_=8,
            to=72,
            textvariable=self.font_size
        ).pack(pady=5)
        
        # Couleur du texte
        ttk.Label(tab, text="Couleur du texte:").pack(pady=(10, 0))
        self.text_color = tk.StringVar(value="#000000")
        ttk.Entry(tab, textvariable=self.text_color).pack(pady=5)
        
        # Alignement
        ttk.Label(tab, text="Alignement:").pack(pady=(10, 0))
        self.alignment = tk.StringVar(value="left")
        ttk.Combobox(
            tab,
            textvariable=self.alignment,
            values=["left", "center", "right"],
            state="readonly"
        ).pack(pady=5)
        
        # Bouton d'application
        ttk.Button(
            tab,
            text="Appliquer les modifications",
            command=self.apply_customization
        ).pack(pady=20)
    
    def setup_generation_tab(self):
        """Configure l'onglet de génération"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Génération")
        
        # Bouton de génération
        self.generate_btn = ttk.Button(
            tab,
            text="🖨️ Générer Tous les Certificats (PDF)",
            command=self.generer_certificats,
            state="disabled"
        )
        self.generate_btn.pack(pady=20)
        
        # Progression
        self.progress = ttk.Progressbar(
            tab,
            orient="horizontal",
            length=400,
            mode="determinate"
        )
        self.progress.pack(pady=10)
        
        # Statut
        self.label_status = ttk.Label(tab, text="")
        self.label_status.pack()
        
        # Bouton d'ouverture du dossier
        ttk.Button(
            tab,
            text="Ouvrir le dossier des certificats",
            command=self.open_output_folder
        ).pack(pady=10)
    
    def setup_logs_tab(self):
        """Configure l'onglet des logs"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Journal")
        
        # Zone de texte pour les logs
        self.log_text = tk.Text(
            tab,
            wrap="word",
            state="disabled",
            bg="black",
            fg="white"
        )
        self.log_text.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Barre de défilement
        scroll = ttk.Scrollbar(
            tab,
            command=self.log_text.yview
        )
        scroll.pack(side="bottom", fill="y")
        self.log_text.config(yscrollcommand=scroll.set)
        
        # Bouton de rafraîchissement
        ttk.Button(
            tab,
            text="Actualiser les logs",
            command=self.refresh_logs
        ).pack(pady=5)
    
    def log_message(self, message, level="info"):
        """Ajoute un message aux logs"""
        self.log_text.config(state="normal")
        self.log_text.insert("end", f"{datetime.now()} - {message}\n")
        self.log_text.config(state="disabled")
        self.log_text.see("end")
        
        if level == "info":
            logging.info(message)
        elif level == "warning":
            logging.warning(message)
        elif level == "error":
            logging.error(message)
    
    def refresh_logs(self):
        """Actualise l'affichage des logs"""
        try:
            with open('certificat_generator.log', 'r') as f:
                content = f.read()
                self.log_text.config(state="normal")
                self.log_text.delete(1.0, "end")
                self.log_text.insert("end", content)
                self.log_text.config(state="disabled")
                self.log_text.see("end")
        except Exception as e:
            self.log_message(f"Erreur lors de la lecture des logs: {str(e)}", "error")
    
    def importer_word(self):
        """Importe un fichier Word comme modèle"""
        filepath = filedialog.askopenfilename(filetypes=[("Word", "*.docx")])
        if filepath:
            self.word_path = filepath
            try:
                # Conversion DOCX -> PDF temporaire
                self.temp_pdf_path = tempfile.mktemp(suffix=".pdf")
                convert(self.word_path, self.temp_pdf_path)
                
                # Affichage du PDF converti
                self.afficher_pdf()
                
                self.preview_btn["state"] = "normal"
                self.log_message(f"Modèle Word importé: {filepath}")
                messagebox.showinfo("Succès", "Modèle Word chargé avec succès !")
            except Exception as e:
                self.log_message(f"Erreur d'importation Word: {str(e)}", "error")
                messagebox.showerror("Erreur", f"Conversion impossible : {e}")
    
    def afficher_pdf(self, page_num=0):
        """Affiche une page spécifique du PDF"""
        if not self.temp_pdf_path:
            return

        try:
            doc = fitz.open(self.temp_pdf_path)
            page = doc.load_page(page_num)
            zoom = 1.5  # Facteur de zoom pour meilleure lisibilité
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            
            # Sauvegarde en image temporaire
            img_path = tempfile.mktemp(suffix=".png")
            pix.save(img_path)
            
            # Affichage dans Tkinter
            img = Image.open(img_path)
            self.current_preview = ImageTk.PhotoImage(img)
            
            # Calculer la taille du canvas
            canvas_width = self.canvas_preview.winfo_width()
            canvas_height = self.canvas_preview.winfo_height()
            
            # Effacer et redessiner
            self.canvas_preview.delete("all")
            self.canvas_preview.config(scrollregion=(0, 0, img.width, img.height))
            self.canvas_preview.create_image(0, 0, anchor="nw", image=self.current_preview)
            
            # Nettoyage
            doc.close()
            os.remove(img_path)
            
            self.log_message(f"Aperçu du PDF généré (page {page_num + 1})")
        except Exception as e:
            self.log_message(f"Erreur d'affichage PDF: {str(e)}", "error")
            messagebox.showerror("Erreur", f"Affichage PDF échoué : {e}")
    
    def previsualiser_certificat(self):
        """Prévisualise le certificat avec les modifications"""
        if not self.word_path:
            return

        try:
            # Créer un document temporaire avec les modifications
            temp_docx = tempfile.mktemp(suffix=".docx")
            doc = Document(self.word_path)
            
            # Appliquer les personnalisations
            self.apply_docx_styles(doc)
            
            # Sauvegarder et convertir
            doc.save(temp_docx)
            convert(temp_docx, self.temp_pdf_path)
            
            # Afficher le résultat
            self.afficher_pdf()
            
            # Nettoyer
            os.remove(temp_docx)
            
            self.log_message("Prévisualisation générée avec les styles appliqués")
        except Exception as e:
            self.log_message(f"Erreur de prévisualisation: {str(e)}", "error")
            messagebox.showerror("Erreur", f"Prévisualisation impossible : {e}")
    
    def apply_docx_styles(self, doc):
        """Version robuste de l'application des styles"""
        try:
            # Convertir la couleur hex en RGB
            try:
                color_hex = self.text_color.get().lstrip('#')
                text_color = RGBColor(*[int(color_hex[i:i+2], 16) for i in (0, 2, 4)])
            except:
                text_color = RGBColor(0, 0, 0)  # Noir par défaut en cas d'erreur
            
            # Appliquer les styles
            for para in doc.paragraphs:
                for run in para.runs:
                    try:
                        run.font.name = self.font_var.get()
                        run.font.size = Pt(min(max(8, self.font_size.get()), 72))  # Borné 8-72
                        run.font.color.rgb = text_color
                    except:
                        continue  # Passe au run suivant en cas d'erreur
        except Exception as e:
            self.log_message(f"Erreur style: {str(e)}", "warning")
    
    def importer_excel(self):
        """Importe un fichier Excel/CSV"""
        filepath = filedialog.askopenfilename(
            filetypes=[("Excel/CSV", "*.xlsx *.xls *.csv")]
        )
        if filepath:
            self.excel_path = filepath
            try:
                ext = os.path.splitext(filepath)[1].lower()
                if ext in (".xlsx", ".xls"):
                    self.df = pd.read_excel(filepath)
                else:
                    self.df = pd.read_csv(filepath)

                # Configurer le Treeview
                self.tree_excel.delete(*self.tree_excel.get_children())
                self.tree_excel["columns"] = list(self.df.columns)
                self.tree_excel["show"] = "headings"
                
                # Configurer les colonnes
                for col in self.df.columns:
                    self.tree_excel.heading(col, text=col)
                    self.tree_excel.column(col, width=100, anchor="center")
                
                # Remplir les données
                for _, row in self.df.iterrows():
                    self.tree_excel.insert("", "end", values=list(row))
                
                self.generate_btn["state"] = "normal"
                self.edited_values = {}  # Réinitialiser les modifications
                
                self.log_message(f"Fichier Excel importé: {filepath}")
                messagebox.showinfo("Succès", "Données Excel chargées !")
            except Exception as e:
                self.log_message(f"Erreur d'importation Excel: {str(e)}", "error")
                messagebox.showerror("Erreur", f"Erreur : {e}")
    
    def show_tree_menu(self, event):
        """Affiche le menu contextuel pour l'arbre Excel"""
        item = self.tree_excel.identify_row(event.y)
        if item:
            self.tree_excel.selection_set(item)
            self.tree_menu.post(event.x_root, event.y_root)
    
    def edit_tree_value(self):
        """Permet d'éditer une valeur dans le Treeview"""
        selected = self.tree_excel.selection()
        if not selected:
            return
            
        item = selected[0]
        column = self.tree_excel.identify_column(self.tree_excel.winfo_pointerx() - self.tree_excel.winfo_rootx())
        col_index = int(column.replace('#', '')) - 1
        
        if col_index < 0 or col_index >= len(self.df.columns):
            return
        
        col_name = self.df.columns[col_index]
        current_value = self.tree_excel.item(item, "values")[col_index]
        
        # Fenêtre d'édition
        edit_win = tk.Toplevel(self)
        edit_win.title(f"Modifier {col_name}")
        
        tk.Label(edit_win, text=f"Nouvelle valeur pour {col_name}:").pack(padx=10, pady=5)
        
        entry = tk.Entry(edit_win, width=30)
        entry.insert(0, current_value)
        entry.pack(padx=10, pady=5)
        
        def save_edit():
            new_value = entry.get()
            values = list(self.tree_excel.item(item, "values"))
            values[col_index] = new_value
            self.tree_excel.item(item, values=values)
            
            # Stocker la modification
            item_id = self.tree_excel.item(item, "text")
            if item_id not in self.edited_values:
                self.edited_values[item_id] = {}
            self.edited_values[item_id][col_name] = new_value
            
            edit_win.destroy()
            self.log_message(f"Valeur modifiée: {col_name}={new_value}")
        
        tk.Button(edit_win, text="Enregistrer", command=save_edit).pack(pady=10)
        entry.focus_set()
        edit_win.transient(self)
        edit_win.grab_set()
        edit_win.wait_window(edit_win)
    
    def apply_customization(self):
        """Applique les personnalisations au document"""
        if not self.word_path:
            messagebox.showwarning("Attention", "Aucun modèle Word chargé !")
            return
            
        try:
            self.previsualiser_certificat()
            messagebox.showinfo("Succès", "Personnalisation appliquée !")
        except Exception as e:
            self.log_message(f"Erreur de personnalisation: {str(e)}", "error")
            messagebox.showerror("Erreur", f"Échec de la personnalisation : {e}")
    
    def generer_certificats(self):
        """Génère tous les certificats avec traitement spécifique des champs et sauvegarde auto-incrémentée"""
        if not all([self.word_path, self.excel_path]):
            messagebox.showwarning("Attention", "Veuillez importer un modèle Word et un fichier Excel")
            return

        output_dir = filedialog.askdirectory(title="Sélectionnez le dossier de sortie pour les PDF")
        if not output_dir:
            return

        try:
            total = len(self.df)
            self.progress["maximum"] = total
            self.label_status.config(text="Génération en cours...")
            self.update_idletasks()  # Meilleure alternative à update()

            # Création du dossier de sortie avec vérification
            os.makedirs(output_dir, exist_ok=True)
            
            # Détection du premier numéro disponible
            existing_files = glob.glob(os.path.join(output_dir, "Certificat_*.pdf"))
            start_num = max([int(f.split("_")[-1].split(".")[0]) for f in existing_files] + [0]) + 1 if existing_files else 1
            
            # Format de numérotation personnalisable
            num_format = "Certificat_{:04d}.pdf"  # Format: 0001, 0002, etc.

            success_count = 0
            error_count = 0
            error_details = []

            for idx, row in self.df.iterrows():
                current_num = start_num + idx
                output_filename = num_format.format(current_num)
                output_path = os.path.join(output_dir, output_filename)
                
                try:
                    # Création du document avec gestion de contexte
                    doc = Document(self.word_path)
                    
                    # Traitement optimisé des tableaux
                    self.process_tables(doc, row)
                    
                    # Sauvegarde temporaire sécurisée
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                        temp_docx = temp_file.name
                        doc.save(temp_docx)
                    
                    # Conversion avec gestion d'erreur
                    try:
                        convert(temp_docx, output_path)
                        success_count += 1
                    except Exception as e:
                        error_msg = f"Erreur conversion ligne {idx+1}: {str(e)}"
                        error_details.append(error_msg)
                        self.log_message(error_msg, "error")
                        
                        # Tentative de conversion alternative
                        if self.fallback_conversion(temp_docx, output_path):
                            success_count += 1
                        else:
                            error_count += 1
                            continue
                    
                    # Mise à jour de la progression
                    self.progress["value"] = idx + 1
                    self.label_status.config(text=f"Traitement {idx+1}/{total} - {success_count} réussis, {error_count} échecs")
                    self.update_idletasks()

                except Exception as e:
                    error_msg = f"Erreur majeure ligne {idx+1}: {str(e)}"
                    error_details.append(error_msg)
                    self.log_message(error_msg, "error")
                    error_count += 1
                    continue
                
                finally:
                    # Nettoyage garantie du fichier temporaire
                    if os.path.exists(temp_docx):
                        try:
                            os.remove(temp_docx)
                        except Exception as e:
                            self.log_message(f"Erreur nettoyage temporaire: {str(e)}", "warning")

            # Rapport final
            report_msg = [
                f"Génération terminée !",
                f"Certificats générés: {success_count}/{total}",
                f"Échecs: {error_count}"
            ]
            
            if error_count > 0:
                report_msg.append("\nDétails des erreurs:")
                report_msg.extend(error_details[:5])  # Affiche les 5 premières erreurs max
                if error_count > 5:
                    report_msg.append(f"...plus {error_count-5} autres erreurs")
                
                log_path = os.path.join(output_dir, "generation_errors.log")
                with open(log_path, "w") as f:
                    f.write("\n".join(error_details))
                report_msg.append(f"\nJournal complet des erreurs sauvegardé dans:\n{log_path}")

            messagebox.showinfo(
                "Rapport de génération",
                "\n".join(report_msg)
            )
            
            # Ouverture automatique du dossier si succès
            if success_count > 0:
                self.open_output_folder(output_dir)

        except Exception as e:
            self.log_message(f"Erreur critique: {str(e)}", "error")
            messagebox.showerror(
                "Erreur critique", 
                f"Le processus a échoué : {str(e)}\n\nVeuillez vérifier les logs pour plus de détails."
            )
        
        finally:
            self.progress["value"] = 0
            self.label_status.config(text="Prêt")

    def process_tables(self, doc, row_data):
        """Traite tous les tableaux du document avec les données"""
        date_fields = ["Date de Naissance", "Date de souscription", "Effet", "Echéance"]
        special_fields = ["N° Référence", "Nom(s) et Prénoms", "Date de Naissance"]
        
        for table in doc.tables:
            for row_idx, table_row in enumerate(table.rows):
                for cell_idx, cell in enumerate(table_row.cells):
                    cell_text = cell.text.strip()
                    
                    for champ in self.champs_cotes + self.champs_dessous:
                        if champ in cell_text and champ in self.df.columns:
                            valeur = self.format_cell_value(champ, str(row_data[champ]))
                            
                            target_cell = self.find_target_cell(
                                table, row_idx, cell_idx, champ, 
                                special_fields, self.champs_cotes
                            )
                            
                            if target_cell:
                                self.update_cell_content(target_cell, valeur)

    def format_cell_value(self, field_name, value):
        """Formate la valeur de la cellule selon son type"""
        date_fields = ["Date de Naissance", "Date de souscription", "Effet", "Echéance"]
        
        if field_name in date_fields:
            try:
                dt = pd.to_datetime(value)
                return dt.strftime("%d-%m-%Y")
            except:
                return value
        return value

    def find_target_cell(self, table, row_idx, cell_idx, field_name, special_fields, cote_fields):
        """Trouve la cellule cible selon les règles de positionnement"""
        if field_name in cote_fields:
            offset = 3 if field_name == "Date de souscription" else (
                2 if field_name in special_fields else 1
            )
            
            if cell_idx + offset < len(table.rows[row_idx].cells):
                return table.rows[row_idx].cells[cell_idx + offset]
        else:
            if row_idx + 1 < len(table.rows):
                next_row = table.rows[row_idx + 1]
                if cell_idx < len(next_row.cells):
                    return next_row.cells[cell_idx]
        return None

    def update_cell_content(self, cell, value):
        """Met à jour le contenu d'une cellule avec le style approprié"""
        # Effacer le contenu existant
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = ""
        
        # Ajouter le nouveau contenu
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(value)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = self.font_var.get()

    def fallback_conversion(self, docx_path, pdf_path):
        """Tentative de conversion alternative avec LibreOffice"""
        try:
            # Solution 1: LibreOffice
            if shutil.which("libreoffice"):
                cmd = [
                    "libreoffice",
                    "--headless",
                    "--convert-to", "pdf",
                    docx_path,
                    "--outdir", os.path.dirname(pdf_path)
                ]
                subprocess.run(cmd, check=True, timeout=60)
                
                # Vérification du résultat
                base_name = os.path.splitext(os.path.basename(docx_path))[0]
                temp_pdf = os.path.join(os.path.dirname(pdf_path), f"{base_name}.pdf")
                
                if os.path.exists(temp_pdf):
                    os.rename(temp_pdf, pdf_path)
                    return True
            
            # Solution 2: Word via COM (Windows seulement)
            try:
                from win32com.client import Dispatch
                word = Dispatch("Word.Application")
                doc = word.Documents.Open(docx_path)
                doc.SaveAs(pdf_path, FileFormat=17)  # 17 = PDF format
                doc.Close()
                word.Quit()
                return True
            except:
                pass
            
            return False
            
        except Exception as e:
            self.log_message(f"Échec conversion alternative: {str(e)}", "error")
            return False

    
    def fallback_conversion(self, docx_path, pdf_path):
        """Méthode alternative si Word échoue"""
        try:
            # Solution 1: Utiliser LibreOffice en ligne de commande
            os.system(f'libreoffice --headless --convert-to pdf "{docx_path}" --outdir "{os.path.dirname(pdf_path)}"')
            
            # Solution 2: Copie de secours si échec
            if not os.path.exists(pdf_path):
                base_name = os.path.splitext(os.path.basename(docx_path))[0]
                fallback_pdf = os.path.join(os.path.dirname(docx_path), f"{base_name}.pdf")
                if os.path.exists(fallback_pdf):
                    os.rename(fallback_pdf, pdf_path)
                else:
                    raise Exception("Échec des méthodes alternatives de conversion")
                
        except Exception as e:
            self.log_message(f"Échec fallback: {str(e)}", "error")
            raise

    def check_prerequisites(self):
        """Vérifie les dépendances nécessaires"""
        missing = []
        
        # Vérifier que Word est installé
        try:
            from win32com.client import Dispatch
            Dispatch("Word.Application").Quit()  # Pensez à fermer Word après vérification
        except Exception as e:
            missing.append("Microsoft Word")
            self.log_message(f"Word non détecté: {str(e)}", "warning")
        
        # Vérifier que LibreOffice est disponible
        if not shutil.which("libreoffice"):
            missing.append("LibreOffice (alternative)")
            self.log_message("LibreOffice non détecté", "warning")
        
        if missing:
            messagebox.showwarning(
                "Dépendances manquantes",
                f"Composants requis manquants:\n- " + "\n- ".join(missing)
            )
        
        return missing
    
    # Alternative si shutil.which n'est pas disponible
    def is_tool_installed(name):
        """Vérifie si un programme est installé (méthode cross-platform)"""
        try:
            from shutil import which
            return which(name) is not None
        except ImportError:
            # Fallback pour anciennes versions de Python
            try:
                subprocess.check_call([name, '--version'], 
                                    stdout=subprocess.DEVNULL, 
                                    stderr=subprocess.DEVNULL)
                return True
            except:
                return False
    
    def open_output_folder(self):
        """Ouvre le dossier de sortie des PDF"""
        if not self.word_path or not self.excel_path:
            return
            
        output_dir = filedialog.askdirectory(title="Sélectionnez le dossier des certificats")
        if output_dir:
            try:
                webbrowser.open(output_dir)
                self.log_message(f"Dossier ouvert: {output_dir}")
            except Exception as e:
                self.log_message(f"Erreur d'ouverture du dossier: {str(e)}", "error")
                messagebox.showerror("Erreur", f"Impossible d'ouvrir le dossier : {e}")


class DataImporterApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Gestion de CP_410 & CP_411 - Outil de Rapprochement")
        self.geometry("1200x700")
        self.configure(bg='#f0f0f0')
        
        # Style configuration
        self.style = ttk.Style()
        #self.style.theme_use('clam')
        
        # Configure styles
        self.style.configure('TNotebook', background='#f0f0f0')
        self.style.configure('TNotebook.Tab', font=('Helvetica', 10, 'bold'), padding=[10, 5])
        self.style.map('TNotebook.Tab', background=[('selected', '#4b8bbe'), ('active', '#5c9ccc')])
        self.style.configure('TFrame', background='#f0f0f0')
        self.style.configure('TButton', font=('Helvetica', 10), padding=5)
        self.style.map('TButton', 
                      foreground=[('active', 'white'), ('!disabled', 'black')],
                      background=[('active', '#4b8bbe'), ('!disabled', '#e0e0e0')])
        self.style.configure('Treeview', font=('Helvetica', 9), rowheight=25)
        self.style.configure('Treeview.Heading', font=('Helvetica', 10, 'bold'))
        self.style.map('Treeview', background=[('selected', '#4b8bbe')])
        
        # Custom colors
        self.primary_color = '#4b8bbe'
        self.secondary_color = '#e0e0e0'
        self.accent_color = '#ff8c00'
        
        # Create main container
        self.main_container = ttk.Frame(self)
        self.main_container.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Create notebook for tabs
        self.tab_control = ttk.Notebook(self.main_container)
        
        # Create tabs
        self.tab_cp410 = ttk.Frame(self.tab_control)
        self.tab_cp411 = ttk.Frame(self.tab_control)
        self.tab_verification = ttk.Frame(self.tab_control)
        self.tab_verification_411 = ttk.Frame(self.tab_control)
        self.tab_rapprochement = ttk.Frame(self.tab_control)
        
        # Add tabs to notebook
        self.tab_control.add(self.tab_cp410, text="CP_410")
        self.tab_control.add(self.tab_cp411, text="CP_411")
        self.tab_control.add(self.tab_verification, text="Vérification 410/411")
        self.tab_control.add(self.tab_verification_411, text="Vérification 411/410")
        self.tab_control.add(self.tab_rapprochement, text="Rapprochement")
        
        self.tab_control.pack(fill='both', expand=True)
        
        # Data storage
        self.cp410_data = None
        self.cp411_data = None
        self.numero_recu_list = []
        
        # Initialize UI components
        self.create_table(self.tab_cp410, "CP_410")
        self.create_table(self.tab_cp411, "CP_411")
        self.create_verification_tab()
        self.create_verification_tab_411()
        self.create_rapprochement_tab()
        
        # Status bar
        self.status_bar = ttk.Label(self, text="Prêt", relief='sunken', anchor='w')
        self.status_bar.pack(side='bottom', fill='x')
        
        # Add menu
        self.create_menu()
    
    def create_menu(self):
        """Create application menu bar"""
        menubar = tk.Menu(self)
        
        # File menu
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="Ouvrir CP_410", command=lambda: self.import_data(self.tab_cp410, "CP_410"))
        file_menu.add_command(label="Ouvrir CP_411", command=lambda: self.import_data(self.tab_cp411, "CP_411"))
        file_menu.add_separator()
        file_menu.add_command(label="Quitter", command=self.quit)
        menubar.add_cascade(label="Fichier", menu=file_menu)
        
        # Tools menu
        tools_menu = tk.Menu(menubar, tearoff=0)
        tools_menu.add_command(label="Vérifier Polices 410/411", command=self.verify_policies)
        tools_menu.add_command(label="Vérifier Polices 411/410", command=self.verify_policies_411)
        tools_menu.add_command(label="Vérifier Références", command=self.verify_ref_piece)
        menubar.add_cascade(label="Outils", menu=tools_menu)
        
        # Help menu
        help_menu = tk.Menu(menubar, tearoff=0)
        help_menu.add_command(label="À propos", command=self.show_about)
        menubar.add_cascade(label="Aide", menu=help_menu)
        
        self.config(menu=menubar)
    
    def show_about(self):
        """Show about dialog"""
        about_window = tk.Toplevel(self)
        about_window.title("À propos")
        about_window.geometry("400x200")
        about_window.resizable(False, False)
        
        about_frame = ttk.Frame(about_window)
        about_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        ttk.Label(about_frame, text="Gestion de CP_410 & CP_411", font=('Helvetica', 14, 'bold')).pack(pady=10)
        ttk.Label(about_frame, text="Outil de rapprochement et vérification").pack()
        ttk.Label(about_frame, text="Version 2.0", font=('Helvetica', 8)).pack(pady=10)
        ttk.Label(about_frame, text="© 2023 - Tous droits réservés").pack()
        
        ttk.Button(about_frame, text="Fermer", command=about_window.destroy).pack(pady=10)
    
    def update_status(self, message):
        """Update status bar message"""
        self.status_bar.config(text=message)
        self.after(5000, lambda: self.status_bar.config(text="Prêt"))
    
    def create_table(self, parent, data_type):
        """Create interface for importing and displaying data with search functionality"""
        main_frame = ttk.Frame(parent)
        main_frame.pack(fill='both', expand=True, padx=5, pady=5)
        
        # Control panel
        control_frame = ttk.Frame(main_frame)
        control_frame.pack(fill='x', pady=(0, 5))
        
        # Import button with icon
        import_btn = ttk.Button(
            control_frame, 
            text=f"Importer {data_type}", 
            command=lambda: self.import_data(parent, data_type),
            style='Accent.TButton'
        )
        import_btn.pack(side='left', padx=(0, 10))
        
        # Search frame
        search_frame = ttk.Frame(control_frame)
        search_frame.pack(side='right', fill='x', expand=True)
        
        ttk.Label(search_frame, text="Rechercher:").pack(side='left', padx=(0, 5))
        search_entry = ttk.Entry(search_frame, width=40)
        search_entry.pack(side='left', fill='x', expand=True)
        search_entry.bind(
            "<KeyRelease>", 
            lambda e, t=parent: self.search_treeview(t.tree, search_entry.get())
        )
        
        # Treeview frame
        tree_frame = ttk.Frame(main_frame)
        tree_frame.pack(fill='both', expand=True)
        
        # Create scrollbars
        tree_scroll_y = ttk.Scrollbar(tree_frame, orient="vertical")
        tree_scroll_x = ttk.Scrollbar(tree_frame, orient="horizontal")
        
        # Create treeview
        tree = ttk.Treeview(
            tree_frame, 
            show="headings", 
            yscrollcommand=tree_scroll_y.set, 
            xscrollcommand=tree_scroll_x.set,
            selectmode='extended'
        )
        
        # Configure scrollbars
        tree_scroll_y.config(command=tree.yview)
        tree_scroll_x.config(command=tree.xview)
        
        # Pack everything
        tree_scroll_y.pack(side="right", fill="y")
        tree_scroll_x.pack(side="bottom", fill="x")
        tree.pack(fill="both", expand=True)
        
        # Store tree reference in parent
        parent.tree = tree
        
        # Configure treeview style
        tree.tag_configure('oddrow', background='#f0f0f0')
        tree.tag_configure('evenrow', background='#ffffff')
    
    def create_verification_tab(self):
        """Create verification tab for 410/411 comparison"""
        main_frame = ttk.Frame(self.tab_verification)
        main_frame.pack(fill='both', expand=True, padx=5, pady=5)
        
        # Control panel
        control_frame = ttk.Frame(main_frame)
        control_frame.pack(fill='x', pady=(0, 5))
        
        verify_btn = ttk.Button(
            control_frame, 
            text="Vérifier les No Police", 
            command=self.verify_policies,
            style='Accent.TButton'
        )
        verify_btn.pack(side='left', padx=(0, 10))
        
        export_btn = ttk.Button(
            control_frame, 
            text="Exporter vers Excel", 
            command=lambda: self.export_to_excel(self.verification_tree, "verification_410_411.xlsx")
        )
        export_btn.pack(side='left')
        
        # Treeview frame
        tree_frame = ttk.Frame(main_frame)
        tree_frame.pack(fill='both', expand=True)
        
        # Create treeview
        columns = ["Police_410_411", "Etat_1", "Police_410_Only", "Etat_2"]
        self.verification_tree = self.create_verification_treeview(tree_frame, columns)
        
        # Configure column colors
        self.verification_tree.tag_configure('match', background='#d4edda')
        self.verification_tree.tag_configure('nomatch', background='#f8d7da')
    
    def create_verification_tab_411(self):
        """Create verification tab for 411/410 comparison"""
        main_frame = ttk.Frame(self.tab_verification_411)
        main_frame.pack(fill='both', expand=True, padx=5, pady=5)
        
        # Control panel
        control_frame = ttk.Frame(main_frame)
        control_frame.pack(fill='x', pady=(0, 5))
        
        verify_btn = ttk.Button(
            control_frame, 
            text="Vérifier les No Police", 
            command=self.verify_policies_411,
            style='Accent.TButton'
        )
        verify_btn.pack(side='left', padx=(0, 10))
        
        export_btn = ttk.Button(
            control_frame, 
            text="Exporter vers Excel", 
            command=lambda: self.export_to_excel(self.verification_tree_411, "verification_411_410.xlsx")
        )
        export_btn.pack(side='left')
        
        # Treeview frame
        tree_frame = ttk.Frame(main_frame)
        tree_frame.pack(fill='both', expand=True)
        
        # Create treeview
        columns = ["Police_411_410", "Etat_1", "Police_411_Only", "Etat_2"]
        self.verification_tree_411 = self.create_verification_treeview(tree_frame, columns)
        
        # Configure column colors
        self.verification_tree_411.tag_configure('match', background='#d4edda')
        self.verification_tree_411.tag_configure('nomatch', background='#f8d7da')
    
    def create_rapprochement_tab(self):
        """Create reconciliation tab"""
        main_frame = ttk.Frame(self.tab_rapprochement)
        main_frame.pack(fill='both', expand=True, padx=5, pady=5)
        
        # Control panel
        control_frame = ttk.Frame(main_frame)
        control_frame.pack(fill='x', pady=(0, 5))
        
        verify_ref_btn = ttk.Button(
            control_frame, 
            text="Vérifier Réf Pièce", 
            command=self.verify_ref_piece,
            style='Accent.TButton'
        )
        verify_ref_btn.pack(side='left', padx=(0, 10))
        
        find_police_btn = ttk.Button(
            control_frame, 
            text="Trouver Polices Associées", 
            command=self.find_policies_for_recu
        )
        find_police_btn.pack(side='left', padx=(0, 10))
        
        export_btn = ttk.Button(
            control_frame, 
            text="Exporter vers Excel", 
            command=lambda: self.export_to_excel(self.rapprochement_tree, "rapprochement.xlsx")
        )
        export_btn.pack(side='left')
        
        # Treeview frame
        tree_frame = ttk.Frame(main_frame)
        tree_frame.pack(fill='both', expand=True)
        
        # Create treeview
        columns = ["Numero_reçu", "Police_Associée"]
        self.rapprochement_tree = self.create_verification_treeview(tree_frame, columns)
    
    def create_verification_treeview(self, parent, columns):
        """Create a styled Treeview for verification results"""
        # Create scrollbars
        tree_scroll_y = ttk.Scrollbar(parent, orient="vertical")
        tree_scroll_x = ttk.Scrollbar(parent, orient="horizontal")
        
        # Create treeview
        tree = ttk.Treeview(
            parent, 
            columns=columns, 
            show="headings", 
            yscrollcommand=tree_scroll_y.set, 
            xscrollcommand=tree_scroll_x.set,
            selectmode='extended'
        )
        
        # Configure scrollbars
        tree_scroll_y.config(command=tree.yview)
        tree_scroll_x.config(command=tree.xview)
        
        # Pack everything
        tree_scroll_y.pack(side="right", fill="y")
        tree_scroll_x.pack(side="bottom", fill="x")
        tree.pack(fill="both", expand=True)
        
        # Configure columns
        for col in columns:
            tree.heading(col, text=col)
            tree.column(col, width=200, anchor='center')
        
        # Configure tags for alternating row colors
        tree.tag_configure('oddrow', background='#f0f0f0')
        tree.tag_configure('evenrow', background='#ffffff')
        
        return tree
    
    def search_treeview(self, tree, query):
        """Filter table content based on search query"""
        items = tree.get_children()
        for item in items:
            tree.delete(item)
        
        if not query:
            return
        
        data = self.cp410_data if tree.master.master == self.tab_cp410 else self.cp411_data
        
        if data is not None:
            try:
                filtered_data = data[data.astype(str).apply(
                    lambda row: row.str.contains(query, case=False, na=False)
                ).any(axis=1)]
                
                for i, (_, row) in enumerate(filtered_data.iterrows()):
                    tag = 'evenrow' if i % 2 == 0 else 'oddrow'
                    tree.insert("", "end", values=list(row), tags=(tag,))
            except Exception as e:
                self.update_status(f"Erreur de recherche: {str(e)}")
    
    def import_data(self, parent, data_type):
        """Import CSV or Excel file and display it"""
        file_path = filedialog.askopenfilename(
            title=f"Importer {data_type}",
            filetypes=[
                ("Fichiers CSV", "*.csv"),
                ("Fichiers Excel", "*.xlsx;*.xls"),
                ("Tous les fichiers", "*.*")
            ]
        )
        
        if not file_path:
            return
        
        try:
            self.update_status(f"Importation du fichier {data_type}...")
            
            # Read file based on extension
            if file_path.endswith(".csv"):
                df = pd.read_csv(file_path, encoding='utf-8', dtype=str)
            else:  # Excel file
                df = pd.read_excel(file_path, dtype=str)
            
            if df.empty:
                messagebox.showwarning("Avertissement", f"Le fichier {data_type} est vide.")
                return
            
            # Clean column names (remove extra spaces)
            df.columns = df.columns.str.strip()
            
            # Store data
            if data_type == "CP_410":
                self.cp410_data = df
            else:
                self.cp411_data = df
            
            # Update treeview
            self.update_treeview(parent.tree, df)
            
            self.update_status(f"{data_type} importé avec succès: {len(df)} enregistrements")
            
        except Exception as e:
            messagebox.showerror(
                "Erreur d'importation", 
                f"Impossible d'importer le fichier:\n{str(e)}"
            )
            self.update_status(f"Erreur d'importation {data_type}")
    
    def update_treeview(self, tree, df):
        """Update treeview with new data"""
        # Clear existing data
        tree.delete(*tree.get_children())
        
        # Set new columns
        tree["columns"] = list(df.columns)
        
        # Configure columns
        for col in df.columns:
            tree.heading(col, text=col)
            tree.column(col, width=150, anchor='w')
        
        # Insert data with alternating row colors
        for i, (_, row) in enumerate(df.iterrows()):
            tag = 'evenrow' if i % 2 == 0 else 'oddrow'
            tree.insert("", "end", values=list(row), tags=(tag,))
    
    def export_to_excel(self, tree, filename):
        """Export treeview data to Excel file"""
        if not tree.get_children():
            messagebox.showwarning("Exportation", "Aucune donnée à exporter.")
            return
        
        # Get data from treeview
        data = []
        columns = tree["columns"]
        
        for item in tree.get_children():
            data.append(tree.item(item)["values"])
        
        # Create DataFrame
        df = pd.DataFrame(data, columns=columns)
        
        # Ask for save location
        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            initialfile=filename,
            filetypes=[("Fichier Excel", "*.xlsx"), ("Tous les fichiers", "*.*")]
        )
        
        if file_path:
            try:
                df.to_excel(file_path, index=False)
                messagebox.showinfo(
                    "Exportation réussie", 
                    f"Les données ont été exportées avec succès:\n{file_path}"
                )
                self.update_status(f"Exportation réussie: {file_path}")
            except Exception as e:
                messagebox.showerror(
                    "Erreur d'exportation", 
                    f"Impossible d'exporter le fichier:\n{str(e)}"
                )
                self.update_status(f"Erreur d'exportation")
    
    def verify_policies(self):
        """Compare policy numbers between CP_410 and CP_411"""
        try:
            self.update_status("Vérification des polices 410/411...")
            
            if self.cp410_data is None or self.cp411_data is None:
                messagebox.showwarning(
                    "Données manquantes", 
                    "Veuillez importer les deux fichiers (CP_410 et CP_411) avant de vérifier."
                )
                return
            
            if "No Police" not in self.cp410_data.columns or "No Police" not in self.cp411_data.columns:
                messagebox.showwarning(
                    "Colonne manquante", 
                    "Les fichiers doivent contenir une colonne 'No Police'."
                )
                return
            
            # Get policy numbers
            cp410_policies = set(self.cp410_data["No Police"].dropna().astype(str).str.strip())
            cp411_policies = set(self.cp411_data["No Police"].dropna().astype(str).str.strip())
            
            # Find matches and differences
            common_policies = sorted(cp410_policies.intersection(cp411_policies))
            only_in_cp410 = sorted(cp410_policies - cp411_policies)
            
            # Prepare data for display
            max_rows = max(len(common_policies), len(only_in_cp410))
            
            # Clear treeview
            self.verification_tree.delete(*self.verification_tree.get_children())
            
            # Insert data with appropriate tags for coloring
            for i in range(max_rows):
                values = [
                    common_policies[i] if i < len(common_policies) else "",
                    "Police retrouvée dans 411" if i < len(common_policies) else "",
                    only_in_cp410[i] if i < len(only_in_cp410) else "",
                    "Police non retrouvée dans 411" if i < len(only_in_cp410) else ""
                ]
                
                tags = ('match',) if i < len(common_policies) else ('nomatch',)
                self.verification_tree.insert("", "end", values=values, tags=tags)
            
            # Update statistics
            stats = (
                f"Total CP_410: {len(cp410_policies)} | "
                f"Total CP_411: {len(cp411_policies)} | "
                f"Correspondances: {len(common_policies)} | "
                f"Différences: {len(only_in_cp410)}"
            )
            self.update_status(f"Vérification terminée. {stats}")
            
            messagebox.showinfo(
                "Vérification terminée", 
                f"Résultats de la vérification:\n\n"
                f"Polices communes: {len(common_policies)}\n"
                f"Polices uniquement dans CP_410: {len(only_in_cp410)}"
            )
            
        except Exception as e:
            messagebox.showerror(
                "Erreur de vérification", 
                f"Une erreur s'est produite lors de la vérification:\n{str(e)}"
            )
            self.update_status("Erreur lors de la vérification")
    
    def verify_policies_411(self):
        """Compare policy numbers between CP_411 and CP_410"""
        try:
            self.update_status("Vérification des polices 411/410...")
            
            if self.cp411_data is None or self.cp410_data is None:
                messagebox.showwarning(
                    "Données manquantes", 
                    "Veuillez importer les deux fichiers (CP_410 et CP_411) avant de vérifier."
                )
                return
            
            if "No Police" not in self.cp411_data.columns or "No Police" not in self.cp410_data.columns:
                messagebox.showwarning(
                    "Colonne manquante", 
                    "Les fichiers doivent contenir une colonne 'No Police'."
                )
                return
            
            # Get policy numbers
            cp411_policies = set(self.cp411_data["No Police"].dropna().astype(str).str.strip())
            cp410_policies = set(self.cp410_data["No Police"].dropna().astype(str).str.strip())
            
            # Find matches and differences
            common_policies = sorted(cp411_policies.intersection(cp410_policies))
            only_in_cp411 = sorted(cp411_policies - cp410_policies)
            
            # Prepare data for display
            max_rows = max(len(common_policies), len(only_in_cp411))
            
            # Clear treeview
            self.verification_tree_411.delete(*self.verification_tree_411.get_children())
            
            # Insert data with appropriate tags for coloring
            for i in range(max_rows):
                values = [
                    common_policies[i] if i < len(common_policies) else "",
                    "Police retrouvée dans 410" if i < len(common_policies) else "",
                    only_in_cp411[i] if i < len(only_in_cp411) else "",
                    "Police non retrouvée dans 410" if i < len(only_in_cp411) else ""
                ]
                
                tags = ('match',) if i < len(common_policies) else ('nomatch',)
                self.verification_tree_411.insert("", "end", values=values, tags=tags)
            
            # Update statistics
            stats = (
                f"Total CP_411: {len(cp411_policies)} | "
                f"Total CP_410: {len(cp410_policies)} | "
                f"Correspondances: {len(common_policies)} | "
                f"Différences: {len(only_in_cp411)}"
            )
            self.update_status(f"Vérification terminée. {stats}")
            
            messagebox.showinfo(
                "Vérification terminée", 
                f"Résultats de la vérification:\n\n"
                f"Polices communes: {len(common_policies)}\n"
                f"Polices uniquement dans CP_411: {len(only_in_cp411)}"
            )
            
        except Exception as e:
            messagebox.showerror(
                "Erreur de vérification", 
                f"Une erreur s'est produite lors de la vérification:\n{str(e)}"
            )
            self.update_status("Erreur lors de la vérification")
    
    def verify_ref_piece(self):
        """Verify reference numbers in CP_411 and collect incorrect numbers"""
        try:
            self.update_status("Vérification des références...")
            
            if self.cp411_data is None:
                messagebox.showwarning(
                    "Données manquantes", 
                    "Veuillez importer le fichier CP_411 avant de vérifier."
                )
                return
            
            if "Réf Pièce" not in self.cp411_data.columns:
                messagebox.showwarning(
                    "Colonne manquante", 
                    "Le fichier CP_411 doit contenir une colonne 'Réf Pièce'."
                )
                return
            
            # Define pattern for valid references
            pattern = r"^\w+-\d+(?:/\d+)?$"
            
            # Find invalid references
            self.numero_recu_list = [
                str(ref) for ref in self.cp411_data["Réf Pièce"] 
                if not re.match(pattern, str(ref)) and pd.notna(ref)
            ]
            
            # Clear and update treeview
            self.rapprochement_tree.delete(*self.rapprochement_tree.get_children())
            
            for i, ref in enumerate(self.numero_recu_list):
                tag = 'evenrow' if i % 2 == 0 else 'oddrow'
                self.rapprochement_tree.insert("", "end", values=(ref, ""), tags=(tag,))
            
            stats = f"Références vérifiées: {len(self.cp411_data)} | Invalides: {len(self.numero_recu_list)}"
            self.update_status(f"Vérification des références terminée. {stats}")
            
            messagebox.showinfo(
                "Vérification terminée", 
                f"Résultats de la vérification:\n\n"
                f"Références totales: {len(self.cp411_data)}\n"
                f"Références invalides: {len(self.numero_recu_list)}"
            )
            
        except Exception as e:
            messagebox.showerror(
                "Erreur de vérification", 
                f"Une erreur s'est produite lors de la vérification:\n{str(e)}"
            )
            self.update_status("Erreur lors de la vérification des références")
    
    def find_policies_for_recu(self):
        """Find associated policies for each receipt number in CP_411"""
        try:
            self.update_status("Recherche des polices associées...")
            
            if not self.numero_recu_list:
                messagebox.showwarning(
                    "Aucune référence", 
                    "Veuillez d'abord vérifier les références."
                )
                return
            
            if self.cp411_data is None:
                messagebox.showwarning(
                    "Données manquantes", 
                    "Veuillez importer CP_411 avant de vérifier."
                )
                return
            
            if "Libellé" not in self.cp411_data.columns or "No Police" not in self.cp411_data.columns:
                messagebox.showwarning(
                    "Colonnes manquantes", 
                    "CP_411 doit contenir les colonnes 'Libellé' et 'No Police'."
                )
                return
            
            # Create dictionary to store results
            police_associee_dict = {numero: [] for numero in self.numero_recu_list}
            
            # Search for each receipt number in Libellé
            for _, row in self.cp411_data.iterrows():
                libelle = str(row["Libellé"])
                police = str(row["No Police"])
                
                for numero_recu in self.numero_recu_list:
                    if numero_recu in libelle:
                        police_associee_dict[numero_recu].append(police)
            
            # Clear and update treeview
            self.rapprochement_tree.delete(*self.rapprochement_tree.get_children())
            
            for i, (numero_recu, polices) in enumerate(police_associee_dict.items()):
                police_associee = ", ".join(sorted(set(polices))) if polices else "Aucune"
                tag = 'evenrow' if i % 2 == 0 else 'oddrow'
                self.rapprochement_tree.insert("", "end", values=(numero_recu, police_associee), tags=(tag,))
            
            self.update_status("Recherche des polices associées terminée")
            
        except Exception as e:
            messagebox.showerror(
                "Erreur de recherche", 
                f"Une erreur s'est produite lors de la recherche:\n{str(e)}"
            )
            self.update_status("Erreur lors de la recherche des polices")

class Gestion41_Tech(tk.Tk):
    def __init__(self, pivot_techniques, pivot_comptables):
        super().__init__()
        self.title("Gestion 41 - Tech/Compta - Tableau de bord")
        self.geometry("1200x750")
        self.minsize(1000, 600)
        
        # Configuration du style
        self.configure_style()
        
        # Données initiales
        self.pivot_techniques = pivot_techniques
        self.pivot_comptables = pivot_comptables
        
        # Configuration de la fenêtre principale
        self.create_main_container()
        
        # Création des onglets
        self.create_tabs()
        
        # Initialisation des données
        self.initialize_data()
        
        # Affichage initial
        self.display_initial_data()
        
        # Barre de statut
        self.create_status_bar()
        
        # Menu
        self.create_menu()
    
    def configure_style(self):
        """Configure les styles visuels de l'application"""
        self.style = ttk.Style()
        #self.style.theme_use('clam')
        
        # Couleurs
        self.primary_color = '#4b8bbe'  # Bleu
        self.secondary_color = '#e0e0e0'  # Gris clair
        self.accent_color = '#ff8c00'  # Orange
        self.success_color = '#d4edda'  # Vert clair
        self.warning_color = '#f8d7da'  # Rouge clair
        
        # Configuration des styles
        self.style.configure('TNotebook', background=self.secondary_color)
        self.style.configure('TNotebook.Tab', 
                           font=('Helvetica', 10, 'bold'), 
                           padding=[10, 5],
                           background=self.secondary_color)
        self.style.map('TNotebook.Tab', 
                      background=[('selected', self.primary_color), 
                                ('active', '#5c9ccc')],
                      foreground=[('selected', 'white')])
        
        self.style.configure('TFrame', background=self.secondary_color)
        self.style.configure('TButton', font=('Helvetica', 10), padding=5)
        self.style.map('TButton', 
                      foreground=[('active', 'white'), ('!disabled', 'black')],
                      background=[('active', self.primary_color), 
                                ('!disabled', self.secondary_color)])
        
        self.style.configure('Treeview', font=('Helvetica', 9), rowheight=25)
        self.style.configure('Treeview.Heading', font=('Helvetica', 10, 'bold'))
        self.style.map('Treeview', background=[('selected', self.primary_color)])
        
        self.style.configure('Accent.TButton', background=self.accent_color)
        self.style.configure('Success.TButton', background=self.success_color)
        self.style.configure('Warning.TButton', background=self.warning_color)
    
    def create_main_container(self):
        """Crée le conteneur principal"""
        self.main_container = ttk.Frame(self)
        self.main_container.pack(fill='both', expand=True, padx=10, pady=10)
    
    def create_tabs(self):
        """Crée les onglets de l'application"""
        self.notebook = ttk.Notebook(self.main_container)
        self.notebook.pack(expand=True, fill="both")
        
        # Noms des onglets
        tab_names = [
            "Données Techniques", 
            "Données Comptables", 
            "Données Compte 41",
            "Vérification Police 41/Tech/Compta", 
            "Vérification Police Tech/Compta/41", 
            "Vérification Police Compta/41/Tech", 
            "Vérification Reliquat"
        ]
        
        # Création des onglets
        self.tabs = {}
        for name in tab_names:
            self.tabs[name] = ttk.Frame(self.notebook)
            self.notebook.add(self.tabs[name], text=name)
            self.create_tab_content(name)
    
    def create_tab_content(self, tab_name):
        """Crée le contenu d'un onglet"""
        frame = self.tabs[tab_name]
        
        # Contrôles supérieurs (recherche, boutons)
        control_frame = ttk.Frame(frame)
        control_frame.pack(fill='x', pady=(0, 10))
        
        # Barre de recherche
        search_frame = ttk.Frame(control_frame)
        search_frame.pack(side='left', fill='x', expand=True)
        
        search_label = ttk.Label(search_frame, text="Rechercher :")
        search_label.pack(side='left', padx=(0, 5))
        
        search_entry = ttk.Entry(search_frame)
        search_entry.pack(side='left', fill='x', expand=True)
        search_entry.bind("<KeyRelease>", lambda e, tn=tab_name: self.search_data(e, tn))
        self.tabs[tab_name + "_search"] = search_entry
        
        # Bouton Réinitialiser
        reset_button = ttk.Button(
            control_frame, 
            text="Réinitialiser", 
            command=lambda: self.reset_search(tab_name)
        )
        reset_button.pack(side='left', padx=5)
        
        # Bouton Importer pour l'onglet Compte 41
        if tab_name == "Données Compte 41":
            import_button = ttk.Button(
                control_frame, 
                text="Importer", 
                command=lambda: self.import_data(tab_name),
                style='Accent.TButton'
            )
            import_button.pack(side='left', padx=5)
        
        # Bouton Exporter pour les onglets de vérification
        if "Vérification" in tab_name:
            export_button = ttk.Button(
                control_frame, 
                text="Exporter", 
                command=lambda: self.export_data(tab_name),
                style='Success.TButton'
            )
            export_button.pack(side='left', padx=5)
            
            # Bouton pour lancer la vérification
            if tab_name != "Vérification Reliquat":
                verify_button = ttk.Button(
                    control_frame,
                    text="Vérifier",
                    command=lambda: self.run_verification(tab_name),
                    style='Warning.TButton'
                )
                verify_button.pack(side='left', padx=5)
        
        # Tableau de données
        tree_frame = ttk.Frame(frame)
        tree_frame.pack(expand=True, fill="both")
        
        # Barres de défilement
        tree_scroll_y = ttk.Scrollbar(tree_frame, orient="vertical")
        tree_scroll_x = ttk.Scrollbar(tree_frame, orient="horizontal")
        
        # Treeview
        treeview = ttk.Treeview(
            tree_frame, 
            yscrollcommand=tree_scroll_y.set, 
            xscrollcommand=tree_scroll_x.set,
            selectmode='extended'
        )
        
        # Configuration des scrollbars
        tree_scroll_y.config(command=treeview.yview)
        tree_scroll_x.config(command=treeview.xview)
        
        # Placement des éléments
        tree_scroll_y.pack(side="right", fill="y")
        tree_scroll_x.pack(side="bottom", fill="x")
        treeview.pack(expand=True, fill="both")
        
        # Configuration des tags pour les lignes
        treeview.tag_configure('oddrow', background=self.secondary_color)
        treeview.tag_configure('evenrow', background='white')
        treeview.tag_configure('match', background=self.success_color)
        treeview.tag_configure('nomatch', background=self.warning_color)
        
        self.tabs[tab_name + "_tree"] = treeview
    
    def initialize_data(self):
        """Initialise les dataframes"""
        self.dataframes = {
            "Données Techniques": self.pivot_techniques,
            "Données Comptables": self.pivot_comptables,
            "Données Compte 41": pd.DataFrame(),
            "Vérification Police 41/Tech/Compta": pd.DataFrame(),
            "Vérification Police Tech/Compta/41": pd.DataFrame(),
            "Vérification Police Compta/41/Tech": pd.DataFrame(),
            "Vérification Reliquat": pd.DataFrame()
        }
    
    def display_initial_data(self):
        """Affiche les données initiales"""
        self.display_data("Données Techniques", self.pivot_techniques)
        self.display_data("Données Comptables", self.pivot_comptables)
    
    def create_status_bar(self):
        """Crée la barre de statut"""
        self.status_bar = ttk.Label(
            self, 
            text="Prêt", 
            relief='sunken', 
            anchor='w',
            font=('Helvetica', 9)
        )
        self.status_bar.pack(side='bottom', fill='x')
    
    def create_menu(self):
        """Crée le menu de l'application"""
        menubar = tk.Menu(self)
        
        # Menu Fichier
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(
            label="Ouvrir Données Techniques", 
            command=lambda: self.import_data("Données Techniques")
        )
        file_menu.add_command(
            label="Ouvrir Données Comptables", 
            command=lambda: self.import_data("Données Comptables")
        )
        file_menu.add_command(
            label="Ouvrir Données Compte 41", 
            command=lambda: self.import_data("Données Compte 41")
        )
        file_menu.add_separator()
        file_menu.add_command(label="Quitter", command=self.quit)
        menubar.add_cascade(label="Fichier", menu=file_menu)
        
        # Menu Vérification
        verify_menu = tk.Menu(menubar, tearoff=0)
        verify_menu.add_command(
            label="Vérifier Police 41/Tech/Compta", 
            command=lambda: self.run_verification("Vérification Police 41/Tech/Compta")
        )
        verify_menu.add_command(
            label="Vérifier Police Tech/Compta/41", 
            command=lambda: self.run_verification("Vérification Police Tech/Compta/41")
        )
        verify_menu.add_command(
            label="Vérifier Police Compta/41/Tech", 
            command=lambda: self.run_verification("Vérification Police Compta/41/Tech")
        )
        verify_menu.add_command(
            label="Vérifier Reliquat", 
            command=lambda: self.run_verification("Vérification Reliquat")
        )
        menubar.add_cascade(label="Vérification", menu=verify_menu)
        
        # Menu Aide
        help_menu = tk.Menu(menubar, tearoff=0)
        help_menu.add_command(label="À propos", command=self.show_about)
        menubar.add_cascade(label="Aide", menu=help_menu)
        
        self.config(menu=menubar)
    
    def show_about(self):
        """Affiche la fenêtre À propos"""
        about_window = tk.Toplevel(self)
        about_window.title("À propos")
        about_window.geometry("400x200")
        about_window.resizable(False, False)
        
        about_frame = ttk.Frame(about_window)
        about_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        ttk.Label(
            about_frame, 
            text="Gestion 41 - Tech/Compta", 
            font=('Helvetica', 14, 'bold')
        ).pack(pady=10)
        
        ttk.Label(
            about_frame, 
            text="Outil de vérification et de rapprochement des données"
        ).pack()
        
        ttk.Label(
            about_frame, 
            text="Version 2.0", 
            font=('Helvetica', 8)
        ).pack(pady=10)
        
        ttk.Label(about_frame, text="© 2023 - Tous droits réservés").pack()
        
        ttk.Button(
            about_frame, 
            text="Fermer", 
            command=about_window.destroy
        ).pack(pady=10)
    
    def update_status(self, message):
        """Met à jour la barre de statut"""
        self.status_bar.config(text=message)
        self.after(5000, lambda: self.status_bar.config(text="Prêt"))
    
    def import_data(self, tab_name):
        """Importe des données depuis un fichier"""
        try:
            file_types = [
                ("Fichiers Excel", "*.xlsx;*.xls"), 
                ("Fichiers CSV", "*.csv"),
                ("Tous les fichiers", "*.*")
            ]
            
            file_path = filedialog.askopenfilename(filetypes=file_types)
            if not file_path:
                return
            
            self.update_status(f"Importation du fichier {os.path.basename(file_path)}...")
            
            # Afficher une fenêtre de progression pour les gros fichiers
            if os.path.getsize(file_path) > 1_000_000:  # > 10MB
                progress_window = tk.Toplevel(self)
                progress_window.title("Importation en cours")
                progress_window.geometry("300x100")
                
                progress_label = ttk.Label(progress_window, text="Chargement du fichier...")
                progress_label.pack(pady=10)
                
                progress_bar = ttk.Progressbar(progress_window, orient="horizontal", mode="indeterminate")
                progress_bar.pack(fill='x', padx=20)
                progress_bar.start()
                
                self.update_idletasks()
            
            # Lecture du fichier
            if file_path.endswith(".csv"):
                df = pd.read_csv(file_path, dtype=str)
            else:
                df = pd.read_excel(file_path, dtype=str)
            
            # Fermer la fenêtre de progression si elle existe
            if 'progress_window' in locals():
                progress_bar.stop()
                progress_window.destroy()
            
            # Nettoyage des noms de colonnes
            df.columns = df.columns.str.strip()
            
            # Stockage des données
            self.dataframes[tab_name] = df
            self.display_data(tab_name, df)
            
            # Vérification automatique si toutes les données sont chargées
            if all(not df.empty for df in [
                self.dataframes["Données Techniques"],
                self.dataframes["Données Comptables"],
                self.dataframes["Données Compte 41"]
            ]):
                self.verify_policies_after_import()
            
            self.update_status(f"Données importées avec succès : {len(df)} enregistrements")
            
        except Exception as e:
            messagebox.showerror(
                "Erreur d'importation", 
                f"Impossible d'importer le fichier :\n{str(e)}"
            )
            self.update_status("Erreur lors de l'importation")
    
    def display_data(self, tab_name, df):
        """Affiche les données dans le Treeview"""
        tree = self.tabs[tab_name + "_tree"]
        tree.delete(*tree.get_children())
        
        if df is None or df.empty:
            return
        
        # Configuration des colonnes
        tree['columns'] = list(df.columns)
        tree['show'] = 'headings'
        
        for col in df.columns:
            tree.heading(col, text=col)
            tree.column(col, width=150, anchor='w')
        
        # Insertion des données avec alternance de couleurs
        for i, (_, row) in enumerate(df.iterrows()):
            tag = 'evenrow' if i % 2 == 0 else 'oddrow'
            tree.insert("", "end", values=list(row), tags=(tag,))
    
    def search_data(self, event, tab_name):
        """Filtre les données selon la recherche"""
        search_text = self.tabs[tab_name + "_search"].get().lower()
        treeview = self.tabs[tab_name + "_tree"]
        df = self.dataframes.get(tab_name)
        
        if df is None or df.empty:
            return
        
        # Effacer les résultats précédents
        treeview.delete(*treeview.get_children())
        
        # Filtrer les données
        try:
            filtered_df = df[df.apply(
                lambda row: row.astype(str).str.contains(search_text, case=False).any(), 
                axis=1
            )]
            
            # Afficher les résultats filtrés
            for i, (_, row) in enumerate(filtered_df.iterrows()):
                tag = 'evenrow' if i % 2 == 0 else 'oddrow'
                treeview.insert("", "end", values=list(row), tags=(tag,))
                
        except Exception as e:
            self.update_status(f"Erreur de recherche : {str(e)}")
    
    def reset_search(self, tab_name):
        """Réinitialise la recherche"""
        self.tabs[tab_name + "_search"].delete(0, tk.END)
        self.display_data(tab_name, self.dataframes.get(tab_name))
        self.update_status("Recherche réinitialisée")
    
    def export_data(self, tab_name):
        """Exporte les données vers Excel"""
        df = self.dataframes.get(tab_name)
        
        if df is None or df.empty:
            messagebox.showwarning(
                "Avertissement", 
                f"Aucune donnée à exporter pour {tab_name}."
            )
            return
        
        try:
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                initialfile=f"{tab_name}.xlsx",
                filetypes=[("Fichiers Excel", "*.xlsx"), ("Tous les fichiers", "*.*")]
            )
            
            if file_path:
                df.to_excel(file_path, index=False)
                messagebox.showinfo(
                    "Exportation réussie", 
                    f"Les données ont été exportées avec succès :\n{file_path}"
                )
                self.update_status(f"Données exportées : {file_path}")
                
        except Exception as e:
            messagebox.showerror(
                "Erreur d'exportation", 
                f"Impossible d'exporter le fichier :\n{str(e)}"
            )
            self.update_status("Erreur lors de l'exportation")
    
    def run_verification(self, tab_name):
        """Lance la vérification appropriée selon l'onglet"""
        if tab_name == "Vérification Police 41/Tech/Compta":
            self.verify_policies_41()
        elif tab_name == "Vérification Police Tech/Compta/41":
            self.verify_policies_tech()
        elif tab_name == "Vérification Police Compta/41/Tech":
            self.verify_policies_compta_41_tech()
        elif tab_name == "Vérification Reliquat":
            self.verify_policies_reliquat()
    
    def verify_policies_after_import(self):
        """Lance toutes les vérifications après un import"""
        self.verify_policies_41()
        self.verify_policies_tech()
        self.verify_policies_compta_41_tech()
        self.verify_policies_reliquat()
        self.update_status("Toutes les vérifications ont été effectuées")
    
    def verify_policies_41(self):
        """Vérifie les polices entre 41, Tech et Compta"""
        try:
            self.update_status("Vérification Police 41/Tech/Compta en cours...")
            
            tech_df = self.dataframes["Données Techniques"]
            compta_df = self.dataframes["Données Comptables"]
            compte41_df = self.dataframes["Données Compte 41"]
            
            # Vérification des colonnes nécessaires
            required_columns = {
                "Données Techniques": "Nouvelle_Police",
                "Données Comptables": "No Police",
                "Données Compte 41": "No Police"
            }
            
            for df_name, col_name in required_columns.items():
                if col_name not in self.dataframes[df_name].columns:
                    raise ValueError(f"La colonne '{col_name}' est manquante dans {df_name}")
            
            # Récupération des polices
            policies_tech = set(tech_df["Nouvelle_Police"].astype(str).str.strip())
            policies_compta = set(compta_df["No Police"].astype(str).str.strip())
            policies_compte41 = set(compte41_df["No Police"].astype(str).str.strip())
            
            # Polices communes et différentes
            found_policies = policies_compte41.intersection(policies_tech, policies_compta)
            not_found_policies = policies_compte41 - found_policies
            
            # Création du DataFrame de résultats
            max_rows = max(len(found_policies), len(not_found_policies))
            verification_data = {
                "Police Trouvée": list(found_policies) + [""] * (max_rows - len(found_policies)),
                "État": ["Police trouvée"] * len(found_policies) + [""] * (max_rows - len(found_policies)),
                "Police Non Trouvée": list(not_found_policies) + [""] * (max_rows - len(not_found_policies)),
                "État.1": ["Police non trouvée"] * len(not_found_policies) + [""] * (max_rows - len(not_found_policies))
            }
            
            verification_df = pd.DataFrame(verification_data)
            self.dataframes["Vérification Police 41/Tech/Compta"] = verification_df
            self.display_data("Vérification Police 41/Tech/Compta", verification_df)
            
            # Mise à jour du statut avec les statistiques
            stats = (
                f"Polices 41: {len(policies_compte41)} | "
                f"Correspondances: {len(found_policies)} | "
                f"Différences: {len(not_found_policies)}"
            )
            self.update_status(f"Vérification Police 41/Tech/Compta terminée. {stats}")
            
            messagebox.showinfo(
                "Vérification terminée",
                f"Résultats de la vérification 41/Tech/Compta:\n\n"
                f"Polices communes: {len(found_policies)}\n"
                f"Polices uniquement dans 41: {len(not_found_policies)}"
            )
            
        except Exception as e:
            messagebox.showerror(
                "Erreur de vérification", 
                f"Erreur lors de la vérification Police 41/Tech/Compta:\n{str(e)}"
            )
            self.update_status("Erreur lors de la vérification")
    
    def verify_policies_tech(self):
        """Vérifie les polices entre Tech, Compta et 41"""
        try:
            self.update_status("Vérification Police Tech/Compta/41 en cours...")
            
            tech_df = self.dataframes["Données Techniques"]
            compta_df = self.dataframes["Données Comptables"]
            compte41_df = self.dataframes["Données Compte 41"]
            
            # Vérification des colonnes nécessaires
            if "Nouvelle_Police" not in tech_df.columns:
                raise ValueError("La colonne 'Nouvelle_Police' est manquante dans Données Techniques")
            
            # Récupération des polices
            policies_tech = set(tech_df["Nouvelle_Police"].astype(str).str.strip())
            policies_compta = set(compta_df["No Police"].astype(str).str.strip())
            policies_compte41 = set(compte41_df["No Police"].astype(str).str.strip())
            
            # Polices communes et différentes
            found_policies = policies_tech.intersection(policies_compta, policies_compte41)
            not_found_policies = policies_tech - found_policies
            
            # Création du DataFrame de résultats
            max_rows = max(len(found_policies), len(not_found_policies))
            verification_data = {
                "Police Trouvée": list(found_policies) + [""] * (max_rows - len(found_policies)),
                "État": ["Police trouvée"] * len(found_policies) + [""] * (max_rows - len(found_policies)),
                "Police Non Trouvée": list(not_found_policies) + [""] * (max_rows - len(not_found_policies)),
                "État.1": ["Police non trouvée"] * len(not_found_policies) + [""] * (max_rows - len(not_found_policies))
            }
            
            verification_df = pd.DataFrame(verification_data)
            self.dataframes["Vérification Police Tech/Compta/41"] = verification_df
            self.display_data("Vérification Police Tech/Compta/41", verification_df)
            
            # Mise à jour du statut avec les statistiques
            stats = (
                f"Polices Tech: {len(policies_tech)} | "
                f"Correspondances: {len(found_policies)} | "
                f"Différences: {len(not_found_policies)}"
            )
            self.update_status(f"Vérification Police Tech/Compta/41 terminée. {stats}")
            
            messagebox.showinfo(
                "Vérification terminée",
                f"Résultats de la vérification Tech/Compta/41:\n\n"
                f"Polices communes: {len(found_policies)}\n"
                f"Polices uniquement dans Tech: {len(not_found_policies)}"
            )
            
        except Exception as e:
            messagebox.showerror(
                "Erreur de vérification", 
                f"Erreur lors de la vérification Police Tech/Compta/41:\n{str(e)}"
            )
            self.update_status("Erreur lors de la vérification")
    
    def verify_policies_compta_41_tech(self):
        """Vérifie les polices entre Compta, 41 et Tech"""
        try:
            self.update_status("Vérification Police Compta/41/Tech en cours...")
            
            tech_df = self.dataframes["Données Techniques"]
            compta_df = self.dataframes["Données Comptables"]
            compte41_df = self.dataframes["Données Compte 41"]
            
            # Vérification des colonnes nécessaires
            if "No Police" not in compta_df.columns:
                raise ValueError("La colonne 'No Police' est manquante dans Données Comptables")
            
            # Récupération des polices
            policies_tech = set(tech_df["Nouvelle_Police"].astype(str).str.strip())
            policies_compta = set(compta_df["No Police"].astype(str).str.strip())
            policies_compte41 = set(compte41_df["No Police"].astype(str).str.strip())
            
            # Polices communes et différentes
            found_policies = policies_compta.intersection(policies_compte41, policies_tech)
            not_found_policies = policies_compta - found_policies
            
            # Création du DataFrame de résultats
            max_rows = max(len(found_policies), len(not_found_policies))
            verification_data = {
                "Police Trouvée": list(found_policies) + [""] * (max_rows - len(found_policies)),
                "État": ["Police trouvée"] * len(found_policies) + [""] * (max_rows - len(found_policies)),
                "Police Non Trouvée": list(not_found_policies) + [""] * (max_rows - len(not_found_policies)),
                "État.1": ["Police non trouvée"] * len(not_found_policies) + [""] * (max_rows - len(not_found_policies))
            }
            
            verification_df = pd.DataFrame(verification_data)
            self.dataframes["Vérification Police Compta/41/Tech"] = verification_df
            self.display_data("Vérification Police Compta/41/Tech", verification_df)
            
            # Mise à jour du statut avec les statistiques
            stats = (
                f"Polices Compta: {len(policies_compta)} | "
                f"Correspondances: {len(found_policies)} | "
                f"Différences: {len(not_found_policies)}"
            )
            self.update_status(f"Vérification Police Compta/41/Tech terminée. {stats}")
            
            messagebox.showinfo(
                "Vérification terminée",
                f"Résultats de la vérification Compta/41/Tech:\n\n"
                f"Polices communes: {len(found_policies)}\n"
                f"Polices uniquement dans Compta: {len(not_found_policies)}"
            )
            
        except Exception as e:
            messagebox.showerror(
                "Erreur de vérification", 
                f"Erreur lors de la vérification Police Compta/41/Tech:\n{str(e)}"
            )
            self.update_status("Erreur lors de la vérification")
    
    def verify_policies_reliquat(self):
        """Vérifie les reliquats des polices"""
        try:
            self.update_status("Vérification des reliquats en cours...")
            
            verification_df = self.dataframes["Vérification Police 41/Tech/Compta"]
            tech_df = self.dataframes["Données Techniques"]
            
            # Vérification des colonnes nécessaires
            if "Police Trouvée" not in verification_df.columns:
                raise ValueError("La colonne 'Police Trouvée' est manquante dans les vérifications")
            
            if "Reliquat A Encaisser" not in tech_df.columns or "Etat Quittance" not in tech_df.columns:
                raise ValueError("Colonnes manquantes dans Données Techniques")
            
            # Récupération des polices vérifiées
            verified_policies = verification_df["Police Trouvée"].dropna().unique()
            
            # Préparation des résultats
            results = []
            
            for police in verified_policies:
                if not police:  # Ignorer les valeurs vides
                    continue
                
                # Recherche de la police dans les données techniques
                tech_row = tech_df[tech_df["Nouvelle_Police"].astype(str).str.strip() == police.strip()]
                
                if not tech_row.empty:
                    reliquat = tech_row["Reliquat A Encaisser"].values[0]
                    etat = tech_row["Etat Quittance"].values[0]
                    
                    try:
                        reliquat_num = float(reliquat) if str(reliquat).strip() else 0.0
                    except ValueError:
                        reliquat_num = 0.0
                    
                    # Détermination du statut
                    if reliquat_num == 0:
                        status = "Payé"
                    elif reliquat_num > 0:
                        status = "Impayé"
                    else:
                        status = "Inconnu"
                    
                    results.append({
                        "Police": police,
                        "Reliquat": reliquat,
                        "Etat Quittance": etat,
                        "Statut": status
                    })
            
            # Création du DataFrame de résultats
            reliquat_df = pd.DataFrame(results)
            self.dataframes["Vérification Reliquat"] = reliquat_df
            self.display_data("Vérification Reliquat", reliquat_df)
            
            # Calcul des statistiques
            if not reliquat_df.empty:
                paye_count = (reliquat_df["Statut"] == "Payé").sum()
                impaye_count = (reliquat_df["Statut"] == "Impayé").sum()
                inconnu_count = (reliquat_df["Statut"] == "Inconnu").sum()
                
                stats = (
                    f"Polices vérifiées: {len(reliquat_df)} | "
                    f"Payées: {paye_count} | "
                    f"Impayées: {impaye_count} | "
                    f"Inconnues: {inconnu_count}"
                )
            else:
                stats = "Aucune police à vérifier"
            
            self.update_status(f"Vérification des reliquats terminée. {stats}")
            
            messagebox.showinfo(
                "Vérification terminée",
                f"Résultats de la vérification des reliquats:\n\n"
                f"Polices vérifiées: {len(reliquat_df)}\n"
                f"Polices payées: {paye_count}\n"
                f"Polices impayées: {impaye_count}\n"
                f"Statut inconnu: {inconnu_count}"
            )
            
        except Exception as e:
            messagebox.showerror(
                "Erreur de vérification", 
                f"Erreur lors de la vérification des reliquats:\n{str(e)}"
            )
            self.update_status("Erreur lors de la vérification des reliquats")

class Application_Analyse(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Analyse des Doublons de Polices")
        self.geometry("1200x700")
        self.minsize(1000, 600)
        
        # Configuration du style
        self.configure_style()
        
        # Configuration de la fenêtre principale
        self.create_main_container()
        
        # Création des onglets
        self.create_notebook()
        
        # Barre de statut
        self.create_status_bar()
        
        # Menu
        self.create_menu()
        
        # Initialisation des données
        self.data_df = pd.DataFrame()
        self.duplicates_df = pd.DataFrame()
        self.no_duplicates_df = pd.DataFrame()
    
    def configure_style(self):
        """Configure les styles visuels de l'application"""
        self.style = ttk.Style()
        #self.style.theme_use('clam')
        
        # Couleurs
        self.primary_color = '#4b8bbe'  # Bleu
        self.secondary_color = '#e0e0e0'  # Gris clair
        self.accent_color = '#ff8c00'  # Orange
        self.success_color = '#d4edda'  # Vert clair
        self.warning_color = '#f8d7da'  # Rouge clair
        
        # Configuration des styles
        self.style.configure('TNotebook', background=self.secondary_color)
        self.style.configure('TNotebook.Tab', 
                           font=('Helvetica', 10, 'bold'), 
                           padding=[10, 5],
                           background=self.secondary_color)
        self.style.map('TNotebook.Tab', 
                      background=[('selected', self.primary_color), 
                                ('active', '#5c9ccc')],
                      foreground=[('selected', 'white')])
        
        self.style.configure('TFrame', background=self.secondary_color)
        self.style.configure('TButton', font=('Helvetica', 10), padding=5)
        self.style.map('TButton', 
                      foreground=[('active', 'white'), ('!disabled', 'black')],
                      background=[('active', self.primary_color), 
                                ('!disabled', self.secondary_color)])
        
        self.style.configure('Treeview', font=('Helvetica', 9), rowheight=25)
        self.style.configure('Treeview.Heading', font=('Helvetica', 10, 'bold'))
        self.style.map('Treeview', background=[('selected', self.primary_color)])
        
        self.style.configure('Accent.TButton', background=self.accent_color)
        self.style.configure('Success.TButton', background=self.success_color)
        self.style.configure('Warning.TButton', background=self.warning_color)
    
    def create_main_container(self):
        """Crée le conteneur principal"""
        self.main_container = ttk.Frame(self)
        self.main_container.pack(fill='both', expand=True, padx=10, pady=10)
    
    def create_notebook(self):
        """Crée le notebook avec les onglets"""
        self.notebook = ttk.Notebook(self.main_container)
        self.notebook.pack(fill='both', expand=True)
        
        # Onglet Données
        self.tab_data = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_data, text="Données Importées")
        self.create_data_tab()
        
        # Onglet Doublons
        self.tab_duplicates = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_duplicates, text="Polices en Doublon")
        self.create_duplicates_tab()
        
        # Onglet Sans Doublons
        self.tab_no_duplicates = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_no_duplicates, text="Polices Uniques")
        self.create_no_duplicates_tab()
    
    def create_data_tab(self):
        """Crée le contenu de l'onglet Données Importées"""
        # Frame pour les boutons
        button_frame = ttk.Frame(self.tab_data)
        button_frame.pack(fill='x', pady=10)
        
        # Bouton Importer
        self.import_button = ttk.Button(
            button_frame, 
            text="Importer des Données", 
            command=self.import_data,
            style='Accent.TButton'
        )
        self.import_button.pack(side='left', padx=5)
        
        # Bouton Afficher Doublons
        self.show_duplicates_button = ttk.Button(
            button_frame,
            text="Afficher les Doublons",
            command=self.show_duplicates,
            style='Warning.TButton'
        )
        self.show_duplicates_button.pack(side='left', padx=5)
        
        # Bouton Afficher Uniques
        self.show_no_duplicates_button = ttk.Button(
            button_frame,
            text="Afficher les Polices Uniques",
            command=self.show_no_duplicates,
            style='Success.TButton'
        )
        self.show_no_duplicates_button.pack(side='left', padx=5)
        
        # Barre de recherche
        self.create_search_bar(self.tab_data, "data")
        
        # Tableau de données
        self.data_table = self.create_table(self.tab_data)
    
    def create_duplicates_tab(self):
        """Crée le contenu de l'onglet Doublons"""
        # Frame pour les boutons
        button_frame = ttk.Frame(self.tab_duplicates)
        button_frame.pack(fill='x', pady=10)
        
        # Bouton Exporter
        self.export_duplicates_button = ttk.Button(
            button_frame,
            text="Exporter vers Excel",
            command=self.export_duplicates_to_excel,
            style='Accent.TButton'
        )
        self.export_duplicates_button.pack(side='left', padx=5)
        
        # Barre de recherche
        self.create_search_bar(self.tab_duplicates, "duplicates")
        
        # Tableau de données
        self.duplicates_table = self.create_table(self.tab_duplicates)
        
        # Étiquette d'information
        self.duplicates_info = ttk.Label(
            self.tab_duplicates,
            text="Aucune donnée à afficher",
            font=('Helvetica', 10, 'italic')
        )
        self.duplicates_info.pack(pady=5)
    
    def create_no_duplicates_tab(self):
        """Crée le contenu de l'onglet Sans Doublons"""
        # Frame pour les boutons
        button_frame = ttk.Frame(self.tab_no_duplicates)
        button_frame.pack(fill='x', pady=10)
        
        # Bouton Exporter
        self.export_no_duplicates_button = ttk.Button(
            button_frame,
            text="Exporter vers Excel",
            command=self.export_no_duplicates_to_excel,
            style='Accent.TButton'
        )
        self.export_no_duplicates_button.pack(side='left', padx=5)
        
        # Barre de recherche
        self.create_search_bar(self.tab_no_duplicates, "no_duplicates")
        
        # Tableau de données
        self.no_duplicates_table = self.create_table(self.tab_no_duplicates)
        
        # Étiquette d'information
        self.no_duplicates_info = ttk.Label(
            self.tab_no_duplicates,
            text="Aucune donnée à afficher",
            font=('Helvetica', 10, 'italic')
        )
        self.no_duplicates_info.pack(pady=5)
    
    def create_status_bar(self):
        """Crée la barre de statut"""
        self.status_bar = ttk.Label(
            self, 
            text="Prêt", 
            relief='sunken', 
            anchor='w',
            font=('Helvetica', 9)
        )
        self.status_bar.pack(side='bottom', fill='x')
    
    def create_menu(self):
        """Crée le menu de l'application"""
        menubar = tk.Menu(self)
        
        # Menu Fichier
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="Importer", command=self.import_data)
        file_menu.add_separator()
        file_menu.add_command(label="Quitter", command=self.quit)
        menubar.add_cascade(label="Fichier", menu=file_menu)
        
        # Menu Analyse
        analyze_menu = tk.Menu(menubar, tearoff=0)
        analyze_menu.add_command(label="Afficher les Doublons", command=self.show_duplicates)
        analyze_menu.add_command(label="Afficher les Polices Uniques", command=self.show_no_duplicates)
        menubar.add_cascade(label="Analyse", menu=analyze_menu)
        
        # Menu Export
        export_menu = tk.Menu(menubar, tearoff=0)
        export_menu.add_command(label="Exporter les Doublons", command=self.export_duplicates_to_excel)
        export_menu.add_command(label="Exporter les Polices Uniques", command=self.export_no_duplicates_to_excel)
        menubar.add_cascade(label="Export", menu=export_menu)
        
        # Menu Aide
        help_menu = tk.Menu(menubar, tearoff=0)
        help_menu.add_command(label="À propos", command=self.show_about)
        menubar.add_cascade(label="Aide", menu=help_menu)
        
        self.config(menu=menubar)
    
    def show_about(self):
        """Affiche la fenêtre À propos"""
        about_window = tk.Toplevel(self)
        about_window.title("À propos")
        about_window.geometry("400x200")
        about_window.resizable(False, False)
        
        about_frame = ttk.Frame(about_window)
        about_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        ttk.Label(
            about_frame, 
            text="Analyse des Doublons de Polices", 
            font=('Helvetica', 14, 'bold')
        ).pack(pady=10)
        
        ttk.Label(
            about_frame, 
            text="Outil de détection et d'analyse des doublons"
        ).pack()
        
        ttk.Label(
            about_frame, 
            text="Version 2.0", 
            font=('Helvetica', 8)
        ).pack(pady=10)
        
        ttk.Label(about_frame, text="© 2023 - Tous droits réservés").pack()
        
        ttk.Button(
            about_frame, 
            text="Fermer", 
            command=about_window.destroy
        ).pack(pady=10)
    
    def update_status(self, message):
        """Met à jour la barre de statut"""
        self.status_bar.config(text=message)
        self.after(5000, lambda: self.status_bar.config(text="Prêt"))
    
    def create_search_bar(self, parent, table_type):
        """Crée une barre de recherche pour un tableau"""
        search_frame = ttk.Frame(parent)
        search_frame.pack(fill='x', pady=5)
        
        search_label = ttk.Label(search_frame, text="Rechercher :")
        search_label.pack(side='left', padx=5)
        
        search_var = tk.StringVar()
        search_entry = ttk.Entry(search_frame, textvariable=search_var)
        search_entry.pack(side='left', fill='x', expand=True, padx=5)
        
        # Bouton Effacer
        clear_button = ttk.Button(
            search_frame,
            text="Effacer",
            command=lambda: self.clear_search(table_type, search_var)
        )
        clear_button.pack(side='left', padx=5)
        
        search_var.trace("w", lambda *args: self.filter_table(search_var.get(), table_type))
    
    def clear_search(self, table_type, search_var):
        """Efface la recherche"""
        search_var.set("")
        self.filter_table("", table_type)
    
    def create_table(self, parent):
        """Crée un tableau avec barres de défilement"""
        table_frame = ttk.Frame(parent)
        table_frame.pack(fill='both', expand=True)
        
        # Barres de défilement
        scroll_y = ttk.Scrollbar(table_frame, orient='vertical')
        scroll_x = ttk.Scrollbar(table_frame, orient='horizontal')
        
        # Tableau
        table = ttk.Treeview(
            table_frame,
            yscrollcommand=scroll_y.set,
            xscrollcommand=scroll_x.set,
            selectmode='extended'
        )
        
        # Configuration des scrollbars
        scroll_y.config(command=table.yview)
        scroll_x.config(command=table.xview)
        
        # Placement des éléments
        scroll_y.pack(side='right', fill='y')
        scroll_x.pack(side='bottom', fill='x')
        table.pack(fill='both', expand=True)
        
        # Configuration des tags pour les lignes
        table.tag_configure('oddrow', background=self.secondary_color)
        table.tag_configure('evenrow', background='white')
        
        return table
    
    def import_data(self):
        """Importe des données depuis un fichier"""
        file_types = [
            ("Fichiers CSV", "*.csv"), 
            ("Fichiers Excel", "*.xlsx;*.xls"),
            ("Tous les fichiers", "*.*")
        ]
        
        file_path = filedialog.askopenfilename(filetypes=file_types)
        if not file_path:
            return
        
        try:
            self.update_status(f"Importation du fichier {file_path}...")
            
            # Lecture du fichier
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path, dtype=str)
            else:
                df = pd.read_excel(file_path, dtype=str)
            
            # Nettoyage des noms de colonnes
            df.columns = df.columns.str.strip()
            
            # Stockage des données
            self.data_df = df
            self.display_data(self.data_table, df)
            
            # Mise à jour de l'interface
            self.update_info_labels()
            self.update_status(f"Données importées avec succès : {len(df)} enregistrements")
            
            messagebox.showinfo(
                "Importation réussie",
                f"Le fichier a été importé avec succès.\n{len(df)} enregistrements chargés."
            )
            
        except Exception as e:
            messagebox.showerror(
                "Erreur d'importation",
                f"Une erreur est survenue lors de l'importation :\n{str(e)}"
            )
            self.update_status("Erreur lors de l'importation")
    
    def display_data(self, table, df):
        """Affiche les données dans un tableau"""
        table.delete(*table.get_children())
        
        if df.empty:
            return
        
        # Configuration des colonnes
        table['columns'] = list(df.columns)
        
        for col in df.columns:
            table.heading(col, text=col)
            table.column(col, width=150, anchor='w')
        
        # Insertion des données avec alternance de couleurs
        for i, (_, row) in enumerate(df.iterrows()):
            tag = 'evenrow' if i % 2 == 0 else 'oddrow'
            table.insert("", "end", values=list(row), tags=(tag,))
    
    def filter_table(self, query, table_type):
        """Filtre un tableau selon la recherche"""
        table = getattr(self, f"{table_type}_table")
        df = getattr(self, f"{table_type}_df", pd.DataFrame())
        
        if not df.empty and 'NUMERO POLICE' in df.columns:
            filtered_df = df[df['NUMERO POLICE'].astype(str).str.contains(query, case=False, na=False)]
            self.display_data(table, filtered_df)
    
    def show_duplicates(self):
        """Affiche les doublons de polices"""
        if self.data_df.empty:
            messagebox.showwarning(
                "Aucune donnée",
                "Veuillez d'abord importer des données."
            )
            return
        
        if 'NUMERO POLICE' not in self.data_df.columns:
            messagebox.showerror(
                "Colonne manquante",
                "La colonne 'NUMERO POLICE' est introuvable dans les données."
            )
            return
        
        # Détection des doublons
        self.duplicates_df = self.data_df[
            self.data_df.duplicated(subset=['NUMERO POLICE'], keep=False)
        ].sort_values(by='NUMERO POLICE')
        
        # Affichage des résultats
        self.display_data(self.duplicates_table, self.duplicates_df)
        self.update_info_labels()
        
        # Message à l'utilisateur
        if self.duplicates_df.empty:
            messagebox.showinfo(
                "Aucun doublon",
                "Aucun doublon de police n'a été trouvé."
            )
        else:
            messagebox.showinfo(
                "Doublons trouvés",
                f"{len(self.duplicates_df)} doublons ont été identifiés."
            )
        
        self.update_status(f"Doublons identifiés : {len(self.duplicates_df)}")
    
    def show_no_duplicates(self):
        """Affiche les polices sans doublons"""
        if self.data_df.empty:
            messagebox.showwarning(
                "Aucune donnée",
                "Veuillez d'abord importer des données."
            )
            return
        
        if 'NUMERO POLICE' not in self.data_df.columns:
            messagebox.showerror(
                "Colonne manquante",
                "La colonne 'NUMERO POLICE' est introuvable dans les données."
            )
            return
        
        # Détection des polices uniques
        self.no_duplicates_df = self.data_df[
            ~self.data_df.duplicated(subset=['NUMERO POLICE'], keep=False)
        ]
        
        # Affichage des résultats
        self.display_data(self.no_duplicates_table, self.no_duplicates_df)
        self.update_info_labels()
        
        # Message à l'utilisateur
        if self.no_duplicates_df.empty:
            messagebox.showinfo(
                "Aucune police unique",
                "Toutes les polices ont des doublons."
            )
        else:
            messagebox.showinfo(
                "Polices uniques",
                f"{len(self.no_duplicates_df)} polices uniques ont été identifiées."
            )
        
        self.update_status(f"Polices uniques identifiées : {len(self.no_duplicates_df)}")
    
    def update_info_labels(self):
        """Met à jour les étiquettes d'information"""
        # Onglet Doublons
        if hasattr(self, 'duplicates_df') and not self.duplicates_df.empty:
            self.duplicates_info.config(
                text=f"{len(self.duplicates_df)} doublons trouvés",
                foreground='black'
            )
        else:
            self.duplicates_info.config(
                text="Aucun doublon à afficher",
                foreground='gray'
            )
        
        # Onglet Sans Doublons
        if hasattr(self, 'no_duplicates_df') and not self.no_duplicates_df.empty:
            self.no_duplicates_info.config(
                text=f"{len(self.no_duplicates_df)} polices uniques trouvées",
                foreground='black'
            )
        else:
            self.no_duplicates_info.config(
                text="Aucune police unique à afficher",
                foreground='gray'
            )
    
    def export_duplicates_to_excel(self):
        """Exporte les doublons vers Excel"""
        if not hasattr(self, 'duplicates_df') or self.duplicates_df.empty:
            messagebox.showwarning(
                "Aucune donnée",
                "Aucun doublon à exporter."
            )
            return
        
        self.export_dataframe(
            self.duplicates_df,
            "Export des Doublons",
            "doublons_polices.xlsx"
        )
    
    def export_no_duplicates_to_excel(self):
        """Exporte les polices uniques vers Excel"""
        if not hasattr(self, 'no_duplicates_df') or self.no_duplicates_df.empty:
            messagebox.showwarning(
                "Aucune donnée",
                "Aucune police unique à exporter."
            )
            return
        
        self.export_dataframe(
            self.no_duplicates_df,
            "Export des Polices Uniques",
            "polices_uniques.xlsx"
        )
    
    def export_dataframe(self, df, title, default_filename):
        """Exporte un DataFrame vers Excel"""
        try:
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                initialfile=default_filename,
                filetypes=[("Fichiers Excel", "*.xlsx")],
                title=title
            )
            
            if file_path:
                df.to_excel(file_path, index=False)
                messagebox.showinfo(
                    "Exportation réussie",
                    f"Les données ont été exportées avec succès vers :\n{file_path}"
                )
                self.update_status(f"Données exportées : {file_path}")
        
        except Exception as e:
            messagebox.showerror(
                "Erreur d'exportation",
                f"Une erreur est survenue lors de l'exportation :\n{str(e)}"
            )
            self.update_status("Erreur lors de l'exportation")

class RoundedButton_1(tk.Canvas):
    def __init__(self, parent, text, command=None, radius=20, **kwargs):
        tk.Canvas.__init__(self, parent, height=radius*2, width=150, bg=parent.cget("bg"), highlightthickness=0)

        self.command = command
        self.original_color = "#009640"  # Couleur d'origine 007bff bleu claire
        self.clicked_color = "#0056b3"    # Couleur après clic

        # Dessiner un bouton arrondi avec des tags
        self.create_oval((0, 0, radius*2, radius*2), fill=self.original_color, outline="", tags="button")
        self.create_oval((150-radius*2, 0, 150, radius*2), fill=self.original_color, outline="", tags="button")
        self.create_rectangle((radius, 0, 150-radius, radius*2), fill=self.original_color, outline="", tags="button")

        # Ajouter le texte
        self.text = self.create_text(75, radius, text=text, fill="white", font=("Arial", 10, "bold"))

        # Ajouter les événements pour cliquer
        #self.bind("<Button-1>", self._on_click)
        self.tag_bind(self.text, "<Button-1>", self._on_click)

    def _on_click(self, event):
        if self.command:
            self.change_color()
            self.command()

    def change_color(self):
        # Changer la couleur du bouton
        self._set_button_color(self.clicked_color)

        # Revenir à la couleur d'origine après 2 secondes
        self.after(2000, self.reset_color)

    def _set_button_color(self, color):
        # Mettre à jour la couleur du bouton
        self.itemconfig("button", fill=color)  # Utiliser le tag "button"

    def reset_color(self):
        # Revenir à la couleur d'origine
        self._set_button_color(self.original_color)
        self.itemconfig(self.text, fill="white")  # Remettre la couleur du texte

class RoundedButton_3(tk.Canvas):
    def __init__(self, parent, text, command=None, radius=20, **kwargs):
        tk.Canvas.__init__(self, parent, height=radius*2, width=150, bg=parent.cget("bg"), highlightthickness=0)

        self.command = command
        self.original_color = "#FF0000"  # Couleur d'origine
        self.clicked_color = "#0056b3"    # Couleur après clic

        # Dessiner un bouton arrondi avec des tags
        self.create_oval((0, 0, radius*2, radius*2), fill=self.original_color, outline="", tags="button")
        self.create_oval((150-radius*2, 0, 150, radius*2), fill=self.original_color, outline="", tags="button")
        self.create_rectangle((radius, 0, 150-radius, radius*2), fill=self.original_color, outline="", tags="button")

        # Ajouter le texte
        self.text = self.create_text(75, radius, text=text, fill="white", font=("Arial", 10, "bold"))

        # Ajouter les événements pour cliquer
        #self.bind("<Button-1>", self._on_click)
        self.tag_bind(self.text, "<Button-1>", self._on_click)

    def _on_click(self, event):
        if self.command:
            self.change_color()
            self.command()

    def change_color(self):
        # Changer la couleur du bouton
        self._set_button_color(self.clicked_color)

        # Revenir à la couleur d'origine après 2 secondes
        self.after(2000, self.reset_color)

    def _set_button_color(self, color):
        # Mettre à jour la couleur du bouton
        self.itemconfig("button", fill=color)  # Utiliser le tag "button"

    def reset_color(self):
        # Revenir à la couleur d'origine
        self._set_button_color(self.original_color)
        self.itemconfig(self.text, fill="white")  # Remettre la couleur du texte

# Classe pour les boutons arrondis
class RoundedButton(tk.Canvas):
    def __init__(self, master, text, command, width=100, height=30, corner_radius=10, bg="#3498db", fg="white", hover_bg="#2980b9", **kwargs):
        super().__init__(master, width=width, height=height, highlightthickness=0, bg=master['bg'], **kwargs)
        self.command = command
        self.bg = bg
        self.fg = fg
        self.hover_bg = hover_bg
        self.current_bg = bg
        
        self.bind("<Enter>", self.on_enter)
        self.bind("<Leave>", self.on_leave)
        self.bind("<Button-1>", self.on_click)
        
        self.draw_button(text)
        
    def draw_button(self, text):
        self.delete("all")
        self.create_rounded_rect(0, 0, self.winfo_reqwidth(), self.winfo_reqheight(), 
                                radius=10, fill=self.current_bg)
        self.create_text(self.winfo_reqwidth()/2, self.winfo_reqheight()/2, 
                         text=text, fill=self.fg, font=("Helvetica", 10, "bold"))
    
    def create_rounded_rect(self, x1, y1, x2, y2, radius=10, **kwargs):
        points = [x1+radius, y1,
                 x2-radius, y1,
                 x2, y1,
                 x2, y1+radius,
                 x2, y2-radius,
                 x2, y2,
                 x2-radius, y2,
                 x1+radius, y2,
                 x1, y2,
                 x1, y2-radius,
                 x1, y1+radius,
                 x1, y1]
        return self.create_polygon(points, **kwargs, smooth=True)
    
    def on_enter(self, event):
        self.current_bg = self.hover_bg
        self.draw_button(self.itemcget("text", "text"))
        
    def on_leave(self, event):
        self.current_bg = self.bg
        self.draw_button(self.itemcget("text", "text"))
        
    def on_click(self, event):
        self.command()

class DashboardApp():
    THEMES = {
        "Light": {
            "bg": "#000033",
            "fg": "white",
            "button_bg": "#4CAF50",
            "button_fg": "white",
            "button_hover": "#45a049",
            "tab_bg": "#9999FF",
            "tab_fg": "black",
            "menu_bg": "#d9d9d9",
            "menu_fg": "black"
        },
    }

    def __init__(self, root):
        self.root = root
        self.root.title("Tableau de Bord - AGC-VIE")
        self.root.geometry("1366x800")

        self.password = "secret"
        self.password_attempts = 0
        self.max_attempts = 5
        
        # Variables pour le thème
        self.current_theme = "Light"
        self.apply_theme(self.current_theme)
        
        self.previous_data = None
        self.tech_table = None
        self.compte_table = None
        self.pivot_comptables = pd.DataFrame()
        self.pivot_techniques = pd.DataFrame()

        self.setup_ui()
        
    def setup_ui(self):
        # Créer un cadre pour le menu vertical (à gauche)
        self.left_frame = GradientFrame(self.root, color1="#ffffff", color2="#ffffff", width=250) #2a5298 1e3c72 1e3c72 009640
        self.left_frame.pack(side="left", fill="y")

        # Charger les icônes
        self.load_icons()

        # Profil utilisateur
        self.user_profile()

        # Menu vertical avec listes déroulantes
        self.vertical_menu()
        
        # Barre de menu horizontale
        self.horizontal_menu()

        # Création du notebook pour les onglets de gestion
        self.notebook = ttk.Notebook(self.root, style="Custom.TNotebook")
        self.notebook.pack(pady=10, expand=True, side="right", fill="both")

        # Création des onglets
        self.tab_technique = tk.Frame(self.notebook, bg=self.theme['tab_bg'])
        self.tab_comptable = tk.Frame(self.notebook, bg=self.theme['tab_bg'])

        self.notebook.add(self.tab_technique, text="Gestion Technique")
        self.notebook.add(self.tab_comptable, text="Gestion Comptable")

        # Contenu des onglets
        self.create_technique_tab()
        self.create_comptable_tab()
        
        # Configurer le style
        self.configure_styles()

    def configure_styles(self):
        style = ttk.Style()
        
        # Style pour le Notebook
        style.configure("Custom.TNotebook", background=self.theme['fg'])
        style.configure("Custom.TNotebook.Tab", 
                       background=self.theme['menu_bg'],
                       foreground=self.theme['menu_fg'],
                       padding=[10, 5],
                       font=('Helvetica', 10, 'bold'))
        style.map("Custom.TNotebook.Tab",
                  background=[("selected", self.theme['button_bg'])],
                  foreground=[("selected", "#009640")])
        
        # Style pour les Treeview
        style.configure("Treeview",
                        background=self.theme['fg'],
                        foreground=self.theme['menu_fg'],
                        fieldbackground=self.theme['bg'])
        style.configure("Treeview.Heading",
                       background=self.theme['bg'],
                       foreground="black",
                       font=('Helvetica', 9, 'bold'))
        style.map("Treeview",
                  background=[('selected', self.theme['button_hover'])],
                  foreground=[('selected', 'white')])
    
    def apply_theme(self, theme_name):
        self.current_theme = theme_name
        self.theme = self.THEMES[theme_name]
        
        if hasattr(self, 'root'):
            self.root.config(bg=self.theme['bg'])
            
            # Reconfigurer les éléments existants
            if hasattr(self, 'left_frame'):
                self.left_frame.config(bg=self.theme['bg'])
                
            if hasattr(self, 'notebook'):
                self.notebook.config(style="Custom.TNotebook")
                
            if hasattr(self, 'tab_technique'):
                self.tab_technique.config(bg=self.theme['tab_bg'])
                self.tab_comptable.config(bg=self.theme['tab_bg'])
                
            self.configure_styles()
            
            # Mettre à jour les boutons du menu vertical
            if hasattr(self, 'left_frame'):
                for widget in self.left_frame.winfo_children():
                    if isinstance(widget, tk.Button):
                        widget.config(bg=self.theme['fg'], fg=self.theme['fg'])
    
    def load_icons(self):
        # Récupérer le chemin du dossier où se trouve le script
        BASE_DIR = os.path.dirname(os.path.abspath(__file__))

        # Construire le chemin absolu de l'image
        image_path = os.path.join(BASE_DIR, "images", "1.png")
        image_path_1 = os.path.join(BASE_DIR, "images", "3.png")
        image_path_2 = os.path.join(BASE_DIR, "images", "4.png")
        image_path_3 = os.path.join(BASE_DIR, "images", "4.png")
        image_path_4 = os.path.join(BASE_DIR, "images", "5.png")

        # Vérifier si l'image existe
        if not os.path.exists(image_path):
            print(f"⚠️ Erreur : L'image {image_path} est introuvable.")
        else:
            self.icon_410_tech_compta = ImageTk.PhotoImage(Image.open(image_path).resize((20, 20)))
            self.icon_compte_410_411 = ImageTk.PhotoImage(Image.open(image_path_1).resize((20, 20)))
            self.icon_gestion_doublon = ImageTk.PhotoImage(Image.open(image_path_2).resize((20, 20)))
            self.icon_gestion_production = ImageTk.PhotoImage(Image.open(image_path_4).resize((20, 20)))
            self.icon_stats = ImageTk.PhotoImage(Image.open(image_path_3).resize((20, 20)))

    def user_profile(self):
        # Récupérer le chemin du dossier où se trouve le script
        BASE_DIR = os.path.dirname(os.path.abspath(__file__))

        # Construire le chemin absolu de l'image
        image_path = os.path.join(BASE_DIR, "images", "logo_1.jpg")

        # Vérifier si l'image existe
        if not os.path.exists(image_path):
            print(f"⚠️ Erreur : L'image {image_path} est introuvable.")
        else:
            self.user_image = Image.open(image_path).resize((80, 80))
            self.user_photo = ImageTk.PhotoImage(self.user_image)
        
        user_photo_label = tk.Label(self.left_frame, image=self.user_photo, bg=self.theme['fg'])
        user_photo_label.pack(pady=20)
        
        # Ajouter un label pour le nom d'utilisateur
        tk.Label(self.left_frame, text="AGC-VIE", bg=self.theme['fg'], fg=self.theme['bg'], 
                font=("Helvetica", 12, "bold")).pack()

    def ask_password(self, callback, menu_name):
        """Boîte de dialogue de mot de passe avec un design amélioré"""
        self.password_window = tk.Toplevel()
        self.password_window.title(f"Accès sécurisé - {menu_name}")
        self.password_window.resizable(False, False)
        self.password_window.grab_set()  # Rend la fenêtre modale
        
        # Centrer la fenêtre
        window_width = 350
        window_height = 200
        screen_width = self.password_window.winfo_screenwidth()
        screen_height = self.password_window.winfo_screenheight()
        x = (screen_width // 2) - (window_width // 2)
        y = (screen_height // 2) - (window_height // 2)
        self.password_window.geometry(f"{window_width}x{window_height}+{x}+{y}")
        
        # Style
        bg_color = "#f0f0f0"
        accent_color = "#4a6fa5"
        self.password_window.configure(bg=bg_color)
        
        # Icône (remplacer par votre propre icône si disponible)
        try:
            self.password_window.iconbitmap('shield.ico')  # Icône de sécurité
        except:
            pass
        
        # Contenu
        tk.Label(
            self.password_window, 
            text=f"Accès à {menu_name}", 
            font=('Helvetica', 12, 'bold'), 
            bg=bg_color
        ).pack(pady=(15, 5))
        
        tk.Label(
            self.password_window, 
            text="Cette section est protégée. Veuillez entrer le mot de passe :", 
            font=('Helvetica', 9), 
            bg=bg_color, 
            wraplength=300
        ).pack(pady=(0, 15))
        
        # Champ de mot de passe
        password_frame = tk.Frame(self.password_window, bg=bg_color)
        password_frame.pack()
        
        self.password_entry = ttk.Entry(
            password_frame, 
            show='•', 
            font=('Helvetica', 10), 
            width=25
        )
        self.password_entry.pack(pady=5)
        self.password_entry.focus_set()
        
        # Boutons
        button_frame = tk.Frame(self.password_window, bg=bg_color)
        button_frame.pack(pady=15)
        
        ttk.Button(
            button_frame, 
            text="Valider", 
            style='Accent.TButton', 
            command=lambda: self.check_password(callback, menu_name)  # Modification ici
        ).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            button_frame, 
            text="Annuler", 
            command=self.password_window.destroy
        ).pack(side=tk.LEFT, padx=5)
        
        # Lier la touche Entrée au bouton Valider
        self.password_window.bind('<Return>', lambda e: self.check_password(callback))
        
        # Style ttk
        style = ttk.Style()
        style.configure('Accent.TButton', foreground='white', background=accent_color)
    
    def check_password(self, callback, menu_name):
        """Vérifie le mot de passe et exécute le callback si correct"""
        entered_password = self.password_entry.get()
    
        if entered_password == self.password:
            self.password_window.destroy()
            if callable(callback):  # Vérification importante
                callback()  # Appel de la fonction seulement si c'est callable
            else:
                messagebox.showerror("Erreur", f"Action invalide pour {menu_name}")
            self.password_attempts = 0
        else:
            self.password_attempts += 1
            remaining_attempts = self.max_attempts - self.password_attempts
            
            if remaining_attempts > 0:
                messagebox.showwarning(
                    "Accès refusé",
                    f"Mot de passe incorrect. Il vous reste {remaining_attempts} tentative(s).",
                    parent=self.password_window
                )
                self.password_entry.delete(0, tk.END)
                self.password_entry.focus_set()
            else:
                messagebox.showerror(
                    "Accès bloqué",
                    "Nombre maximal de tentatives atteint. Accès refusé.",
                    parent=self.password_window
                )
                self.password_window.destroy()


    def vertical_menu(self):
        # Menu vertical avec des menus déroulants
        menu_items = [
            ("GESTION TECHNIQUE", [
                ("Gestion Doublons", lambda: self.ask_password(self.open_gestion_doublons, "Gestion Doublons")),
                ("Gestion 410", lambda: self.ask_password(self.open_compte_410_tech_compta, "Gestion 410 Technique"))
            ]),
            ("GESTION COMPTABLES", [
                ("Gestion 410 & 411", lambda: self.ask_password(self.open_compte_410_411, "Gestion 410 & 411"))
            ]),
            ("GESTION PRODUCTION", [
                ("Gestion Production", lambda: self.ask_password(self.open_gestion_production, "Gestion Production"))
            ]),
            ("STATISTIQUES", [
                ("Statistiques", lambda: self.ask_password(self.show_statistics, "Statistiques"))
            ])
        ]
        
        for title, items in menu_items:
            dropdown = DropdownMenu(
                self.left_frame,
                title=title,
                items=items,
                bg=self.theme['bg']
            )
            dropdown.pack(fill="x", pady=5, padx=10)
    
    def horizontal_menu(self):
        # Barre de menu horizontale (en haut)
        menubar = tk.Menu(self.root, bg=self.theme['menu_bg'], fg=self.theme['menu_fg'])
        self.root.config(menu=menubar)
        
        # Menu Fichier
        file_menu = tk.Menu(menubar, tearoff=0, bg=self.theme['menu_bg'], fg=self.theme['menu_fg'])
        file_menu.add_command(label="Importer des données techniques", command=self.load_tech_data)
        file_menu.add_command(label="Importer des données comptables", command=self.load_compte_data)   
        file_menu.add_separator()
        file_menu.add_command(label="Quitter", command=self.root.quit)
        menubar.add_cascade(label="Fichier", menu=file_menu)

        # Menu Rapprochement
        rapprochement_menu = tk.Menu(menubar, tearoff=0, bg=self.theme['menu_bg'], fg=self.theme['menu_fg'])
        rapprochement_menu.add_command(label="Rapprochement Technique", command=self.rapprochement_technique)
        rapprochement_menu.add_command(label="Rapprochement Comptable", command=self.rapprochement_comptable)
        menubar.add_cascade(label="Rapprochement", menu=rapprochement_menu)

        # Menu Aide
        help_menu = tk.Menu(menubar, tearoff=0, bg=self.theme['menu_bg'], fg=self.theme['menu_fg'])
        help_menu.add_command(label="À propos", command=self.show_about)
        menubar.add_cascade(label="Aide", menu=help_menu)
    
        
    # Tableau des données techniques 
    def create_technique_tab(self):
        tk.Label(self.tab_technique, text="Technique", font=("Helvetica", 14), bg=self.theme['tab_bg'], fg=self.theme['fg']).pack(pady=5)

        # Créer le tableau technique
        self.tech_table = ttk.Treeview(self.tab_technique)
        
        # Barre de recherche
        search_frame_1 = tk.Frame(self.tab_technique)
        search_frame_1.pack(pady=10)

        tk.Label(search_frame_1, text="Rechercher:").pack(side="left")
        self.search_entry_1 = tk.Entry(search_frame_1)
        self.search_entry_1.pack(side="left")
        tk.Button(search_frame_1, text="Chercher", command=self.search_techniques).pack(side="left")

        # Frame pour les boutons
        button_frame = tk.Frame(self.tab_technique)
        button_frame.pack(side="bottom", pady=15, fill="x")

        # Ajouter les boutons arrondis

        RoundedButton_1(button_frame, "Exporter", self.export_filtered_data).pack(side="left", padx=200, pady=10)
        RoundedButton_3(button_frame, "Supprimer", self.delete_selected_row).pack(side="left", padx=100, pady=10)

        # Créer un cadre pour le tableau et la barre de défilement
        frame = tk.Frame(self.tab_technique)
        frame.pack(pady=10, fill="both", expand=True)

        # Créer une barre de défilement verticale
        scrollbar = ttk.Scrollbar(frame, orient="vertical")
        scrollbar.pack(side="right", fill="y")

        # Créer une barre de défilement horizontale
        scrollbar_x = ttk.Scrollbar(frame, orient="horizontal")
        scrollbar_x.pack(side="bottom", fill="x")

        # Tableau pour les données techniques
        columns = ("Nouvelle_Police", "Emissions", "Ristournes", "Chiffre affaire", "Etat Quittance", "Reliquat A Encaisser")
        self.tech_table = ttk.Treeview(frame, columns=columns, show="headings", height=30, yscrollcommand=scrollbar.set, xscrollcommand=scrollbar_x.set)
        scrollbar.config(command=self.tech_table.yview)
        scrollbar_x.config(command=self.tech_table.xview)

        for col in columns:
            self.tech_table.heading(col, text=col)
            self.tech_table.column(col, anchor="center")
        self.tech_table.pack(side="left", fill="both", expand=True)

    def search_techniques(self):
        search_term_1 = self.search_entry_1.get().lower()
        for row in self.tech_table.get_children():
            self.tech_table.delete(row)
        for _, row in pivot_techniques.iterrows():
            if search_term_1 in row['Nouvelle_Police'].lower():
                self.tech_table.insert("", "end", values=(row["Nouvelle_Police"], row["Emissions"], row["Ristournes"], row["Chiffre affaire"], row["Etat Quittance"], row["Reliquat A Encaisser"]))
    
    def delete_selected_row(self):
        selected_item = self.tech_table.selection()
        if selected_item:
            self.tech_table.delete(selected_item)
        else:
            messagebox.showwarning("Avertissement", "Veuillez sélectionner une ligne à supprimer.")

    def export_filtered_data(self):
        try:
            file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Fichiers Excel", "*.xlsx"), ("Fichiers CSV", "*.csv")])
            if file_path:
                filtered_data = [self.tech_table.item(item, 'values') for item in self.tech_table.get_children()]
                df = pd.DataFrame(filtered_data, columns=self.tech_table["columns"])
                df.to_excel(file_path, index=False)
                messagebox.showinfo("Succès", "Données filtrées exportées avec succès.")
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de l'exportation des données : {e}")

    def load_tech_data(self):
        try:
            global pivot_techniques
            file_path = filedialog.askopenfilename(filetypes=[("Fichiers Excel", "*.xlsx *.xls")])
            if file_path:
                # Création de la fenêtre de chargement
                loading_window = tk.Toplevel(self.root)
                loading_window.title("Chargement")
                loading_label = tk.Label(loading_window, text="Importation des données techniques en cours...", padx=20, pady=20)
                loading_label.pack()
                self.root.update()
                df = pd.read_excel(file_path)

                # Ajout de la colonne "Nouvelle_Police" si nécessaire
                if 'Num avenant' in df.columns and 'Code intermédiaire' in df.columns and 'N° police' in df.columns:
                    df['Nouvelle_Police'] = df.apply(
                        lambda row: f"{row['Code intermédiaire']}-{row['N° police']}/{row['Num avenant']}" 
                        if pd.notnull(row['Num avenant']) 
                        else f"{row['Code intermédiaire']}-{row['N° police']}", axis=1
                    )

                # Convertir la colonne en chaîne et enlever ".0"
                df['Nouvelle_Police'] = df['Nouvelle_Police'].astype(str).str.replace('.0', '')

                # Calcul des Ristournes, Emissions, et Totale
                df['Ristournes'] = df.apply(lambda row: row['Chiffre affaire'] if row['Type quittance'] == 'Ristourne' else 0, axis=1)
                df['Emissions'] = df.apply(lambda row: row['Chiffre affaire'] if row['Type quittance'] == 'Emission' else 0, axis=1)
                #df['Chiffre affaire'] = df['Emissions'].fillna(0) + df['Ristournes'].fillna(0)

                pivot_df = pd.pivot_table(df, index=['Nouvelle_Police'], values=['Emissions', 'Ristournes', 'Chiffre affaire', 'Etat Quittance', 'Reliquat A Encaisser'], aggfunc='sum', fill_value=0).reset_index()
                pivot_techniques = pivot_df
                self.pivot_tech = df

            # Insertion des données dans le tableau technique
            self.tech_table.delete(*self.tech_table.get_children())
            for _, row in pivot_df.iterrows():
                self.tech_table.insert("", "end", values=(row["Nouvelle_Police"], row["Emissions"], row["Ristournes"], row["Chiffre affaire"], row["Etat Quittance"], row["Reliquat A Encaisser"]))

            loading_window.destroy()
            messagebox.showinfo("Succès", "Tableau des données techniques généré avec succès")
            return pivot_techniques

        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de l'importation des données techniques : {e}")

    def create_comptable_tab(self):
        tk.Label(self.tab_comptable, text="Comptabilité", font=("Helvetica", 14), bg=self.theme['tab_bg'], fg=self.theme['fg']).pack(pady=5)
    
        # Barre de recherche
        search_frame = tk.Frame(self.tab_comptable)
        search_frame.pack(pady=10)

        tk.Label(search_frame, text="Rechercher:").pack(side="left")
        self.search_entry = tk.Entry(search_frame)
        self.search_entry.pack(side="left")
        tk.Button(search_frame, text="Chercher", command=self.search_comptable).pack(side="left")

        # Frame pour les boutons
        button_frame = tk.Frame(self.tab_comptable)
        button_frame.pack(side="bottom", pady=15, fill="x")

        # Ajouter les boutons arrondis
        RoundedButton_3(button_frame, "Supprimer", self.delete_selected_row_comptable).pack(side="left", padx=100, pady=10)
        RoundedButton_1(button_frame, "Exporter", self.export_filtered_data_comptable).pack(side="left", padx=200, pady=10)

        # Créer un cadre pour le tableau et la barre de défilement
        frame = tk.Frame(self.tab_comptable)
        frame.pack(pady=10, fill="both", expand=True)

        # Créer une barre de défilement verticale
        scrollbar = ttk.Scrollbar(frame, orient="vertical")
        scrollbar.pack(side="right", fill="y")

        # Créer une barre de défilement horizontal
        scrollbar_x = ttk.Scrollbar(frame, orient="horizontal")
        scrollbar_x.pack(side="bottom", fill="x")

        # Tableau pour les données Comptable
        columns = ("No Police", "Débit", "Crédit")
        self.compte_table = ttk.Treeview(frame, columns=columns, show="headings", height=27, yscrollcommand=scrollbar.set, xscrollcommand=scrollbar_x.set)
        scrollbar.config(command=self.compte_table.yview)
        scrollbar_x.config(command=self.compte_table.xview)

        for col in columns:
            self.compte_table.heading(col, text=col)
            self.compte_table.column(col, anchor="center")
        self.compte_table.pack(side="left", fill="both", expand=True)

    def search_comptable(self):
        search_term = self.search_entry.get().lower()
        for row in self.compte_table.get_children():
            self.compte_table.delete(row)
        for _, row in pivot_comptables.iterrows():
            if search_term in row['No Police'].lower():
                self.compte_table.insert("", "end", values=(row["No Police"], row["Débit"], row["Crédit"]))

    def delete_selected_row_comptable(self):
        selected_item = self.compte_table.selection()
        if selected_item:
            self.compte_table.delete(selected_item)
        else:
            messagebox.showwarning("Avertissement", "Veuillez sélectionner une ligne à supprimer.")

    def export_filtered_data_comptable(self):
        try:
            file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Fichiers Excel", "*.xlsx"), ("Fichiers CSV", "*.csv")])
            if file_path:
                filtered_data = [self.compte_table.item(item, 'values') for item in self.compte_table.get_children()]
                df = pd.DataFrame(filtered_data, columns=self.compte_table["columns"])
                df.to_excel(file_path, index=False)
                messagebox.showinfo("Succès", "Données filtrées exportées avec succès.")
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de l'exportation des données : {e}")
        
    def load_compte_data(self):
        try:
            global pivot_comptables
            file_path = filedialog.askopenfilename(filetypes=[("Fichiers Excel", "*.xlsx *.xls")])
            if file_path:
                # Création de la fenêtre de chargement
                loading_window = tk.Toplevel(self.root)
                loading_window.title("Chargement")
                loading_label = tk.Label(loading_window, text="Importation des données comptables en cours...", padx=20, pady=20)
                loading_label.pack()
                self.root.update()
                df = pd.read_excel(file_path)

                # Convertir la colonne en chaîne et enlever ".0"
                df['No Police'] = df['No Police'].astype(str).str.replace('.0', '')

                # Tableau croisé dynamique pour Données Comptables (à adapter selon les colonnes disponibles)
                pivot_df = pd.pivot_table(df, index=['No Police'], values=['Débit', 'Crédit'], aggfunc='sum', fill_value=0).reset_index()  # Ex: somme des montants par Compte
                pivot_comptables = pivot_df
            
            # Insertion des données dans le tableau technique
            self.compte_table.delete(*self.compte_table.get_children())
            for _, row in pivot_df.iterrows():
                self.compte_table.insert("", "end", values=(row["No Police"], row["Débit"], row["Crédit"]))

            loading_window.destroy()
            messagebox.showinfo("Succès", "Tableau des données comptables généré avec succès")
            return pivot_comptables
        
        except Exception as e:
            messagebox.showerror("Erreur", f"Vous n'avez pas encore importer les Données Comptables {e}")


    def navigate(self, section):
        messagebox.showinfo("Navigation", f"Vous avez cliqué sur '{section}'")

    def show_about(self):
        messagebox.showinfo("À propos", "Tableau de bord de gestion technique et comptable\nVersion Finale")

    def rapprochement_technique(self):
        try:
            global pivot_techniques, pivot_comptables
            
            # Fenêtre de chargement - Style minimaliste
            loading_window = tk.Toplevel(self.root)
            loading_window.title("Chargement")
            loading_window.geometry("400x100")  # Plus compact
            loading_window.resizable(False, False)
            
            # Style épuré
            loading_window.configure(bg='#f5f5f5')  # Fond clair
            loading_label = tk.Label(loading_window, 
                                text="Rapprochement technique en cours...", 
                                font=("Helvetica", 12, "normal"),  # Police plus légère
                                bg='#f5f5f5', 
                                fg='#333333')  # Texte foncé sur fond clair
            loading_label.pack(pady=10)
            
            # Barre de progression simplifiée
            progress = ttk.Progressbar(loading_window, 
                                    orient="horizontal", 
                                    length=250, 
                                    mode="indeterminate",
                                    style="grey.Horizontal.TProgressbar")
            progress.pack()
            progress.start()
            
            self.root.update()

            # Fonction pour insérer Débit et Crédit dans la Technique
            def recuperer_debit_credit(pivot_techniques, pivot_comptables):
                pivot_comptables.columns = pivot_comptables.columns.str.strip()

                if 'No Police' not in pivot_comptables.columns:
                    print("Erreur : 'No Police' n'est pas trouvé dans les colonnes")
                    return

                pivot_techniques['Débit'] = 'Introuvé'
                pivot_techniques['Crédit'] = 'Introuvé'
                
                for index, row in pivot_techniques.iterrows():
                    nouvelle_police = row['Nouvelle_Police']
                    correspondance = pivot_comptables[pivot_comptables['No Police'] == nouvelle_police]
                    
                    if not correspondance.empty:
                        pivot_techniques.at[index, 'Débit'] = correspondance['Débit'].values[0]
                        pivot_techniques.at[index, 'Crédit'] = correspondance['Crédit'].values[0]
                    else:
                        pivot_techniques.at[index, 'Débit'] = 'Introuvé'
                        pivot_techniques.at[index, 'Crédit'] = 'Introuvé'
                
                return pivot_techniques

            pivot_techniques = recuperer_debit_credit(pivot_techniques, pivot_comptables)

            # Fonction pour la vérification
            def verifier_polices(pivot_techniques):
                pivot_techniques['Crédit'] = pd.to_numeric(pivot_techniques['Crédit'], errors='coerce')
                pivot_techniques['Débit'] = pd.to_numeric(pivot_techniques['Débit'], errors='coerce')
                pivot_techniques['Emissions'] = pd.to_numeric(pivot_techniques['Emissions'], errors='coerce')
                pivot_techniques['Ristournes'] = pd.to_numeric(pivot_techniques['Ristournes'], errors='coerce')

                # Calcul du chiffre d'affaire
                total_emissions = pivot_techniques['Emissions'].sum()
                total_ristournes = pivot_techniques['Ristournes'].sum()
                total_CA = abs(total_emissions + total_ristournes)

                # Calcul du CA comptable (somme des crédits - débits)
                total_credit_comptable = pivot_comptables['Crédit'].sum()
                total_debit_comptable = pivot_comptables['Débit'].sum()
                total_CA_comptable = abs(total_credit_comptable - total_debit_comptable)
                
                # Calcul de l'écart
                ecart = abs(total_CA - total_CA_comptable)

                pivot_techniques['Rapprochement'] = pivot_techniques.apply(
                    lambda row: 'Rapprochement réussi'
                    if abs((row['Emissions'] + row['Ristournes']) - abs(row['Crédit'] - row['Débit'])) == 0 else 'Rapprochement non réussi', axis=1)
                
                tableau_listing_police_invalide = pivot_techniques[pivot_techniques['Rapprochement'] == 'Rapprochement non réussi'][['Nouvelle_Police', 'Rapprochement']].copy()
                tableau_listing_valide = pivot_techniques[pivot_techniques['Rapprochement'] == 'Rapprochement réussi'][['Nouvelle_Police', 'Emissions', 'Ristournes', 'Rapprochement', 'Débit', 'Crédit']].copy()
                
                return tableau_listing_police_invalide, tableau_listing_valide, total_emissions, total_ristournes, total_CA, ecart, total_credit_comptable, total_debit_comptable

            tableau_listing_police_invalide, tableau_listing_valide, total_emissions, total_ristournes, total_CA, ecart, total_credit_comptable, total_debit_comptable = verifier_polices(pivot_techniques)

            # Fenêtre principale - Style simplifié
            top = tk.Toplevel(self.root)
            top.title("Rapprochement Technique")
            top.geometry("1100x750")  # Légèrement plus grand pour meilleure lisibilité
            top.configure(bg='#f5f5f5')  # Fond clair uniforme

            self.pivot_techniques = pd.DataFrame(pivot_techniques)
            self.tableau_listing_police_invalide = pd.DataFrame(tableau_listing_police_invalide)
            self.tableau_listing_valide = pd.DataFrame(tableau_listing_valide)
                
            # Style des onglets
            style = ttk.Style()
            #style.theme_use('clam')  # Thème plus moderne
            style.configure("TNotebook", background='#f5f5f5', borderwidth=0)
            style.configure("TNotebook.Tab", 
                         background='#e0e0e0', 
                         foreground='#333333',
                         padding=[10, 5],
                         font=('Helvetica', 10, 'bold'))
            style.map("TNotebook.Tab", 
                    background=[("selected", '#ffffff')],
                    expand=[("selected", [1, 1, 1, 0])])
            
            notebook = ttk.Notebook(top)
            notebook.pack(fill="both", expand=True, padx=5, pady=5)

            # Cadres des onglets
            self.onglet_techniques = tk.Frame(notebook, bg='#ffffff')
            self.onglet_invalide_tech = tk.Frame(notebook, bg='#ffffff')
            self.onglet_valide_tech = tk.Frame(notebook, bg='#ffffff')

            # Indicateurs CA - Style unifié
            ca_frame = tk.Frame(self.onglet_techniques, bg='#ffffff')
            ca_frame.pack(fill="x", pady=(10, 5), padx=10)
            
            indicator_style = {
                'font': ('Helvetica', 11),
                'bd': 0,
                'relief': 'flat',
                'padx': 8,
                'pady': 6,
                'highlightthickness': 1,
                'highlightbackground': '#e0e0e0'
            }
            
            # Indicateurs avec couleurs douces
            ca_emissions_label = tk.Label(ca_frame, 
                                        text=f"CA Emission: {total_emissions:,.1f} FCFA", 
                                        bg='#e8f5e9',  # Vert très clair
                                        fg='#2e7d32',  # Vert foncé
                                        **indicator_style)
            ca_emissions_label.pack(side="left", padx=5, fill="x", expand=True)
            
            ca_ristournes_label = tk.Label(ca_frame, 
                                        text=f"CA Ristourne: {total_ristournes:,.1f} FCFA", 
                                        bg='#e3f2fd',  # Bleu très clair
                                        fg='#1565c0',  # Bleu foncé
                                        **indicator_style)
            ca_ristournes_label.pack(side="left", padx=5, fill="x", expand=True)

            ca_net_label = tk.Label(ca_frame, 
                                text=f"CA Net: {total_CA:,.1f} FCFA", 
                                bg='#f3e5f5',  # Violet très clair
                                fg='#6a1b9a',  # Violet foncé
                                **indicator_style)
            ca_net_label.pack(side="left", padx=5, fill="x", expand=True)

            ecart_label = tk.Label(ca_frame, 
                                text=f"Écart: {ecart:,.1f} FCFA", 
                                bg='#ffebee' if ecart > 0 else '#e8f5e9',  # Rouge/vert très clair
                                fg='#c62828' if ecart > 0 else '#2e7d32',  # Rouge/vert foncé
                                **indicator_style)
            ecart_label.pack(side="left", padx=5, fill="x", expand=True)

            # Barres de recherche - Style cohérent
            def create_search_frame(parent):
                frame = tk.Frame(parent, bg='#ffffff')
                frame.pack(pady=5)
                
                tk.Label(frame, 
                        text="Recherche:", 
                        bg='#ffffff', 
                        fg='#333333',
                        font=('Helvetica', 10)).pack(side="left")
                
                entry = tk.Entry(frame, 
                                font=("Helvetica", 11),
                                relief='flat',
                                highlightthickness=1,
                                highlightbackground='#bdbdbd',
                                highlightcolor='#2196f3')
                entry.pack(side="left", padx=5)
                
                return frame, entry

            # Barre recherche principale
            search_frame, self.search_entry_tech = create_search_frame(self.onglet_techniques)
            search_btn = RoundedButton(search_frame, 
                                    "Chercher", 
                                    self.search_technique_1,
                                    width=90,
                                    height=25,
                                    bg='#2196f3',
                                    fg='white',
                                    hover_bg='#1976d2')
            search_btn.pack(side="left")

            # Bouton export - Style moderne
            export_frame = tk.Frame(self.onglet_techniques, bg='#ffffff')
            export_frame.pack(pady=5)
            
            export_btn = RoundedButton(export_frame, 
                                    "Exporter les Données", 
                                    self.exporter_donnees,
                                    width=140,
                                    height=30,
                                    bg='#ff9800',
                                    fg='white',
                                    hover_bg='#fb8c00')
            export_btn.pack()

            # Barres de recherche pour onglets secondaires
            search_frame_invalide, self.search_entry_tech_1 = create_search_frame(self.onglet_invalide_tech)
            search_btn_invalide = RoundedButton(search_frame_invalide, 
                                            "Chercher", 
                                            self.search_technique_invalide,
                                            width=90,
                                            height=25,
                                            bg='#2196f3',
                                            fg='white',
                                            hover_bg='#1976d2')
            search_btn_invalide.pack(side="left")

            search_frame_valide, self.search_entry_tech_2 = create_search_frame(self.onglet_valide_tech)
            search_btn_valide = RoundedButton(search_frame_valide, 
                                            "Chercher", 
                                            self.search_technique_valide,
                                            width=90,
                                            height=25,
                                            bg='#2196f3',
                                            fg='white',
                                            hover_bg='#1976d2')
            search_btn_valide.pack(side="left")

            # Ajout des onglets
            notebook.add(self.onglet_techniques, text="Données Complètes")
            notebook.add(self.onglet_invalide_tech, text="Polices Non Validées")
            notebook.add(self.onglet_valide_tech, text="Polices Validées")

            def configure_treeview(tree):
                # Style des tableaux
                style.configure("Treeview",
                            background='#ffffff',
                            foreground='#333333',
                            fieldbackground='#ffffff',
                            rowheight=28,
                            borderwidth=0,
                            highlightthickness=0)
                style.configure("Treeview.Heading",
                            background='#e0e0e0',
                            foreground='#333333',
                            font=('Helvetica', 10, 'normal'),
                            relief='flat')
                style.map("Treeview",
                        background=[('selected', '#e3f2fd')],  # Bleu très clair pour la sélection
                        foreground=[('selected', '#000000')])
                
                # Barres de défilement discrètes
                scroll_y = ttk.Scrollbar(tree, orient="vertical")
                scroll_x = ttk.Scrollbar(tree, orient="horizontal")
                
                tree.configure(yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)
                scroll_y.config(command=tree.yview)
                scroll_x.config(command=tree.xview)
                
                scroll_y.pack(side="right", fill="y")
                scroll_x.pack(side="bottom", fill="x")
                
                tree.pack(fill="both", expand=True, padx=10, pady=5)

            # Création des tableaux
            self.table_view_tech = ttk.Treeview(self.onglet_techniques, style="Custom.Treeview")
            self.table_view_tech.pack(fill="both", expand=True, padx=5, pady=5)
            configure_treeview(self.table_view_tech)
            
            self.table_view_invalide_tech = ttk.Treeview(self.onglet_invalide_tech, style="Custom.Treeview")
            self.table_view_invalide_tech.pack(fill="both", expand=True, padx=5, pady=5)
            configure_treeview(self.table_view_invalide_tech)
            
            self.table_view_valide_tech = ttk.Treeview(self.onglet_valide_tech, style="Custom.Treeview")
            self.table_view_valide_tech.pack(fill="both", expand=True, padx=5, pady=5)
            configure_treeview(self.table_view_valide_tech)

            # Afficher les données dans chaque tableau
            if pivot_techniques is not None:
                afficher_donnees(pivot_techniques, self.table_view_tech)
                
                # Ajouter une ligne de totaux au tableau principal
                if len(pivot_techniques) > 0:
                    # Créer une ligne synthétique pour les totaux
                    total_row = {
                        'Nouvelle_Police': 'TOTAUX',
                        'Chiffre affaire': f"{total_CA:,.1f}",
                        'Emissions': f"{total_emissions:,.1f}",
                        'Etat Quittance': '',
                        'Reliquat A Encaisser': '',
                        'Ristournes': f"{total_ristournes:,.1f}",
                        'Débit': f"{total_debit_comptable:,.1f}",
                        'Crédit': f"{total_credit_comptable:,.1f}",
                        'Rapprochement': f"{ecart:,.1f}",
                    }
                    
                    # Insérer la ligne de totaux
                    self.table_view_tech.insert("", "end", values=list(total_row.values()), tags=('total',))
                    
                    # Style de la ligne de totaux
                    self.table_view_tech.tag_configure('total', 
                                                    background='#e3f2fd',
                                                    font=('Helvetica', 10, 'bold'),
                                                    foreground='#0d47a1')

            if tableau_listing_police_invalide is not None:
                afficher_donnees(tableau_listing_police_invalide, self.table_view_invalide_tech)
            if tableau_listing_valide is not None:
                afficher_donnees(tableau_listing_valide, self.table_view_valide_tech)

            loading_window.destroy()
            messagebox.showinfo("Succès", "Rapprochement technique terminé.")
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors du rapprochement : {str(e)}")

    def exporter_donnees(self):
        try:
            with pd.ExcelWriter("Rapprochement_Technique.xlsx") as writer:
                self.pivot_techniques.to_excel(writer, sheet_name="Données Techniques", index=False)
                self.tableau_listing_police_invalide.to_excel(writer, sheet_name="Techniques Invalides", index=False)
                self.tableau_listing_valide.to_excel(writer, sheet_name="Techniques Valides", index=False)
            messagebox.showinfo("Succès", "Les données ont été exportées avec succès dans 'Rapprochement_Technique.xlsx'.")
        except Exception as e:
                messagebox.showerror("Erreur", f"Erreur lors de l'exportation : {e}")

    def search_technique_1(self):
        search_term = self.search_entry_tech.get().lower()
        for row in self.table_view_tech.get_children():
            self.table_view_tech.delete(row)
        for _, row in pivot_techniques.iterrows():
            if search_term in row['Nouvelle_Police'].lower():  # Assurez-vous que 'Nouvelle_Police' est une colonne dans pivot_techniques
                self.table_view_tech.insert("", "end", values=list(row))
    
    def search_technique_invalide(self):
        search_term_1 = self.search_entry_tech_1.get().lower()
        for row in self.table_view_invalide_tech.get_children():
            self.table_view_invalide_tech.delete(row)
        for _, row in self.tableau_listing_police_invalide.iterrows():
            if search_term_1 in row['Nouvelle_Police'].lower():  # Assurez-vous que 'Nouvelle_Police' est une colonne dans pivot_techniques
                self.table_view_invalide_tech.insert("", "end", values=list(row))

    def search_technique_valide(self):
        search_term_2 = self.search_entry_tech_2.get().lower()
        for row in self.table_view_valide_tech.get_children():
            self.table_view_valide_tech.delete(row)
        for _, row in self.tableau_listing_valide.iterrows():
            if search_term_2 in row['Nouvelle_Police'].lower():  # Assurez-vous que 'Nouvelle_Police' est une colonne dans pivot_techniques
                self.table_view_valide_tech.insert("", "end", values=list(row))
    
    def rapprochement_comptable(self):
        try: 
            global pivot_techniques, pivot_comptables
            
            # Fenêtre de chargement - Style minimaliste
            loading_window = tk.Toplevel(self.root)
            loading_window.title("Chargement")
            loading_window.geometry("400x100")  # Plus compact
            loading_window.resizable(False, False)
            loading_window.configure(bg='#f5f5f5')  # Fond clair
            
            # Style épuré pour le label
            loading_label = tk.Label(
                loading_window, 
                text="Rapprochement comptable en cours...", 
                font=("Helvetica", 12, "normal"), 
                bg='#f5f5f5', 
                fg='#333333'
            )
            loading_label.pack(pady=10)
            
            # Barre de progression simplifiée
            progress = ttk.Progressbar(
                loading_window, 
                orient="horizontal", 
                length=250, 
                mode="indeterminate",
                style="grey.Horizontal.TProgressbar"
            )
            progress.pack()
            progress.start()
            self.root.update()

            # Fonction pour calculer le chiffre d'affaire comptable
            def recuperer_annulations_emisions(pivot_techniques, pivot_comptables):
                pivot_techniques.columns = pivot_techniques.columns.str.strip()

                if 'Nouvelle_Police' not in pivot_techniques.columns:
                    print("Erreur : 'Nouvelle_Police' n'est pas trouvé dans les colonnes")

                pivot_comptables['Ristournes'] = 'Introuvé'
                pivot_comptables['Emissions'] = 'Introuvé'
                
                for index, row in pivot_comptables.iterrows():
                    nouvelle_police = row['No Police']
                    correspondance = pivot_techniques[pivot_techniques['Nouvelle_Police'] == nouvelle_police]
                    
                    if not correspondance.empty:
                        pivot_comptables.at[index, 'Ristournes'] = correspondance['Ristournes'].values[0]
                        pivot_comptables.at[index, 'Emissions'] = correspondance['Emissions'].values[0]
                    else:
                        pivot_comptables.at[index, 'Ristournes'] = 'Introuvé'
                        pivot_comptables.at[index, 'Emissions'] = 'Introuvé'
                
                return pivot_comptables

            pivot_comptables = recuperer_annulations_emisions(pivot_techniques, pivot_comptables)

            def verifier_polices_comptable(pivot_comptables):
                pivot_comptables['Crédit'] = pd.to_numeric(pivot_comptables['Crédit'], errors='coerce')
                pivot_comptables['Débit'] = pd.to_numeric(pivot_comptables['Débit'], errors='coerce')
                pivot_comptables['Emissions'] = pd.to_numeric(pivot_comptables['Emissions'], errors='coerce')
                pivot_comptables['Ristournes'] = pd.to_numeric(pivot_comptables['Ristournes'], errors='coerce')

                pivot_comptables['Rapprochement'] = pivot_comptables.apply(
                    lambda row: 'Rapprochement réussi' if abs((row['Crédit'] - row['Débit']) - (row['Emissions'] + row['Ristournes'])) == 0 
                            else 'Rapprochement non réussi', axis=1
                )
                
                # Calcul des totaux
                total_debit = pivot_comptables['Débit'].sum()
                total_credit = pivot_comptables['Crédit'].sum()
                total_ecart = pivot_comptables['Rapprochement'].sum()
                total_CA = abs(total_credit - total_debit)

                # Calcul du CA technique (somme emissions - ristournes)
                total_emissions_tech = pivot_techniques['Emissions'].sum()
                total_ristournes_tech = pivot_techniques['Ristournes'].sum()
                total_CA_technique = abs(total_emissions_tech + total_ristournes_tech)
                
                # Calcul de l'écart
                ecart = abs(total_CA_technique - total_CA)
                
                tableau_listing_police_invalide_comptable = pivot_comptables[
                    pivot_comptables['Rapprochement'] == 'Rapprochement non réussi'
                ][['No Police', 'Rapprochement']].copy()
                
                tableau_listing_valide_comptable = pivot_comptables[
                    pivot_comptables['Rapprochement'] == 'Rapprochement réussi'
                ][['No Police', 'Crédit', 'Débit', 'Rapprochement', 'Ristournes', 'Emissions']].copy()
                
                return tableau_listing_police_invalide_comptable, tableau_listing_valide_comptable, total_debit, total_credit, total_CA, ecart, total_ristournes_tech, total_emissions_tech, total_ecart

            tableau_listing_police_invalide_comptable, tableau_listing_valide_comptable, total_debit, total_credit, total_CA, ecart, total_emissions_tech, total_ristournes_tech, total_ecart = verifier_polices_comptable(pivot_comptables)

            # Fenêtre principale avec style clair
            top = tk.Toplevel(self.root)
            top.title("Rapprochement Comptable")
            top.geometry("1100x750")
            top.configure(bg='#f5f5f5')

            self.pivot_comptables = pd.DataFrame(pivot_comptables)
            self.tableau_listing_police_invalide_comptable = pd.DataFrame(tableau_listing_police_invalide_comptable)
            self.tableau_listing_valide_comptable = pd.DataFrame(tableau_listing_valide_comptable)

            # Style des onglets
            style = ttk.Style()
            style.configure("TNotebook", background='#f5f5f5', borderwidth=0)
            style.configure("TNotebook.Tab", 
                            background='#e0e0e0', 
                            foreground='#333333',
                            padding=[10, 5],
                            font=('Helvetica', 10, 'bold'))
            style.map("TNotebook.Tab", 
                    background=[("selected", '#ffffff')],
                    expand=[("selected", [1, 1, 1, 0])])

            notebook = ttk.Notebook(top)
            notebook.pack(fill="both", expand=True, padx=5, pady=5)

            # Onglets avec fond clair
            self.onglet_comptables = tk.Frame(notebook, bg='#ffffff')
            self.onglet_invalide_compt = tk.Frame(notebook, bg='#ffffff')
            self.onglet_valide_compt = tk.Frame(notebook, bg='#ffffff')
            
            notebook.add(self.onglet_comptables, text="Données Comptables")
            notebook.add(self.onglet_invalide_compt, text="Polices Invalides")
            notebook.add(self.onglet_valide_compt, text="Polices Validées")

            # Frame des totaux
            totals_frame = tk.Frame(self.onglet_comptables, bg='#ffffff')
            totals_frame.pack(fill="x", pady=(10, 5), padx=10)

            # Indicateurs de totaux
            indicator_style = {
                'font': ('Helvetica', 11),
                'bd': 0,
                'relief': 'flat',
                'padx': 8,
                'pady': 6,
                'highlightthickness': 1,
                'highlightbackground': '#e0e0e0'
            }
            
            tk.Label(totals_frame, text=f"Débit Total: {total_debit:,.1f} FCFA", 
                    bg='#e8f5e9', fg='#2e7d32', **indicator_style).pack(side="left", padx=5, fill="x", expand=True)
            tk.Label(totals_frame, text=f"Crédit Total: {total_credit:,.1f} FCFA", 
                    bg='#e3f2fd', fg='#1565c0', **indicator_style).pack(side="left", padx=5, fill="x", expand=True)
            tk.Label(totals_frame, text=f"CA Net: {total_CA:,.1f} FCFA", 
                    bg='#f3e5f5', fg='#6a1b9a', **indicator_style).pack(side="left", padx=5, fill="x", expand=True)
            tk.Label(totals_frame, text=f"Écart CA: {ecart:,.1f} FCFA", 
                    bg='#ffebee' if ecart > 0 else '#e8f5e9', 
                    fg='#c62828' if ecart > 0 else '#2e7d32', **indicator_style).pack(side="left", padx=5, fill="x", expand=True)

            # Barre de recherche simplifiée
            search_frame = tk.Frame(self.onglet_comptables, bg='#ffffff')
            search_frame.pack(pady=10)

            self.search_entry_compt = tk.Entry(search_frame, 
                                            font=("Helvetica", 11),
                                            relief='flat',
                                            bd=1,
                                            highlightthickness=1,
                                            highlightcolor='#4a6fa5',
                                            highlightbackground='#cccccc',
                                            width=30)
            self.search_entry_compt.pack(side="left", padx=5)
            
            search_btn = tk.Button(search_frame, 
                                text="Rechercher",
                                command=self.search_comptable_1,
                                bg='#4a6fa5',
                                fg='white',
                                bd=0,
                                padx=15,
                                pady=5,
                                font=('Helvetica', 9))
            search_btn.pack(side="left")

            # Barre de recherche simplifiée
            search_frame_1 = tk.Frame(self.onglet_invalide_compt, bg='#ffffff')
            search_frame_1.pack(pady=10)

            self.search_entry_compt_1 = tk.Entry(search_frame_1, 
                                            font=("Helvetica", 11),
                                            relief='flat',
                                            bd=1,
                                            highlightthickness=1,
                                            highlightcolor='#4a6fa5',
                                            highlightbackground='#cccccc',
                                            width=30)
            self.search_entry_compt_1.pack(side="left", padx=5)

            search_btn_1 = tk.Button(search_frame_1, 
                                text="Rechercher",
                                command=self.search_comptable_invalide,
                                bg='#4a6fa5',
                                fg='white',
                                bd=0,
                                padx=15,
                                pady=5,
                                font=('Helvetica', 9))
            search_btn_1.pack(side="left")

            # Barre de recherche simplifiée
            search_frame_2 = tk.Frame(self.onglet_valide_compt, bg='#ffffff')
            search_frame_2.pack(pady=10)

            self.search_entry_compt_2 = tk.Entry(search_frame_2, 
                                            font=("Helvetica", 11),
                                            relief='flat',
                                            bd=1,
                                            highlightthickness=1,
                                            highlightcolor='#4a6fa5',
                                            highlightbackground='#cccccc',
                                            width=30)
            self.search_entry_compt_2.pack(side="left", padx=5)

            search_btn_2 = tk.Button(search_frame_2, 
                                text="Rechercher",
                                command=self.search_comptable_valide,
                                bg='#4a6fa5',
                                fg='white',
                                bd=0,
                                padx=15,
                                pady=5,
                                font=('Helvetica', 9))
            search_btn_2.pack(side="left")

            # Bouton d'export style plat
            export_frame = tk.Frame(self.onglet_comptables, bg='#ffffff')
            export_frame.pack(pady=10)
            
            export_btn = tk.Button(export_frame, 
                                text="Exporter Excel",
                                command=self.exporter_donnees_comptable,
                                bg='#38a169',
                                fg='white',
                                bd=0,
                                padx=20,
                                pady=8,
                                font=('Helvetica', 10))
            export_btn.pack()

            # Configuration des Treeview - style minimal
            def configure_treeview(tree):
                style.configure("Treeview",
                              background='white',
                              foreground='#333333',
                              fieldbackground='white',
                              rowheight=28,
                              font=('Helvetica', 9))
                
                style.configure("Treeview.Heading",
                             background='#4a6fa5',
                             foreground='black',
                             font=('Helvetica', 10, 'bold'),
                             padding=5)
                
                style.map("Treeview",
                        background=[('selected', '#e2e8f0')],
                        foreground=[('selected', '#333333')])
                
                # Barres de défilement discrètes
                scroll_y = ttk.Scrollbar(tree, orient="vertical")
                
                tree.configure(yscrollcommand=scroll_y.set)
                scroll_y.config(command=tree.yview)
                scroll_y.pack(side="right", fill="y")
                
                tree.pack(fill="both", expand=True, padx=10, pady=5)

            # Tableau principal
            self.table_view_compt = ttk.Treeview(self.onglet_comptables, style="Custom.Treeview")
            self.table_view_compt.pack(fill="both", expand=True, padx=5, pady=5)
            configure_treeview(self.table_view_compt)
            
            # Tableau des invalides
            self.table_view_invalide_compt = ttk.Treeview(self.onglet_invalide_compt, style="Custom.Treeview")
            self.table_view_invalide_compt.pack(fill="both", expand=True, padx=5, pady=5)
            configure_treeview(self.table_view_invalide_compt)
            
            # Tableau des valides
            self.table_view_valide_compt = ttk.Treeview(self.onglet_valide_compt, style="Custom.Treeview")
            self.table_view_valide_compt.pack(fill="both", expand=True, padx=5, pady=5)
            configure_treeview(self.table_view_valide_compt)

            # Afficher les données dans chaque tableau
            if pivot_comptables is not None:
                afficher_donnees(pivot_comptables, self.table_view_compt)
                
                # Ajouter une ligne de totaux au tableau principal
                if len(pivot_comptables) > 0:
                    # Créer une ligne synthétique pour les totaux
                    total_row_1 = {
                        'No Police': 'TOTAUX',
                        'Crédit': f"{total_credit:,.1f}",
                        'Débit': f"{total_debit:,.1f}",
                        'Emissions': f"{total_emissions_tech:,.1f}",
                        'Ristournes': f"{total_ristournes_tech:,.1f}",
                        'Rapprochement': f"{ecart:,.1f}",
                    }
                    
                    # Insérer la ligne de totaux
                    self.table_view_compt.insert("", "end", values=list(total_row_1.values()), tags=('total_1',))
                    
                    # Style de la ligne de totaux
                    self.table_view_compt.tag_configure('total_1', 
                                                    background='#e3f2fd',
                                                    font=('Helvetica', 10, 'bold'),
                                                    foreground='#0d47a1')

            if tableau_listing_police_invalide_comptable is not None:
                afficher_donnees(tableau_listing_police_invalide_comptable, self.table_view_invalide_compt)
            if tableau_listing_valide_comptable is not None:
                afficher_donnees(tableau_listing_valide_comptable, self.table_view_valide_compt)

            loading_window.destroy()
            messagebox.showinfo("Succès", "Votre Rapprochement Comptable a été réalisé avec succès.")
            
        except Exception as e:
            messagebox.showerror("Erreur", f"Une erreur est survenue : {str(e)}")

    def exporter_donnees_comptable(self):
        try:
            with pd.ExcelWriter("Rapprochement_Comptables.xlsx") as writer:
                pivot_comptables.to_excel(writer, sheet_name="Données Comptables", index=False)
                self.tableau_listing_police_invalide_comptable.to_excel(writer, sheet_name="Comptables Invalides", index=False)
                self.tableau_listing_valide_comptable.to_excel(writer, sheet_name="Comptables Valides", index=False)
            messagebox.showinfo("Succès", "Les données ont été exportées avec succès dans 'Rapprochement_Comptables.xlsx'.")
        except Exception as e:
                messagebox.showerror("Erreur", f"Erreur lors de l'exportation : {e}")

    def search_comptable_1(self):
        search_term = self.search_entry_compt.get().lower()
        for row in self.table_view_compt.get_children():
            self.table_view_compt.delete(row)
        for _, row in pivot_comptables.iterrows():
            if search_term in row['No Police'].lower():  # Assurez-vous que 'No Police' est une colonne dans pivot_comptables
                self.table_view_compt.insert("", "end", values=list(row))

    def search_comptable_invalide(self):
        search_term_1 = self.search_entry_compt_1.get().lower()
        for row in self.table_view_invalide_compt.get_children():
            self.table_view_invalide_compt.delete(row)
        for _, row in self.tableau_listing_police_invalide_comptable.iterrows():
            if search_term_1 in row['No Police'].lower():  # Assurez-vous que 'No Police' est une colonne dans tableau_listing_police_invalide_comptable
                self.table_view_invalide_compt.insert("", "end", values=list(row))

    def search_comptable_valide(self):
        search_term_2 = self.search_entry_compt_2.get().lower()
        for row in self.table_view_valide_compt.get_children():
            self.table_view_valide_compt.delete(row)
        for _, row in self.tableau_listing_valide_comptable.iterrows():
            if search_term_2 in row['No Police'].lower():  # Assurez-vous que 'No Police' est une colonne dans tableau_listing_valide_comptable
                self.table_view_valide_compt.insert("", "end", values=list(row))

    def navigate(self, section):
        if section == "  Accueil":
            self.open_accueil()
        elif section == "  Gestion 410":
            self.open_compte_410_tech_compta()
        elif section == "  Gestion 410 & 411":
            self.open_compte_410_411()
        elif section == "  Gestion Doublons":
            self.open_gestion_doublons()
        elif section == "  Gestion Production":
            self.open_gestion_production()
        elif section == "  Statistiques":
            self.show_statistics()

    def open_gestion_production(self):
        CertificatApp()

    def open_accueil(self):
        top = tk.Toplevel(self.root)
        top.title("Accueil")
        top.geometry("600x400")
        top.configure(bg='#9999FF')
        tk.Label(top, text="Bienvenue sur le tableau de bord!", font=("Helvetica", 16), bg='#9999FF', fg='#ffffff').pack(pady=20)

        
    def open_compte_410_tech_compta(self):
        Gestion41_Tech(pivot_techniques, pivot_comptables)

    def open_compte_410_411(self):
        DataImporterApp()

    def open_gestion_doublons(self):
        Application_Analyse()

    def show_statistics(self):
        try:
            if self.pivot_tech is None:
                messagebox.showerror("Erreur", "Vous devez d'abord importer les données techniques.")
                return
            df = self.pivot_tech
            top = tk.Toplevel(self.root)
            top.title("Statistiques")
            top.geometry("900x600")
            
            category_counts = df['Libellé catégorie'].value_counts()
            top_clients = df['souscripteur'].value_counts().nlargest(5)

            high_usage = category_counts.nlargest(5)
            medium_usage = category_counts.iloc[5:10]
            low_usage = category_counts.nsmallest(5)
            
            tk.Label(top, text="Statistiques des Produits et Clients", font=("Helvetica", 16)).pack(pady=10)

            main_frame = tk.Frame(top)
            main_frame.pack(fill="both", expand=True, padx=30, pady=30)
            
            frame_high = tk.Frame(main_frame)
            frame_high.grid(row=0, column=0, padx=10)

            frame_medium = tk.Frame(main_frame)
            frame_medium.grid(row=0, column=1, padx=10)

            frame_low = tk.Frame(main_frame)
            frame_low.grid(row=0, column=2, padx=10)
            
            tk.Label(frame_high, text="Produits les plus utilisés", font=("Helvetica", 12, "bold")).pack(anchor="center")
            tk.Label(frame_medium, text="Produits moyennement utilisés", font=("Helvetica", 12, "bold")).pack(anchor="center")
            tk.Label(frame_low, text="Produits moins utilisés", font=("Helvetica", 12, "bold")).pack(anchor="center")

            for product, count in high_usage.items():
                tk.Label(frame_high, text=f"{product}: {count} utilisations").pack(anchor="center", padx=100)

            for product, count in medium_usage.items():
                tk.Label(frame_medium, text=f"{product}: {count} utilisations").pack(anchor="center", padx=100)

            for product, count in low_usage.items():
                tk.Label(frame_low, text=f"{product}: {count} utilisations").pack(anchor="center", padx=100)

            fig, axs = plt.subplots(1, 2, figsize=(15, 6))

            def update(frame):
                axs[0].cla()
                axs[1].cla()

                axs[0].pie(high_usage, labels=high_usage.index, autopct='%6.1f%%', startangle=100, colors=plt.cm.Paired.colors)
                axs[0].set_title("Top 5 des Produits les Plus Utilisés")

                axs[1].bar(top_clients.index.astype(str), top_clients.values, color='lightblue')
                axs[1].set_title("Top 5 des Clients les Plus Fidèles")
                axs[1].set_ylabel("Nombre de Souscription")
                axs[1].tick_params(axis='x', rotation=20)
                axs[0].tick_params(axis='y', rotation=20)
                plt.tight_layout()

            ani = FuncAnimation(fig, update, interval=1000, cache_frame_data=False)

            canvas = FigureCanvasTkAgg(fig, top)
            canvas.draw()
            canvas.get_tk_widget().pack()

        except Exception as e:
            messagebox.showerror("Erreur", f"Une erreur s'est produite : {e}")

    def show_about(self):
        about_window = tk.Toplevel(self.root)
        about_window.title("À propos")
        about_window.geometry("400x300")
        about_window.resizable(False, False)
        
        # Appliquer le thème à la fenêtre about
        about_window.config(bg=self.theme['fg'])
        
        tk.Label(about_window, text="Tableau de bord de gestion", 
                font=("Helvetica", 16, "bold"), 
                bg=self.theme['fg'], fg=self.theme['bg']).pack(pady=20)
                
        tk.Label(about_window, text="Version Finale", 
                font=("Helvetica", 12), 
                bg=self.theme['fg'], fg=self.theme['bg']).pack()
        
        tk.Label(about_window, text="Réalisé par Ingénieur en Génie Numérique Frédéric BAYONNE MAVOUNGOU", 
                font=("Helvetica", 12), 
                bg=self.theme['fg'], fg=self.theme['bg']).pack()
                
        tk.Label(about_window, text="© 2025 AGC-VIE", 
                font=("Helvetica", 10), 
                bg=self.theme['fg'], fg=self.theme['bg']).pack(pady=20)
                
        close_btn = RoundedButton(about_window, "Fermer", about_window.destroy,
                                    width=100, height=30,
                                    bg=self.theme['button_bg'],
                                    fg=self.theme['button_fg'],
                                    hover_bg=self.theme['button_hover'])
        close_btn.pack(pady=20)

# Fonction pour afficher les données dans le tableau (Treeview)
def afficher_donnees(df, table_view):
    # Supprimer les anciennes données
    table_view.delete(*table_view.get_children())

    # Ajouter les colonnes et les données
    table_view["columns"] = list(df.columns)
    table_view["show"] = "headings"  # N'affiche pas la colonne de l'index

    # Configuration des en-têtes de colonnes
    for col in df.columns:
        table_view.heading(col, text=col)
        table_view.column(col, width=150)

    for index, row in df.iterrows():
        tags = ()
        if 'Écart' in row and row['Écart'] > 0:
            tags = ('error',)
        table_view.insert("", "end", values=list(row), tags=tags)
    
    table_view.tag_configure('error', background='#ffdddd')  # Fond rouge clair pour les écarts

# ------------------------ MAIN ENTRY POINT ------------------------
#if __name__ == "__main__":
    #root = tk.Tk()
    #root.withdraw()
    
    # Appliquer le thème Sun Valley
    #set_theme("light")  # "light" ou "dark"
    
    # Start with first run setup
    #FirstRunSetup(root)
    
    #root.mainloop()

# Création de l'application
if __name__ == "__main__":
    root = tk.Tk()
    app = DashboardApp(root)
    root.mainloop()